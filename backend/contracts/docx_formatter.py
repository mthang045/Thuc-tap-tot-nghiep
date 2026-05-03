"""
DOCX Formatter for Vietnamese Legal Documents
Tạo file Word có căn chỉnh đàng hoàng từ text template
"""

import io
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None


# ============== Document Formatting Constants ==============

# Page margins (in inches)
PAGE_MARGIN_TOP = 1.2
PAGE_MARGIN_BOTTOM = 1.0
PAGE_MARGIN_LEFT = 1.5
PAGE_MARGIN_RIGHT = 1.0

# Font settings
DEFAULT_FONT_NAME = "Times New Roman"
DEFAULT_FONT_SIZE = 13  # 13pt for body text
HEADING_FONT_SIZE = 14

# Line spacing
DEFAULT_LINE_SPACING = 1.5  # 1.5 line spacing
SINGLE_LINE_SPACING = 1.0


def set_cell_shading(cell, color_hex):
    """Set background color for table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()


def create_heading_paragraph(doc, text, level=1, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_before=0, space_after=200):
    """Create a formatted heading paragraph"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(DEFAULT_LINE_SPACING * DEFAULT_FONT_SIZE)
    
    run = p.add_run(text)
    run.bold = bold
    run.font.name = DEFAULT_FONT_NAME
    run.font.size = Pt(HEADING_FONT_SIZE if level == 1 else 13)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)
    
    return p


def create_body_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, space_before=0, space_after=0, first_line_indent=0.5):
    """Create a formatted body paragraph with proper Vietnamese formatting"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(DEFAULT_LINE_SPACING * DEFAULT_FONT_SIZE)
    
    if first_line_indent > 0:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = DEFAULT_FONT_NAME
    run.font.size = Pt(DEFAULT_FONT_SIZE)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)
    
    return p


def create_empty_paragraph(doc, space_after=100):
    """Create an empty paragraph for spacing"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(DEFAULT_LINE_SPACING * DEFAULT_FONT_SIZE)
    return p


def format_legal_document(text_content, doc_title):
    """
    Format raw text content into a properly structured Vietnamese legal document.
    
    Returns:
        BytesIO buffer containing the DOCX file
    """
    if Document is None:
        raise RuntimeError('python-docx is not installed')
    
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(PAGE_MARGIN_TOP)
        section.bottom_margin = Inches(PAGE_MARGIN_BOTTOM)
        section.left_margin = Inches(PAGE_MARGIN_LEFT)
        section.right_margin = Inches(PAGE_MARGIN_RIGHT)
    
    # Set default paragraph style
    style = doc.styles['Normal']
    style.font.name = DEFAULT_FONT_NAME
    style.font.size = Pt(DEFAULT_FONT_SIZE)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)
    
    # Parse and format the content
    lines = text_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines at the start
        if not stripped and i == 0:
            i += 1
            continue
        
        # Handle special Vietnamese headers
        if _is_vietnamese_header(stripped):
            _format_vietnamese_header(doc, stripped)
        # Handle article/điều headers
        elif _is_article_header(stripped):
            _format_article_header(doc, stripped, lines, i)
        # Handle signature lines
        elif _is_signature_line(stripped):
            _format_signature_block(doc, stripped, lines, i)
            # Skip the next few lines as they're part of signature
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines) and 'ký' in lines[i].lower():
                i += 1
        # Handle regular paragraphs
        elif stripped:
            _format_body_paragraph(doc, stripped)
        # Empty line
        else:
            create_empty_paragraph(doc, space_after=60)
        
        i += 1
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _is_vietnamese_header(text):
    """Check if text is a Vietnamese legal document header"""
    header_keywords = [
        'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM',
        'Độc lập – Tự do – Hạnh phúc',
        'HỘI ĐỒNG QUẢN TRỊ',
        'CÔNG TY',
    ]
    upper_text = text.upper()
    return any(kw in upper_text for kw in header_keywords)


def _is_article_header(text):
    """Check if text is an article header (Điều X: ... )"""
    return bool(re.match(r'^Điều\s*\d+[:.]?', text, re.IGNORECASE))


def _is_signature_line(text):
    """Check if text is a signature line"""
    sig_keywords = ['NGƯỜI LAO ĐỘNG', 'NGƯỜI SỬ DỤNG', 'CÁN BỘ', 'GIÁM ĐỐC', 'KẾ TOÁN']
    upper = text.upper()
    return any(kw in upper for kw in sig_keywords) and ('ký' in upper.lower() or len(text) < 60)


def _format_vietnamese_header(doc, text):
    """Format Vietnamese legal document header"""
    # Determine alignment
    if 'CỘNG HÒA' in text.upper() or 'Độc lập' in text:
        alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Determine if bold
    bold = 'CỘNG HÒA' in text.upper() or 'Điều' in text[:10]
    
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(80)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    run = p.add_run(text)
    run.bold = bold
    run.font.name = DEFAULT_FONT_NAME
    run.font.size = Pt(13)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)


def _format_article_header(doc, text, lines, current_idx):
    """Format article header with its content"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(200)
    p.paragraph_format.space_after = Pt(100)
    p.paragraph_format.line_spacing = Pt(DEFAULT_LINE_SPACING * DEFAULT_FONT_SIZE)
    p.paragraph_format.first_line_indent = Inches(0)
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = DEFAULT_FONT_NAME
    run.font.size = Pt(DEFAULT_FONT_SIZE)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)
    
    # Check if next lines are sub-items
    j = current_idx + 1
    while j < len(lines):
        next_line = lines[j].strip()
        if not next_line:
            j += 1
            continue
        
        # Stop if we hit another article or major header
        if _is_article_header(next_line) or _is_vietnamese_header(next_line):
            break
        
        # Check if it's a numbered sub-item
        if re.match(r'^\d+\.\d*', next_line) or re.match(r'^\([a-z0-9]+\)', next_line, re.IGNORECASE):
            _format_sub_item(doc, next_line)
        elif next_line.startswith('- '):
            _format_bullet_item(doc, next_line[2:])
        else:
            _format_body_paragraph(doc, next_line)
        
        j += 1


def _format_sub_item(doc, text):
    """Format a numbered sub-item"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(60)
    p.paragraph_format.line_spacing = Pt(DEFAULT_LINE_SPACING * DEFAULT_FONT_SIZE)
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    run = p.add_run(text)
    run.font.name = DEFAULT_FONT_NAME
    run.font.size = Pt(DEFAULT_FONT_SIZE)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)


def _format_bullet_item(doc, text):
    """Format a bullet point item"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(60)
    p.paragraph_format.line_spacing = Pt(DEFAULT_LINE_SPACING * DEFAULT_FONT_SIZE)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.left_indent = Inches(0.5)
    
    run = p.add_run(f"• {text}")
    run.font.name = DEFAULT_FONT_NAME
    run.font.size = Pt(DEFAULT_FONT_SIZE)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)


def _format_body_paragraph(doc, text):
    """Format a regular body paragraph"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(DEFAULT_LINE_SPACING * DEFAULT_FONT_SIZE)
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    run = p.add_run(text)
    run.font.name = DEFAULT_FONT_NAME
    run.font.size = Pt(DEFAULT_FONT_SIZE)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)


def _format_signature_block(doc, text, lines, current_idx):
    """Format signature block with table for alignment"""
    # Create a 2-column table for signatures
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for cell in table.rows[0].cells:
        cell.width = Inches(3.0)
    
    # Left cell - Employee/Worker
    left_cell = table.rows[0].cells[0]
    left_cell.text = ""
    p = left_cell.paragraphs[0]
    run = p.add_run(text if 'LAO ĐỘNG' in text.upper() else "")
    run.font.name = DEFAULT_FONT_NAME
    run.font.size = Pt(12)
    
    # Right cell - Employer
    right_cell = table.rows[0].cells[1]
    right_cell.text = ""
    
    # Look for employer signature in next lines
    j = current_idx + 1
    employer_text = ""
    while j < len(lines) and j < current_idx + 3:
        line = lines[j].strip()
        if 'SỬ DỤNG' in line.upper() or 'CÔNG TY' in line.upper() or 'GIÁM ĐỐC' in line.upper():
            employer_text = line
            break
        j += 1
    
    p2 = right_cell.paragraphs[0]
    run2 = p2.add_run(employer_text)
    run2.font.name = DEFAULT_FONT_NAME
    run2.font.size = Pt(12)
    
    # Add spacing
    create_empty_paragraph(doc, space_after=100)
    
    # Add signature lines
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    left_sig = sig_para.add_run("(Ký, ghi rõ họ tên)")
    left_sig.font.name = DEFAULT_FONT_NAME
    left_sig.font.size = Pt(11)
    left_sig.italic = True
    
    sig_para.add_run("                    ")
    
    right_sig = sig_para.add_run("(Ký, ghi rõ họ tên, đóng dấu)")
    right_sig.font.name = DEFAULT_FONT_NAME
    right_sig.font.size = Pt(11)
    right_sig.italic = True


def create_simple_document(text_content, title):
    """
    Simple formatting for basic documents like Đơn xin nghỉ việc.
    """
    if Document is None:
        raise RuntimeError('python-docx is not installed')
    
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = DEFAULT_FONT_NAME
    style.font.size = Pt(DEFAULT_FONT_SIZE)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)
    
    # Process lines
    lines = text_content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        
        # Skip multiple empty lines
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(60)
            continue
        
        # Headers (centered)
        if stripped.isupper() and len(stripped) < 50:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(100)
            p.paragraph_format.space_after = Pt(100)
            run = p.add_run(stripped)
            run.bold = True
            run.font.name = DEFAULT_FONT_NAME
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)
        # Signature lines
        elif '(Ký' in stripped or 'ký' in stripped.lower():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(200)
            run = p.add_run(stripped)
            run.font.name = DEFAULT_FONT_NAME
            run.font.size = Pt(12)
            run.italic = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(60)
            p.paragraph_format.space_after = Pt(60)
            p.paragraph_format.line_spacing = Pt(1.5 * DEFAULT_FONT_SIZE)
            p.paragraph_format.first_line_indent = Inches(0.5)
            run = p.add_run(stripped)
            run.font.name = DEFAULT_FONT_NAME
            run.font.size = Pt(DEFAULT_FONT_SIZE)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def detect_document_type(text_content):
    """Detect the type of document based on content"""
    text_upper = text_content.upper()
    
    if 'ĐƠN XIN' in text_upper:
        return 'don'
    elif 'QUYẾT ĐỊNH' in text_upper:
        return 'quyet_dinh'
    elif 'HỢP ĐỒNG' in text_upper:
        return 'hop_dong'
    elif 'BIÊN BẢN' in text_upper:
        return 'bien_ban'
    elif 'THỎA THUẬN' in text_upper:
        return 'thoa_thuan'
    elif 'GIẤY ỦY QUYỀN' in text_upper:
        return 'giay_uy_quyen'
    else:
        return 'generic'


def format_document(text_content, title):
    """
    Auto-detect document type and apply appropriate formatting.
    """
    doc_type = detect_document_type(text_content)
    
    if doc_type in ['don', 'giay_uy_quyen']:
        # Simple documents - just clean formatting
        return create_simple_document(text_content, title)
    else:
        # Complex legal documents - full formatting
        return format_legal_document(text_content, title)
