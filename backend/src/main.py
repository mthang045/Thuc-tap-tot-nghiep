"""
Legal AI Agent API
- Full-text search Vietnamese law database
- Claude OAuth for AI processing
- Multi-tenant API key authentication
- User authentication and management
"""
from fastapi import FastAPI, HTTPException, Depends, Header, Request, Query, UploadFile, File, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.staticfiles import StaticFiles
import pathlib
from pydantic import BaseModel, Field
from typing import Optional, List, AsyncGenerator
import psycopg2
from psycopg2.extras import RealDictCursor
import httpx
import json
import hashlib
import time
import re as re_module
import os
import uuid
import tempfile
from io import BytesIO
from datetime import datetime
from contextlib import contextmanager
from collections import defaultdict
import jwt
from dotenv import load_dotenv

# Optional imports with fallback
try:
    import uvicorn  # type: ignore
except ImportError:
    uvicorn = None

try:
    from pypdf import PdfReader  # type: ignore
except ImportError:
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except ImportError:
        PdfReader = None

# ============================================
# Rate Limiter (in-memory)
# ============================================
rate_limits: dict = defaultdict(list)

def check_rate_limit(company_id: str, max_requests: int = 20, window_seconds: int = 60):
    """Check if company has exceeded rate limit"""
    now = time.time()
    rate_limits[company_id] = [t for t in rate_limits[company_id] if now - t < window_seconds]
    if len(rate_limits[company_id]) >= max_requests:
        raise HTTPException(status_code=429, detail=f"Rate limit exceeded. Max {max_requests} requests per minute.")
    rate_limits[company_id].append(now)

# ============================================
# Search Cache (in-memory, TTL 1 hour)
# ============================================
search_cache: dict = {}
CACHE_TTL = 3600

# Load environment variables from .env file
load_dotenv()

# FIX 1: Import and use validated JWT secret
from .security_utils import validate_jwt_secret, sanitize_log, rate_limiter as global_rate_limiter

JWT_SECRET = validate_jwt_secret()

# Import new routes
from .routes import auth, company, keys, usage, chats, documents, admin, contracts, templates, crawler, llm_oauth, pricing, platform_admin
# from .middleware.logging import PlatformLoggingMiddleware  # disabled for deploy

# Import agent (initialized after DB functions are defined)
from .agents import legal_agent

# Import LLM Provider Manager
from .services.llm_provider import LLMProviderManager

app = FastAPI(
    title="Legal AI Agent API",
    description="AI-powered Vietnamese Legal Assistant - Tư vấn pháp luật, soạn thảo văn bản, rà soát hợp đồng",
    version="2.0.0"
)

# Logging middleware disabled for production deploy
# app.add_middleware(
#     PlatformLoggingMiddleware,
#     exclude_paths=["/health", "/docs", "/openapi.json", "/redoc", "/static", "/"]
# )

# FIX 11: CORS - Restrict origins
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:8080").split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "PATCH"],
    allow_headers=["*"],
)

# FIX 13: Security headers middleware
@app.middleware("http")
async def security_headers(request: Request, call_next):
    """Add security headers to all responses"""
    response = await call_next(request)
    
    # HSTS - Force HTTPS (only if running on HTTPS)
    response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    
    # Prevent MIME sniffing
    response.headers["X-Content-Type-Options"] = "nosniff"
    
    # Prevent clickjacking
    response.headers["X-Frame-Options"] = "DENY"
    
    # XSS Protection (legacy but harmless)
    response.headers["X-XSS-Protection"] = "1; mode=block"
    
    # CSP (Content Security Policy) - Basic policy
    response.headers["Content-Security-Policy"] = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' 'unsafe-eval'; "
        "style-src 'self' 'unsafe-inline'; "
        "img-src 'self' data: https:; "
        "font-src 'self' data:; "
        "connect-src 'self' https://api.anthropic.com"
    )
    
    return response

@app.middleware("http")
async def log_requests(request: Request, call_next):
    """Log slow requests - FIX 7: Sanitize logs"""
    start = time.time()
    response = await call_next(request)
    duration = time.time() - start
    if duration > 5:  # Log slow requests
        # FIX 7: Sanitize URL and headers before logging
        sanitized_path = sanitize_log(str(request.url.path))
        print(f"SLOW REQUEST: {request.method} {sanitized_path} took {duration:.2f}s")
    return response

# Include new routers
app.include_router(auth.router)
app.include_router(company.router)
app.include_router(keys.router)
app.include_router(usage.router)
app.include_router(chats.router)
app.include_router(documents.router)
app.include_router(admin.router)
app.include_router(contracts.router)
app.include_router(templates.router)
app.include_router(crawler.router)
app.include_router(llm_oauth.router)
app.include_router(pricing.router)
app.include_router(platform_admin.router)

# Startup event - seed templates
@app.on_event("startup")
async def startup_event():
    """Seed default templates and ensure audit table on startup"""
    templates.seed_default_templates()
    ensure_audit_table()

# Static files
static_dir = pathlib.Path(__file__).parent.parent.parent / "static"
if static_dir.exists():
    app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")

@app.middleware("http")
async def no_cache_static(request, call_next):
    response = await call_next(request)
    if request.url.path.startswith("/static/"):
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
    return response

@app.get("/", include_in_schema=False)
async def landing_page():
    html_file = static_dir / "index.html"
    if html_file.exists():
        return FileResponse(str(html_file))
    return {"name": "Legal AI Agent API", "version": "1.0.0"}

@app.get("/app", include_in_schema=False)
async def app_page():
    """Main App - always serve fresh"""
    html_file = static_dir / "app.html"
    if html_file.exists():
        from starlette.responses import Response
        content = html_file.read_text()
        return Response(content=content, media_type="text/html", headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0"
        })
    return {"error": "app.html not found"}

@app.get("/platform-admin", include_in_schema=False)
async def platform_admin_page():
    """Platform Super Admin Panel (superadmin only)"""
    html_file = static_dir / "platform-admin.html"
    if html_file.exists():
        return FileResponse(str(html_file))
    raise HTTPException(status_code=404, detail="Platform admin page not found")

# ============================================
# Database
# ============================================

DB_CONFIG = {
    "host": os.getenv("SUPABASE_DB_HOST", "localhost"),
    "port": int(os.getenv("SUPABASE_DB_PORT", "5432")),
    "dbname": os.getenv("DB_NAME", "postgres"),
    "user": os.getenv("DB_USER", "postgres"),
    "password": os.getenv("SUPABASE_DB_PASSWORD", ""),
    "sslmode": os.getenv("DB_SSL_MODE", "require"),
}

@contextmanager
def get_db():
    conn = psycopg2.connect(**DB_CONFIG)
    try:
        yield conn
    finally:
        conn.close()

# ============================================
# Audit Log
# ============================================

def log_audit(company_id: str, user_id: str, action: str, resource_type: str, resource_id: str = None, details: dict = None):
    """Log an audit event"""
    try:
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO audit_logs (company_id, user_id, action, resource_type, resource_id, details, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, NOW())
            """, (company_id, user_id, action, resource_type, resource_id, json.dumps(details or {})))
            conn.commit()
    except Exception as e:
        print(f"Audit log error: {e}")

def ensure_audit_table():
    try:
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute("""
                CREATE TABLE IF NOT EXISTS audit_logs (
                    id SERIAL PRIMARY KEY,
                    company_id UUID,
                    user_id UUID,
                    action VARCHAR(50) NOT NULL,
                    resource_type VARCHAR(50),
                    resource_id VARCHAR(100),
                    details JSONB DEFAULT '{}',
                    created_at TIMESTAMPTZ DEFAULT NOW()
                )
            """)
            cur.execute("CREATE INDEX IF NOT EXISTS idx_audit_company ON audit_logs(company_id, created_at DESC)")
            conn.commit()
    except Exception as e:
        print(f"Audit table error: {e}")

# ============================================
# Auth
# ============================================

async def verify_api_key(
    x_api_key: Optional[str] = Header(None, alias="X-API-Key"),
    authorization: Optional[str] = Header(None)
):
    """Verify API key OR Bearer token and return company info"""
    
    # Try Bearer token first (from dashboard login)
    if not x_api_key and authorization and authorization.startswith("Bearer "):
        try:
            token = authorization.split(" ", 1)[1]
            payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
            user_id = payload.get("user_id")
            if user_id:
                with get_db() as conn:
                    cur = conn.cursor(cursor_factory=RealDictCursor)
                    cur.execute("""
                        SELECT u.id as user_id, u.company_id, u.role,
                               c.name as company_name, c.plan, c.monthly_quota, c.used_quota
                        FROM users u
                        JOIN companies c ON c.id = u.company_id
                        WHERE u.id = %s
                    """, (user_id,))
                    user = cur.fetchone()
                    if user:
                        if user["used_quota"] >= user["monthly_quota"]:
                            raise HTTPException(status_code=429, detail="Monthly quota exceeded")
                        return {**dict(user), "permissions": ["read","ask","review","draft"], "rate_limit": 60}
        except Exception:
            pass
    
    if not x_api_key:
        raise HTTPException(status_code=401, detail="API key or Bearer token required")
    
    key_hash = hashlib.sha256(x_api_key.encode()).hexdigest()
    key_prefix = x_api_key[:8]
    
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("""
            SELECT ak.id, ak.company_id, ak.permissions, ak.rate_limit,
                   c.name as company_name, c.plan, c.monthly_quota, c.used_quota
            FROM api_keys ak
            JOIN companies c ON c.id = ak.company_id
            WHERE ak.key_prefix = %s AND ak.key_hash = %s AND ak.is_active = true
        """, (key_prefix, key_hash))
        result = cur.fetchone()
        
        if not result:
            raise HTTPException(status_code=401, detail="Invalid API key")
        
        if result["used_quota"] >= result["monthly_quota"]:
            raise HTTPException(status_code=429, detail="Monthly quota exceeded")
        
        # Update last_used
        cur.execute("UPDATE api_keys SET last_used_at = now() WHERE id = %s", (result["id"],))
        conn.commit()
        
        return dict(result)

# ============================================
# Models
# ============================================

class FileContext(BaseModel):
    filename: str
    content: str
    file_type: str

class LegalQuery(BaseModel):
    question: str = Field(..., min_length=1, max_length=5000, description="Câu hỏi pháp luật")
    domains: Optional[List[str]] = Field(None, description="Lĩnh vực: lao_dong, doanh_nghiep, dan_su, thue, dat_dai...")
    max_sources: int = Field(10, ge=1, le=30, description="Số nguồn tham chiếu tối đa")
    stream: bool = Field(False, description="Stream response")
    session_id: Optional[str] = Field(None, description="Chat session ID for multi-turn conversation")
    file_context: Optional[FileContext] = Field(None, description="File content uploaded in chat")

class ContractReview(BaseModel):
    contract_text: str = Field(..., min_length=50, max_length=100000, description="Nội dung hợp đồng cần rà soát")
    contract_type: Optional[str] = Field(None, description="Loại hợp đồng: hop_dong_lao_dong, hop_dong_thuong_mai...")
    focus_areas: Optional[List[str]] = Field(None, description="Các điểm cần chú ý đặc biệt")

class DocumentDraft(BaseModel):
    doc_type: str = Field(..., description="Loại văn bản: hop_dong_lao_dong, quyet_dinh, cong_van, noi_quy...")
    variables: dict = Field(..., description="Thông tin cần điền vào văn bản")
    instructions: Optional[str] = Field(None, description="Yêu cầu bổ sung")

class LegalResponse(BaseModel):
    answer: str
    citations: List[dict]
    confidence: float
    tokens_used: int
    model: str
    session_id: Optional[str] = None

# ============================================
# Claude OAuth Integration
# ============================================

CLAUDE_OAUTH_TOKEN = os.getenv("CLAUDE_OAUTH_TOKEN", "")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
CLAUDE_API_URL = "https://api.anthropic.com/v1/messages"

async def call_claude(system_prompt: str, user_message: str, max_tokens: int = 4096, history: list = None) -> dict:
    """Call Claude via OAuth token or API key, with optional conversation history"""
    api_key = ANTHROPIC_API_KEY
    oauth_token = CLAUDE_OAUTH_TOKEN
    
    headers = {
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }
    
    if oauth_token:
        headers["Authorization"] = f"Bearer {oauth_token}"
        headers["anthropic-beta"] = "oauth-2025-04-20"
    elif api_key:
        headers["x-api-key"] = api_key
    else:
        raise ValueError("No Claude API key or OAuth token configured")
    
    # Build messages with history for multi-turn conversations
    messages = []
    if history:
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": user_message})
    
    payload = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": max_tokens,
        "system": system_prompt,
        "messages": messages
    }
    
    async with httpx.AsyncClient(timeout=120) as client:
        try:
            response = await client.post(CLAUDE_API_URL, headers=headers, json=payload)
            response.raise_for_status()
            data = response.json()
            
            return {
                "content": data["content"][0]["text"],
                "input_tokens": data["usage"]["input_tokens"],
                "output_tokens": data["usage"]["output_tokens"],
                "model": data["model"]
            }
        except httpx.TimeoutException:
            print(f"Claude API timeout")
            raise HTTPException(status_code=504, detail="Hệ thống đang bận, vui lòng thử lại")
        except httpx.HTTPStatusError as e:
            print(f"Claude API error: {e.response.status_code} - {e.response.text[:200]}")
            if e.response.status_code == 429:
                raise HTTPException(status_code=429, detail="Đã vượt giới hạn API, vui lòng chờ")
            elif e.response.status_code == 529:
                raise HTTPException(status_code=503, detail="Hệ thống đang bận, vui lòng thử lại")
            raise HTTPException(status_code=502, detail=f"Lỗi AI engine: {e.response.status_code}")
        except psycopg2.Error as e:
            print(f"Database error in Claude call: {e}")
            raise HTTPException(status_code=500, detail="Lỗi kết nối database")
        except ValueError as e:
            print(f"Claude config error: {e}")
            raise HTTPException(status_code=500, detail=str(e))
        except Exception as e:
            print(f"Claude call error: {e}")
            raise HTTPException(status_code=500, detail="Hệ thống đang bận, vui lòng thử lại")


async def call_claude_stream(system_prompt: str, user_message: str, max_tokens: int = 8192, history: list = None) -> AsyncGenerator[str, None]:
    """Call Claude with streaming via SSE, yielding text deltas"""
    api_key = ANTHROPIC_API_KEY
    oauth_token = CLAUDE_OAUTH_TOKEN

    headers = {
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }

    if oauth_token:
        headers["Authorization"] = f"Bearer {oauth_token}"
        headers["anthropic-beta"] = "oauth-2025-04-20"
    elif api_key:
        headers["x-api-key"] = api_key
    else:
        raise ValueError("No Claude API key or OAuth token configured")

    messages = []
    if history:
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": user_message})

    payload = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": max_tokens,
        "system": system_prompt,
        "messages": messages,
        "stream": True
    }

    async with httpx.AsyncClient(timeout=180) as client:
        async with client.stream("POST", CLAUDE_API_URL, headers=headers, json=payload) as response:
            response.raise_for_status()
            async for line in response.aiter_lines():
                if line.startswith("data: "):
                    data_str = line[6:]
                    if data_str.strip() == "[DONE]":
                        break
                    try:
                        event = json.loads(data_str)
                        event_type = event.get("type", "")
                        if event_type == "content_block_delta":
                            delta = event.get("delta", {})
                            if delta.get("type") == "text_delta":
                                yield delta.get("text", "")
                        elif event_type == "message_stop":
                            break
                        elif event_type == "message_delta":
                            # Contains usage info at the end
                            pass
                    except json.JSONDecodeError:
                        continue


# ============================================
# Context Enrichment - Company Documents & Contracts
# ============================================

def fetch_company_context(company_id: str, question: str, limit: int = 5) -> str:
    """Search company's uploaded documents and contracts for relevant context"""
    context_parts = []
    question_lower = question.lower()

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Search company documents with extracted text
        cur.execute("""
            SELECT name, extracted_text, doc_type, analysis
            FROM documents
            WHERE company_id = %s
              AND extracted_text IS NOT NULL
              AND length(extracted_text) > 50
            ORDER BY created_at DESC
            LIMIT %s
        """, (company_id, limit))

        docs = cur.fetchall()
        for doc in docs:
            text = doc["extracted_text"] or ""
            # Check if document is relevant to the question
            if any(kw in text.lower() for kw in question_lower.split() if len(kw) > 2):
                excerpt = text[:1500]
                context_parts.append(
                    f"📄 TÀI LIỆU CÔNG TY: {doc['name']} (Loại: {doc.get('doc_type', 'N/A')})\n{excerpt}"
                )

        # Search company contracts
        cur.execute("""
            SELECT name, contract_type, extracted_text, parties, start_date, end_date, notes
            FROM contracts
            WHERE company_id = %s
              AND status != 'deleted'
              AND extracted_text IS NOT NULL
              AND length(extracted_text) > 50
            ORDER BY created_at DESC
            LIMIT %s
        """, (company_id, limit))

        contracts = cur.fetchall()
        for contract in contracts:
            text = contract["extracted_text"] or ""
            if any(kw in text.lower() for kw in question_lower.split() if len(kw) > 2):
                excerpt = text[:1500]
                parties_str = ""
                if contract.get("parties"):
                    try:
                        parties = json.loads(contract["parties"]) if isinstance(contract["parties"], str) else contract["parties"]
                        parties_str = f" | Các bên: {', '.join(str(p) for p in parties)}"
                    except:
                        pass
                context_parts.append(
                    f"📋 HỢP ĐỒNG: {contract['name']} (Loại: {contract.get('contract_type', 'N/A')}{parties_str})\n{excerpt}"
                )

    return "\n\n".join(context_parts[:5]) if context_parts else ""


# ============================================
# Law Search
# ============================================

# ============================================
# Vietnamese Diacritics Restoration
# ============================================

NO_ACCENT_MAP = {
    # Common legal phrases
    "thu viec": "thử việc",
    "nghi phep": "nghỉ phép",
    "hop dong lao dong": "hợp đồng lao động",
    "thue tndn": "thuế TNDN",
    "thue thu nhap": "thuế thu nhập",
    "sa thai": "sa thải",
    "luong": "lương",
    "bao hiem": "bảo hiểm",
    "nghi viec": "nghỉ việc",
    "ky luat": "kỷ luật",
    "thai san": "thai sản",
    "tang ca": "tăng ca",
    "lam them gio": "làm thêm giờ",
    "nghi le": "nghỉ lễ",
    "cham dut hop dong": "chấm dứt hợp đồng",
    "boi thuong": "bồi thường",
    "tranh chap": "tranh chấp",
    "thanh lap cong ty": "thành lập công ty",
    "doanh nghiep": "doanh nghiệp",
    "co phan": "cổ phần",
    "thue gia tri gia tang": "thuế giá trị gia tăng",
    "dat dai": "đất đai",
    "quyen su dung dat": "quyền sử dụng đất",
    # Common single words
    "thoi gian": "thời gian",
    "toi da": "tối đa",
    "toi thieu": "tối thiểu",
    "quy dinh": "quy định",
    "noi dung": "nội dung",
    "hinh thuc": "hình thức",
    "hop dong": "hợp đồng",
    "cong ty": "công ty",
    "dieu": "điều",
    "khoan": "khoản",
    "luat": "luật",
    "bo luat": "bộ luật",
    "nghi dinh": "nghị định",
    "thong tu": "thông tư",
    "quyet dinh": "quyết định",
    "muc": "mức",
    "so": "số",
    "nam": "năm",
    "thang": "tháng",
    "ngay": "ngày",
    "gio": "giờ",
    "viec": "việc",
    "nguoi": "người",
    "phep": "phép",
    "thue": "thuế",
    "suat": "suất",
    "tien": "tiền",
}

def has_vietnamese_diacritics(text: str) -> bool:
    """Check if text contains Vietnamese diacritics"""
    import re
    # Vietnamese diacritics pattern
    vietnamese_chars = r'[àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ]'
    return bool(re.search(vietnamese_chars, text.lower()))

def restore_diacritics(query: str) -> str:
    """Restore Vietnamese diacritics from common non-diacritics legal terms"""
    import re
    
    # If query already has diacritics, return as-is
    if has_vietnamese_diacritics(query):
        return query
    
    # Try to match and replace phrases from NO_ACCENT_MAP
    restored = query.lower()
    
    # Sort by length (longest first) to match longer phrases first
    sorted_mappings = sorted(NO_ACCENT_MAP.items(), key=lambda x: len(x[0]), reverse=True)
    
    for no_accent, with_accent in sorted_mappings:
        # Use word boundaries to avoid partial matches
        pattern = r'\b' + re.escape(no_accent) + r'\b'
        restored = re.sub(pattern, with_accent, restored, flags=re.IGNORECASE)
    
    return restored

def extract_search_query(question: str) -> str:
    """Extract key legal terms from Vietnamese question"""
    import re
    
    # Remove Vietnamese question words
    question_words = [
        r'\bbao lâu\b', r'\bbao nhiêu\b', r'\bthế nào\b', r'\bnhư thế nào\b',
        r'\blà gì\b', r'\bcó phải\b', r'\bcó được\b', r'\blà\b', r'\bcó\b',
        r'\bkhông\b', r'\bhay không\b', r'\?', r'\.'
    ]
    
    cleaned = question.lower()
    for pattern in question_words:
        cleaned = re.sub(pattern, ' ', cleaned)
    
    # Remove extra spaces
    cleaned = ' '.join(cleaned.split())
    
    return cleaned.strip()

def expand_synonyms(query: str) -> List[str]:
    """Expand Vietnamese legal term synonyms — returns the EXPANDED TERMS to search for"""
    expansions = []
    query_lower = query.lower()
    
    synonym_map = {
        "tndn": "thu nhập doanh nghiệp",
        "tncn": "thu nhập cá nhân",
        "gtgt": "giá trị gia tăng",
        "vat": "giá trị gia tăng",
        "bhxh": "bảo hiểm xã hội",
        "bhyt": "bảo hiểm y tế",
        "hđlđ": "hợp đồng lao động",
        "nsdlđ": "người sử dụng lao động",
        "nlđ": "người lao động",
        "nghỉ phép": "nghỉ hằng năm",
        "phép năm": "nghỉ hằng năm",
        "sa thải": "kỷ luật sa thải",
        "đuổi việc": "kỷ luật sa thải",
    }
    
    for abbr, full in synonym_map.items():
        if abbr in query_lower:
            expansions.append(full)  # Return just the expanded term
    
    return expansions

def detect_domain(question: str) -> Optional[List[str]]:
    """Auto-detect legal domain from question keywords"""
    question_lower = question.lower()
    
    domain_keywords = {
        "lao_dong": ["lao động", "hợp đồng lao động", "thử việc", "nghỉ phép", "tăng ca", "lương", "sa thải", "bảo hiểm xã hội", "bhxh", "thôi việc", "chấm dứt hợp đồng"],
        "thue": ["thuế", "tndn", "vat", "tncn", "kê khai thuế", "hoàn thuế", "miễn thuế", "giảm thuế", "thuế suất"],
        "doanh_nghiep": ["thành lập công ty", "cổ phần", "doanh nghiệp", "giải thể", "phá sản", "điều lệ", "đại hội cổ đông", "hội đồng quản trị"],
        "dan_su": ["di sản", "thừa kế", "hôn nhân", "ly hôn", "nuôi con", "nhà ở", "quyền sở hữu", "tài sản chung"],
        "dat_dai": ["đất đai", "quyền sử dụng đất", "sổ đỏ", "chuyển nhượng đất", "thuê đất"],
        "hinh_su": ["hình sự", "án tù", "tội phạm", "vi phạm hình sự", "truy tố"],
        "hanh_chinh": ["vi phạm hành chính", "phạt hành chính", "khiếu nại", "tố cáo"]
    }
    
    detected = []
    for domain, keywords in domain_keywords.items():
        for keyword in keywords:
            if keyword in question_lower:
                detected.append(domain)
                break
    
    return detected if detected else None

def search_laws(query: str, domains: Optional[List[str]] = None, limit: int = 10) -> List[dict]:
    """Search Vietnamese law database"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        if domains:
            domain_array = "{" + ",".join(domains) + "}"
            cur.execute(
                "SELECT * FROM search_law(%s, %s::legal_domain[], %s)",
                (query, domain_array, limit)
            )
        else:
            cur.execute(
                "SELECT * FROM search_law(%s, NULL, %s)",
                (query, limit)
            )
        
        return [dict(r) for r in cur.fetchall()]

def multi_query_search(question: str, domains: Optional[List[str]] = None, limit: int = 15) -> List[dict]:
    """Smart multi-query search: domain detection + ILIKE phrase + tsvector + diacritics fallback"""
    
    # Auto-detect domain
    if not domains:
        domains = detect_domain(question)
    
    # Check if query needs diacritics restoration
    original_query = question
    restored_query = restore_diacritics(question)
    queries_to_search = [original_query]
    
    # If restoration produced a different query, search with both
    if restored_query.lower() != original_query.lower():
        queries_to_search.append(restored_query)
    
    # Extract clean keywords from the best available query
    keywords = extract_search_query(restored_query if restored_query != original_query else question)
    words = [w for w in keywords.split() if len(w) > 1]
    
    # Build meaningful phrases (skip common words like thời gian, quy định)
    common_prefixes = {"thời", "gian", "quy", "định", "mức", "tối", "đa", "số", "ngày"}
    key_words = [w for w in words if w not in common_prefixes]
    if not key_words:
        key_words = words
    
    # Build search phrases from key words
    phrases = []
    if len(key_words) >= 2:
        phrases.append(" ".join(key_words[:3]))  # Top 3 key words
        phrases.append(" ".join(key_words[:2]))  # Top 2 key words
    elif key_words:
        phrases.append(key_words[0])
    
    # Also try full keyword string
    if len(words) >= 2:
        phrases.append(" ".join(words[:4]))
    
    # Phase 1: ILIKE phrase search with domain filter (most precise)
    phrase_results = []
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Search with original phrases
        for phrase in phrases[:3]:
            domain_filter = ""
            params = [f"%{phrase}%"]
            if domains:
                domain_filter = "AND lc.domains && %s::legal_domain[]"
                params.append("{" + ",".join(domains) + "}")
            
            cur.execute(f"""
                SELECT lc.id as chunk_id, lc.law_id, ld.title as law_title, 
                       ld.law_number, lc.article, lc.title as chunk_title,
                       lc.content, lc.domains, 1.0::float as rank
                FROM law_chunks lc
                JOIN law_documents ld ON ld.id = lc.law_id
                WHERE lc.content ILIKE %s {domain_filter}
                ORDER BY 
                    CASE WHEN ld.title LIKE 'Bo Luat%%' OR ld.title LIKE 'Bộ luật%%' THEN 0
                         WHEN ld.title LIKE 'Luat %%' OR ld.title LIKE 'Luật %%' THEN 1
                         WHEN ld.title LIKE 'Nghi dinh%%' OR ld.title LIKE 'Nghị định%%' THEN 2
                         ELSE 3 END,
                    length(lc.content) DESC
                LIMIT {limit}
            """, params)
            phrase_results.extend([dict(r) for r in cur.fetchall()])
        
        # If we have a restored query that's different, search with restored phrases too
        if restored_query.lower() != original_query.lower():
            restored_keywords = extract_search_query(restored_query)
            restored_words = [w for w in restored_keywords.split() if len(w) > 1]
            common_prefixes = {"thời", "gian", "quy", "định", "mức", "tối", "đa", "số", "ngày"}
            restored_key_words = [w for w in restored_words if w not in common_prefixes]
            if not restored_key_words:
                restored_key_words = restored_words
            
            restored_phrases = []
            if len(restored_key_words) >= 2:
                restored_phrases.append(" ".join(restored_key_words[:3]))
                restored_phrases.append(" ".join(restored_key_words[:2]))
            elif restored_key_words:
                restored_phrases.append(restored_key_words[0])
            
            for phrase in restored_phrases[:2]:
                domain_filter = ""
                params = [f"%{phrase}%"]
                if domains:
                    domain_filter = "AND lc.domains && %s::legal_domain[]"
                    params.append("{" + ",".join(domains) + "}")
                
                cur.execute(f"""
                    SELECT lc.id as chunk_id, lc.law_id, ld.title as law_title, 
                           ld.law_number, lc.article, lc.title as chunk_title,
                           lc.content, lc.domains, 1.2::float as rank
                    FROM law_chunks lc
                    JOIN law_documents ld ON ld.id = lc.law_id
                    WHERE lc.content ILIKE %s {domain_filter}
                    ORDER BY 
                        CASE WHEN ld.title LIKE 'Bo Luat%%' OR ld.title LIKE 'Bộ luật%%' THEN 0
                             WHEN ld.title LIKE 'Luat %%' OR ld.title LIKE 'Luật %%' THEN 1
                             WHEN ld.title LIKE 'Nghi dinh%%' OR ld.title LIKE 'Nghị định%%' THEN 2
                             ELSE 3 END,
                        length(lc.content) DESC
                    LIMIT {limit}
                """, params)
                phrase_results.extend([dict(r) for r in cur.fetchall()])
    
    # Phase 1.5: Synonym expansion search (e.g., "tndn" → "thu nhập doanh nghiệp")
    synonyms = expand_synonyms(keywords)
    for syn_term in synonyms[:2]:
        sp = syn_term  # Use the expanded term directly as search phrase
        if True:
            domain_filter = ""
            params = [f"%{sp}%"]
            if domains:
                domain_filter = "AND lc.domains && %s::legal_domain[]"
                params.append("{" + ",".join(domains) + "}")
            with get_db() as conn:
                cur = conn.cursor(cursor_factory=RealDictCursor)
                cur.execute(f"""
                    SELECT lc.id as chunk_id, lc.law_id, ld.title as law_title, 
                           ld.law_number, lc.article, lc.title as chunk_title,
                           lc.content, lc.domains, 1.0::float as rank
                    FROM law_chunks lc
                    JOIN law_documents ld ON ld.id = lc.law_id
                    WHERE lc.content ILIKE %s {domain_filter}
                    ORDER BY CASE WHEN ld.title LIKE 'Bo Luat%%' THEN 0 WHEN ld.title LIKE 'Luat%%' THEN 1 ELSE 2 END
                    LIMIT {limit}
                """, params)
                phrase_results.extend([dict(r) for r in cur.fetchall()])
    
    # Phase 2: tsvector search (broader coverage)
    tsv_results = search_laws(keywords, domains, limit)
    
    # Merge: phrase results first, then tsvector
    seen_ids = set()
    merged = []
    
    # Build title matching keywords from original question
    title_keywords = [w for w in words if len(w) > 2]
    synonym_terms = expand_synonyms(keywords)
    for st in synonym_terms:
        title_keywords.extend(st.split())
    
    for result in phrase_results:
        chunk_id = result.get("chunk_id")
        if chunk_id and chunk_id not in seen_ids:
            seen_ids.add(chunk_id)
            title = result.get("law_title", "").lower()
            base_rank = 15.0
            if any(x in result.get("law_title", "") for x in ["Bo Luat", "Bộ luật"]):
                base_rank = 30.0
            elif any(x in result.get("law_title", "") for x in ["Luat ", "Luật "]):
                base_rank = 25.0
            elif any(x in result.get("law_title", "") for x in ["Nghi dinh", "Nghị định"]):
                base_rank = 20.0
            elif result.get("law_title", "").startswith("Legal Document"):
                base_rank = 8.0
            
            # Boost if law title contains search keywords
            title_match_bonus = 0
            for kw in title_keywords:
                if kw.lower() in title:
                    title_match_bonus += 3.0
            
            result["rank"] = base_rank + title_match_bonus
            merged.append(result)
    
    for result in tsv_results:
        chunk_id = result.get("chunk_id") or result.get("id")
        if chunk_id and chunk_id not in seen_ids:
            seen_ids.add(chunk_id)
            title = result.get("law_title", "")
            base = result.get("rank", 1.0)
            if any(x in title for x in ["Bo Luat", "Bộ luật"]):
                result["rank"] = base + 10.0
            elif any(x in title for x in ["Luat ", "Luật "]):
                result["rank"] = base + 5.0
            elif title.startswith("Legal Document"):
                result["rank"] = max(base - 5.0, 0.1)
            merged.append(result)
    
    merged.sort(key=lambda x: x.get("rank", 0), reverse=True)
    return merged[:limit]


def cached_search(query: str, domains=None, limit=10):
    """Cached wrapper around multi_query_search with 1-hour TTL"""
    domains_key = ",".join(sorted(domains)) if domains else "none"
    cache_key = hashlib.md5(f"{query}:{domains_key}:{limit}".encode()).hexdigest()
    now = time.time()
    if cache_key in search_cache:
        result, timestamp = search_cache[cache_key]
        if now - timestamp < CACHE_TTL:
            return result
    result = multi_query_search(query, domains, limit)
    search_cache[cache_key] = (result, now)
    # Limit cache size
    if len(search_cache) > 1000:
        oldest = min(search_cache, key=lambda k: search_cache[k][1])
        del search_cache[oldest]
    return result


# ============================================
# Initialize Agent with shared functions
# ============================================

# Initialize LLM Provider Manager
llm_manager = LLMProviderManager(get_db)

legal_agent.init_agent(
    get_db_fn=get_db,
    multi_query_search_fn=multi_query_search,
    search_laws_fn=search_laws,
    detect_domain_fn=detect_domain,
    fetch_company_context_fn=fetch_company_context,
    llm_provider_manager_fn=llm_manager
)

# ============================================
# API Endpoints
# ============================================

# Root endpoint moved to landing page above

@app.get("/favicon.ico", include_in_schema=False)
@app.get("/favicon.svg", include_in_schema=False)
async def favicon():
    return FileResponse(str(static_dir / "favicon.svg"), media_type="image/svg+xml")

@app.get("/health")
@app.get("/v1/health")
async def health():
    status = {"status": "ok", "version": "2.0.0", "timestamp": time.time(), "ai_engine": "claude-sonnet-4"}
    try:
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute("SELECT count(*) FROM law_documents")
            status["documents"] = cur.fetchone()[0]
            cur.execute("SELECT count(*) FROM law_chunks")
            status["chunks"] = cur.fetchone()[0]
            cur.execute("SELECT count(*) FROM companies")
            status["companies"] = cur.fetchone()[0]
            status["database"] = "connected"
    except Exception:
        status["database"] = "error"
        status["status"] = "degraded"
    return status

@app.post("/v1/chat/upload")
async def chat_upload_file(file: UploadFile = File(...), company: dict = Depends(verify_api_key)):
    """Upload file in chat - extracts text and uploads to Supabase Storage"""
    from .routes.contracts import extract_file_text
    from pathlib import Path
    from .services.file_storage import upload_file

    # Validate file type
    filename = file.filename or "unknown"
    file_ext = os.path.splitext(filename)[1].lower()
    allowed = {".pdf", ".docx", ".doc", ".txt"}
    if file_ext not in allowed:
        raise HTTPException(status_code=400, detail=f"Loại file không hỗ trợ: {file_ext}. Chỉ chấp nhận: {', '.join(allowed)}")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:  # 10MB
        raise HTTPException(status_code=400, detail="File quá lớn. Tối đa 10MB.")

    company_id = str(company["company_id"])
    
    # Upload to Supabase Storage (or local fallback)
    storage_result = await upload_file(content, company_id, filename)
    storage_path = storage_result["storage_path"]
    storage_provider = storage_result["provider"]
    
    # For text extraction, we need a local temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
        tmp_file.write(content)
        tmp_path = tmp_file.name
    
    # Extract text
    extracted_text = None
    try:
        if file_ext == ".txt":
            extracted_text = content.decode('utf-8', errors='ignore')
        else:
            try:
                extracted_text = extract_file_text(tmp_path, file_ext, content)
            except Exception as e:
                print(f"Extraction error: {e}")
                extracted_text = f"[Không thể trích xuất text tự động. File đã được lưu: {filename}]"

        if not extracted_text or len(extracted_text.strip()) < 10:
            extracted_text = f"[File uploaded: {filename}]"
    finally:
        # Cleanup temp file
        try:
            os.unlink(tmp_path)
        except:
            pass
    
    # Save to database with storage path
    doc_id = None
    try:
        import unicodedata
        normalized_text = unicodedata.normalize('NFC', extracted_text)
        
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO documents (company_id, name, extracted_text, doc_type, status, file_path, file_size, mime_type, uploaded_by)
                VALUES (%s, %s, %s, 'other', 'uploaded', %s, %s, %s, %s)
                RETURNING id
            """, (company_id, filename, normalized_text, storage_path, len(content), storage_result.get("content_type", "application/octet-stream"), company.get("user_id")))
            doc_id = str(cur.fetchone()[0])
            conn.commit()
    except Exception as e:
        print(f"Auto-save uploaded doc error: {e}")

    return {
        "filename": filename,
        "content": extracted_text,
        "file_type": file_ext,
        "document_id": doc_id,
        "storage_path": storage_path,
        "storage_provider": storage_provider
    }


@app.get("/v1/documents/{doc_id}/download")
async def download_document(doc_id: str, company: dict = Depends(verify_api_key)):
    """Download original or edited document file from Supabase Storage"""
    from pathlib import Path
    from .services.file_storage import download_file
    
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("""
            SELECT name, file_path, mime_type, company_id
            FROM documents
            WHERE id = %s
        """, (doc_id,))
        
        row = cur.fetchone()
        
        if not row:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Check ownership
        if str(row[3]) != str(company["company_id"]):
            raise HTTPException(status_code=403, detail="Access denied")
        
        file_name = row[0]
        storage_path = row[1]
        mime_type = row[2] or "application/octet-stream"
        
        # Handle old 'chat-upload' entries (no actual file)
        if storage_path == 'chat-upload' or not storage_path:
            raise HTTPException(status_code=404, detail="File not available for download")
        
        # Try to download from Supabase Storage first, fallback to local
        try:
            file_bytes = await download_file(storage_path)
            
            return StreamingResponse(
                BytesIO(file_bytes),
                media_type=mime_type,
                headers={"Content-Disposition": f'attachment; filename="{file_name}"'}
            )
        except Exception as e:
            print(f"Download error: {e}")
            # Fallback: try local path (for backwards compatibility)
            full_path = Path(storage_path) if Path(storage_path).is_absolute() else Path.cwd() / storage_path
            
            if full_path.exists():
                return FileResponse(
                    path=str(full_path),
                    filename=file_name,
                    media_type=mime_type
                )
            
            raise HTTPException(status_code=404, detail="File not found")


@app.get("/v1/documents/{doc_id}/preview")
async def preview_document(doc_id: str, company: dict = Depends(verify_api_key)):
    """Convert DOCX to PDF for web preview using LibreOffice"""
    from pathlib import Path
    from .services.file_storage import download_file
    from .services.libreoffice_editor import convert_to_pdf
    
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("""
            SELECT name, file_path, mime_type, company_id
            FROM documents
            WHERE id = %s
        """, (doc_id,))
        
        row = cur.fetchone()
        
        if not row:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Check ownership
        if str(row[3]) != str(company["company_id"]):
            raise HTTPException(status_code=403, detail="Access denied")
        
        file_name = row[0]
        storage_path = row[1]
        mime_type = row[2] or ""
        
        # Only convert DOCX files
        if not mime_type.startswith("application/") or "word" not in mime_type.lower():
            if not file_name.endswith('.docx'):
                raise HTTPException(status_code=400, detail="Only DOCX files can be previewed as PDF")
        
        # Download file from storage
        try:
            file_bytes = await download_file(storage_path)
        except Exception as e:
            print(f"Download for preview error: {e}")
            raise HTTPException(status_code=404, detail="File not found")
        
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            tmp_docx.write(file_bytes)
            tmp_docx_path = tmp_docx.name
        
        try:
            # Convert to PDF
            pdf_path = convert_to_pdf(tmp_docx_path)

            # Read PDF bytes into memory before cleanup so FileResponse doesn't race the finally block
            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()

        except Exception as e:
            print(f"PDF conversion error: {e}")
            raise HTTPException(status_code=500, detail=f"Cannot convert to PDF: {str(e)}")

        finally:
            # Cleanup temp files
            try:
                os.unlink(tmp_docx_path)
                if 'pdf_path' in locals() and os.path.exists(pdf_path):
                    os.unlink(pdf_path)
            except:
                pass

        return StreamingResponse(
            BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'inline; filename="{os.path.splitext(file_name)[0]}.pdf"'}
        )


@app.post("/v1/documents/{doc_id}/edit-docx")
async def edit_docx_endpoint(doc_id: str, edits: dict = Body(...), company: dict = Depends(verify_api_key)):
    """Apply text edits to a DOCX file using LibreOffice for 99% format preservation"""
    from pathlib import Path
    from .services.libreoffice_editor import edit_docx
    from .services.file_storage import download_file, upload_file
    
    with get_db() as conn:
        cur = conn.cursor()
        
        # Get document
        cur.execute("""
            SELECT name, file_path, mime_type, company_id
            FROM documents
            WHERE id = %s
        """, (doc_id,))
        
        row = cur.fetchone()
        
        if not row:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Check ownership
        if str(row[3]) != str(company["company_id"]):
            raise HTTPException(status_code=403, detail="Access denied")
        
        file_name = row[0]
        storage_path = row[1]
        mime_type = row[2]
        company_id = str(company["company_id"])
        
        # Validate it's a DOCX file
        if not file_name.endswith('.docx') and 'wordprocessingml' not in (mime_type or ''):
            raise HTTPException(status_code=400, detail="Only .docx files can be edited")
        
        if storage_path == 'chat-upload' or not storage_path:
            raise HTTPException(status_code=404, detail="Original file not available")
        
        # Download from Supabase Storage
        try:
            file_bytes = await download_file(storage_path)
        except Exception as e:
            print(f"Download error: {e}")
            raise HTTPException(status_code=404, detail="Original file not found")
        
        # Create temp files for editing
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_input:
            tmp_input.write(file_bytes)
            input_path = tmp_input.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_output:
            output_path = tmp_output.name
        
        # Apply edits
        try:
            edits_list = edits.get("edits", [])
            if not edits_list:
                raise HTTPException(status_code=400, detail="No edits provided")
            
            result = edit_docx(
                input_path=input_path,
                output_path=output_path,
                edits=edits_list
            )
            
            # Read edited file
            with open(output_path, 'rb') as f:
                edited_bytes = f.read()
            
            # Upload edited file to Supabase Storage
            base_name = os.path.splitext(file_name)[0]
            edited_filename = f"{base_name}_edited.docx"
            storage_result = await upload_file(edited_bytes, company_id, edited_filename)
            new_storage_path = storage_result["storage_path"]
            
            # Update database with new storage path
            cur.execute("""
                UPDATE documents
                SET file_path = %s,
                    status = 'edited',
                    file_size = %s
                WHERE id = %s
            """, (new_storage_path, len(edited_bytes), doc_id))
            conn.commit()
            
            return {
                "message": "Document edited successfully with LibreOffice",
                "document_id": doc_id,
                "changes_made": result["changes_made"],
                "method": result.get("method", "unknown"),
                "download_url": f"/v1/documents/{doc_id}/download",
                "preview_url": f"/v1/documents/{doc_id}/preview",
                "output_filename": edited_filename
            }
            
        except Exception as e:
            print(f"Edit error: {e}")
            raise HTTPException(status_code=500, detail=f"Edit failed: {str(e)}")
        
        finally:
            # Cleanup temp files
            try:
                os.unlink(input_path)
                os.unlink(output_path)
            except:
                pass


@app.post("/v1/legal/ask", response_model=LegalResponse)
async def legal_ask(query: LegalQuery, company: dict = Depends(verify_api_key)):
    """Tư vấn pháp luật - Legal Q&A (Agent-based with tool use)"""
    check_rate_limit(str(company["company_id"]))
    
    # Load chat history for multi-turn conversation
    chat_history = []
    session_id = None
    user_id = company.get("user_id")
    
    if query.session_id and user_id:
        try:
            with get_db() as conn:
                cur = conn.cursor(cursor_factory=RealDictCursor)
                cur.execute("""
                    SELECT role, content FROM messages
                    WHERE session_id = %s AND company_id = %s
                    ORDER BY created_at ASC
                    LIMIT 50
                """, (query.session_id, company["company_id"]))
                rows = cur.fetchall()
                for row in rows:
                    chat_history.append({"role": row["role"], "content": row["content"]})
                session_id = query.session_id
        except Exception as e:
            print(f"Error loading chat history: {e}")
    
    # Augment question with file context if provided
    actual_question = query.question
    if query.file_context:
        actual_question = f"""[Người dùng đã upload file: {query.file_context.filename}]

NỘI DUNG FILE:
{query.file_context.content[:30000]}

CÂU HỎI: {query.question}"""

    # Run the agent
    result = await legal_agent.run_agent(
        question=actual_question,
        company_id=str(company["company_id"]),
        user_id=str(user_id) if user_id else None,
        session_id=str(session_id) if session_id else None,
        chat_history=chat_history
    )
    
    citations = result.get("citations", [])
    
    # Save to chat history
    if user_id:
        with get_db() as conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            
            if not session_id:
                cur.execute("""
                    SELECT id FROM chat_sessions
                    WHERE user_id = %s AND company_id = %s AND agent_type = 'qa' AND status = 'active'
                    ORDER BY last_message_at DESC NULLS LAST
                    LIMIT 1
                """, (user_id, company["company_id"]))
                
                session = cur.fetchone()
                if session:
                    session_id = session["id"]
                else:
                    cur.execute("""
                        INSERT INTO chat_sessions (company_id, user_id, agent_type, title, status)
                        VALUES (%s, %s, 'qa', %s, 'active')
                        RETURNING id
                    """, (company["company_id"], user_id, f"Q&A - {query.question[:50]}..."))
                    session_id = cur.fetchone()["id"]
            
            cur.execute("""
                INSERT INTO messages (session_id, company_id, role, content, tokens_used, model)
                VALUES (%s, %s, 'user', %s, 0, '')
            """, (session_id, company["company_id"], query.question))
            
            total_tokens = result.get("input_tokens", 0) + result.get("output_tokens", 0)
            cur.execute("""
                INSERT INTO messages (session_id, company_id, role, content, citations, confidence, tokens_used, model)
                VALUES (%s, %s, 'assistant', %s, %s, %s, %s, %s)
            """, (
                session_id,
                company["company_id"],
                result["answer"], 
                json.dumps(citations),
                0.85 if citations else 0.5,
                total_tokens,
                result.get("model", "claude-sonnet-4-20250514")
            ))
            
            cur.execute("""
                UPDATE chat_sessions 
                SET message_count = message_count + 2, last_message_at = now()
                WHERE id = %s
            """, (session_id,))
            
            conn.commit()
    
    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s", (company["company_id"],))
        cur.execute("""
            INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
            VALUES (%s, '/v1/legal/ask', 'agent', %s, %s, 200)
        """, (company["company_id"], result.get("input_tokens", 0), result.get("output_tokens", 0)))
        conn.commit()
    
    # Audit log for chat
    log_audit(str(company["company_id"]), str(user_id) if user_id else None, "chat", "ai", str(session_id) if session_id else None, {"question_length": len(query.question)})

    return LegalResponse(
        answer=result["answer"],
        citations=citations,
        confidence=0.85 if citations else 0.5,
        tokens_used=result.get("input_tokens", 0) + result.get("output_tokens", 0),
        model=result.get("model", "claude-sonnet-4-20250514"),
        session_id=str(session_id) if session_id else None
    )


@app.post("/v1/legal/ask-stream")
async def legal_ask_stream(query: LegalQuery, company: dict = Depends(verify_api_key)):
    """Tư vấn pháp luật với streaming SSE - Agent-based with tool status events"""
    check_rate_limit(str(company["company_id"]))

    # Load chat history
    chat_history = []
    session_id = None
    user_id = company.get("user_id")

    if user_id:
        target_session = query.session_id
        # If no session_id, auto-find latest active session (within 30 min) for continuity
        if not target_session:
            try:
                with get_db() as conn:
                    cur = conn.cursor(cursor_factory=RealDictCursor)
                    cur.execute("""
                        SELECT id FROM chat_sessions
                        WHERE user_id = %s AND company_id = %s AND status = 'active'
                          AND last_message_at > now() - interval '30 minutes'
                        ORDER BY last_message_at DESC LIMIT 1
                    """, (user_id, company["company_id"]))
                    recent = cur.fetchone()
                    if recent:
                        target_session = str(recent["id"])
            except Exception:
                pass

        if target_session:
            try:
                with get_db() as conn:
                    cur = conn.cursor(cursor_factory=RealDictCursor)
                    cur.execute("""
                        SELECT role, content FROM messages
                        WHERE session_id = %s AND company_id = %s
                        ORDER BY created_at ASC
                        LIMIT 50
                    """, (target_session, company["company_id"]))
                    rows = cur.fetchall()
                    for row in rows:
                        chat_history.append({"role": row["role"], "content": row["content"]})
                    session_id = target_session
            except Exception as e:
                print(f"Error loading chat history: {e}")

    # Augment question with file context if provided
    actual_question = query.question
    if query.file_context:
        actual_question = f"""[Người dùng đã upload file: {query.file_context.filename}]

NỘI DUNG FILE:
{query.file_context.content[:30000]}

CÂU HỎI: {query.question}"""

    company_id_str = str(company["company_id"])

    async def sse_generator():
        """Generate SSE events using agent streaming"""
        full_response = []
        all_citations = []

        try:
            async for event_str in legal_agent.run_agent_stream_final_text(
                question=actual_question,
                company_id=company_id_str,
                user_id=str(user_id) if user_id else None,
                session_id=str(session_id) if session_id else None,
                chat_history=chat_history
            ):
                yield event_str

                # Parse to collect full text and citations for saving
                if event_str.startswith("data: "):
                    try:
                        evt = json.loads(event_str[6:].strip())
                        if evt.get("type") == "delta":
                            full_response.append(evt.get("text", ""))
                        elif evt.get("type") == "citations":
                            all_citations = evt.get("citations", [])
                        elif evt.get("type") == "done":
                            if evt.get("citations"):
                                all_citations = evt["citations"]
                    except:
                        pass

        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"
            return

        complete_text = "".join(full_response)

        # Save to chat history
        saved_session_id = session_id
        if user_id and complete_text:
            try:
                with get_db() as conn:
                    cur = conn.cursor(cursor_factory=RealDictCursor)
                    if not saved_session_id:
                        cur.execute("""
                            SELECT id FROM chat_sessions
                            WHERE user_id = %s AND company_id = %s AND agent_type = 'qa' AND status = 'active'
                            ORDER BY last_message_at DESC NULLS LAST LIMIT 1
                        """, (user_id, company["company_id"]))
                        session = cur.fetchone()
                        if session:
                            saved_session_id = session["id"]
                        else:
                            cur.execute("""
                                INSERT INTO chat_sessions (company_id, user_id, agent_type, title, status)
                                VALUES (%s, %s, 'qa', %s, 'active') RETURNING id
                            """, (company["company_id"], user_id, f"Q&A - {query.question[:50]}..."))
                            saved_session_id = cur.fetchone()["id"]

                    cur.execute("""
                        INSERT INTO messages (session_id, company_id, role, content, tokens_used, model)
                        VALUES (%s, %s, 'user', %s, 0, '')
                    """, (saved_session_id, company["company_id"], actual_question))

                    cur.execute("""
                        INSERT INTO messages (session_id, company_id, role, content, citations, confidence, tokens_used, model)
                        VALUES (%s, %s, 'assistant', %s, %s, %s, 0, 'claude-sonnet-4-20250514')
                    """, (
                        saved_session_id, company["company_id"],
                        complete_text, json.dumps(all_citations),
                        0.85 if all_citations else 0.5
                    ))

                    cur.execute("""
                        UPDATE chat_sessions SET message_count = message_count + 2, last_message_at = now()
                        WHERE id = %s
                    """, (saved_session_id,))
                    conn.commit()
                    # Emit session_id so frontend can track conversation
                    if saved_session_id:
                        yield f"data: {json.dumps({'type': 'session_id', 'content': str(saved_session_id)})}\n\n"
            except Exception as e:
                print(f"Error saving stream chat history: {e}")

        # Update usage
        try:
            with get_db() as conn:
                cur = conn.cursor()
                cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s",
                            (company["company_id"],))
                cur.execute("""
                    INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
                    VALUES (%s, '/v1/legal/ask-stream', 'agent', 0, 0, 200)
                """, (company["company_id"],))
                conn.commit()
        except Exception as e:
            print(f"Error updating usage: {e}")

    return StreamingResponse(
        sse_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        }
    )


@app.get("/v1/legal/search-detailed")
async def search_detailed(
    q: str,
    domains: Optional[str] = None,
    limit: int = Query(20, ge=1, le=50),
    company: dict = Depends(verify_api_key)
):
    """Tìm kiếm luật chi tiết - Detailed Law Search with categorization and highlighting"""
    import time as _time

    start_time = _time.time()

    domain_list = domains.split(",") if domains else None

    # Use cached_search for better results with caching
    results = cached_search(q, domain_list, min(limit, 50))

    elapsed = _time.time() - start_time

    # Categorize results by law type
    categories = {}
    highlighted_results = []
    query_words = [w.lower() for w in q.split() if len(w) > 1]

    for r in results:
        law_title = r.get("law_title", "")
        law_number = r.get("law_number", "")
        content = r.get("content", "")

        # Determine category
        if any(x in law_title for x in ["Bộ luật", "Bo Luat"]):
            cat = "Bộ luật"
        elif any(x in law_title for x in ["Luật ", "Luat "]):
            cat = "Luật"
        elif any(x in law_title for x in ["Nghị định", "Nghi dinh"]):
            cat = "Nghị định"
        elif any(x in law_title for x in ["Thông tư", "Thong tu"]):
            cat = "Thông tư"
        elif any(x in law_title for x in ["Quyết định", "Quyet dinh"]):
            cat = "Quyết định"
        else:
            cat = "Khác"

        if cat not in categories:
            categories[cat] = []

        # Highlight matching text
        highlighted_content = content
        for word in query_words:
            if len(word) > 2:
                pattern = re_module.compile(re_module.escape(word), re_module.IGNORECASE)
                highlighted_content = pattern.sub(f"<mark>{word}</mark>", highlighted_content)

        result_item = {
            "law_title": law_title,
            "law_number": law_number,
            "article": r.get("article"),
            "chunk_title": r.get("chunk_title", ""),
            "content": content[:1000],
            "highlighted_content": highlighted_content[:1000],
            "rank": float(r.get("rank", 0)),
            "category": cat,
            "domains": r.get("domains", [])
        }

        categories[cat].append(result_item)
        highlighted_results.append(result_item)

    # Find related articles (same law, adjacent articles) for top results
    related_articles = []
    if results:
        top_law_ids = list(set(r.get("law_id") for r in results[:3] if r.get("law_id")))
        top_articles = [r.get("article") for r in results[:3] if r.get("article")]

        if top_law_ids:
            try:
                with get_db() as conn:
                    cur = conn.cursor(cursor_factory=RealDictCursor)
                    # Get adjacent articles from the same laws
                    for law_id in top_law_ids[:2]:
                        cur.execute("""
                            SELECT lc.article, lc.title as chunk_title, ld.title as law_title,
                                   ld.law_number, LEFT(lc.content, 300) as content_preview
                            FROM law_chunks lc
                            JOIN law_documents ld ON ld.id = lc.law_id
                            WHERE lc.law_id = %s
                              AND lc.article IS NOT NULL
                            ORDER BY lc.article
                            LIMIT 10
                        """, (law_id,))
                        rows = cur.fetchall()
                        for row in rows:
                            if row["article"] not in top_articles:
                                related_articles.append({
                                    "law_title": row["law_title"],
                                    "law_number": row["law_number"],
                                    "article": row["article"],
                                    "chunk_title": row["chunk_title"],
                                    "content_preview": row["content_preview"]
                                })
            except Exception as e:
                print(f"Error fetching related articles: {e}")

    # Category counts
    category_stats = {cat: len(items) for cat, items in categories.items()}

    return {
        "query": q,
        "stats": {
            "total_results": len(results),
            "search_time_seconds": round(elapsed, 3),
            "categories": category_stats,
            "domains_searched": domain_list
        },
        "results": highlighted_results,
        "categories": {cat: items[:10] for cat, items in categories.items()},
        "related_articles": related_articles[:10]
    }


@app.post("/v1/legal/review")
async def contract_review(review: ContractReview, company: dict = Depends(verify_api_key)):
    """Rà soát hợp đồng - Contract Review"""
    check_rate_limit(str(company["company_id"]))
    
    # Search relevant laws based on contract type
    search_terms = {
        "hop_dong_lao_dong": "hợp đồng lao động quyền nghĩa vụ",
        "hop_dong_thuong_mai": "hợp đồng thương mại mua bán",
        "hop_dong_dich_vu": "hợp đồng dịch vụ thuê khoán",
    }
    search_query = search_terms.get(review.contract_type, "hợp đồng điều khoản")
    sources = cached_search(search_query, None, 15)
    
    context = "\n\n".join([
        f"[{src['law_title']}] {src.get('article', '')}\n{src['content'][:1500]}"
        for src in sources
    ])
    
    system_prompt = """Bạn là luật sư chuyên rà soát hợp đồng theo pháp luật Việt Nam.

Nhiệm vụ: Rà soát hợp đồng và đánh giá theo các tiêu chí:
1. **Tính hợp pháp**: Có điều khoản nào vi phạm pháp luật không?
2. **Tính đầy đủ**: Có thiếu điều khoản bắt buộc nào không?
3. **Rủi ro**: Những điều khoản nào có rủi ro cao cho bên nào?
4. **Đề xuất**: Các sửa đổi cần thiết

Trả về JSON format:
{
    "risk_score": 1-100 (100 = rủi ro cao nhất),
    "issues": [{"type": "violation|missing|risk|suggestion", "severity": "critical|high|medium|low", "clause": "điều khoản liên quan", "description": "mô tả", "legal_basis": "căn cứ pháp lý", "recommendation": "đề xuất sửa"}],
    "summary": "Tóm tắt đánh giá",
    "overall_assessment": "Đánh giá tổng thể"
}"""

    user_message = f"""HỢP ĐỒNG CẦN RÀ SOÁT:
{review.contract_text[:50000]}

PHÁP LUẬT LIÊN QUAN:
{context}

{f"YÊU CẦU ĐẶC BIỆT: {', '.join(review.focus_areas)}" if review.focus_areas else ""}

Hãy rà soát hợp đồng trên và trả về kết quả theo format JSON."""

    result = await call_claude(system_prompt, user_message, max_tokens=8192)
    
    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s", (company["company_id"],))
        cur.execute("""
            INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
            VALUES (%s, '/v1/legal/review', 'review', %s, %s, 200)
        """, (company["company_id"], result["input_tokens"], result["output_tokens"]))
        conn.commit()
    
    # Try to parse JSON from response
    try:
        review_data = json.loads(result["content"])
    except:
        review_data = {"raw_analysis": result["content"]}
    
    # Audit log for contract review
    log_audit(str(company["company_id"]), str(company.get("user_id")) if company.get("user_id") else None, "review", "contract", None, {"contract_type": review.contract_type})

    return {
        "review": review_data,
        "tokens_used": result["input_tokens"] + result["output_tokens"],
        "model": result["model"]
    }

@app.post("/v1/legal/draft")
async def document_draft(draft: DocumentDraft, company: dict = Depends(verify_api_key)):
    """Soạn thảo văn bản - Document Drafting"""
    check_rate_limit(str(company["company_id"]))
    
    # Search for templates and relevant laws
    sources = cached_search(draft.doc_type.replace("_", " "), None, 10)
    
    context = "\n\n".join([
        f"[{src['law_title']}] {src.get('article', '')}\n{src['content'][:1500]}"
        for src in sources
    ])
    
    system_prompt = """Bạn là chuyên gia soạn thảo văn bản pháp lý Việt Nam.

Nhiệm vụ: Soạn thảo văn bản hoàn chỉnh, đúng format, đúng pháp luật.

Quy tắc:
1. Sử dụng đúng format văn bản hành chính Việt Nam
2. Tuân thủ quy định tại Nghị định 30/2020/NĐ-CP về công tác văn thư
3. Điền đầy đủ thông tin từ biến số được cung cấp
4. Các điều khoản phải tuân thủ pháp luật hiện hành
5. Ghi rõ căn cứ pháp lý"""

    variables_str = json.dumps(draft.variables, ensure_ascii=False, indent=2)
    
    user_message = f"""LOẠI VĂN BẢN: {draft.doc_type}

THÔNG TIN:
{variables_str}

{f"YÊU CẦU BỔ SUNG: {draft.instructions}" if draft.instructions else ""}

PHÁP LUẬT LIÊN QUAN:
{context}

Hãy soạn thảo văn bản hoàn chỉnh."""

    result = await call_claude(system_prompt, user_message, max_tokens=8192)
    
    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s", (company["company_id"],))
        cur.execute("""
            INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
            VALUES (%s, '/v1/legal/draft', 'draft', %s, %s, 200)
        """, (company["company_id"], result["input_tokens"], result["output_tokens"]))
        conn.commit()
    
    return {
        "document": result["content"],
        "doc_type": draft.doc_type,
        "tokens_used": result["input_tokens"] + result["output_tokens"],
        "model": result["model"]
    }

@app.get("/v1/legal/search")
async def search(q: str, domains: Optional[str] = None, limit: int = 10, company: dict = Depends(verify_api_key)):
    """Tìm kiếm luật - Law Search"""
    domain_list = domains.split(",") if domains else None
    results = cached_search(q, domain_list, min(limit, 30))
    
    return {
        "query": q,
        "count": len(results),
        "results": [{
            "law_title": r["law_title"],
            "law_number": r["law_number"],
            "article": r.get("article"),
            "content": r["content"][:500],
            "rank": float(r.get("rank", 0))
        } for r in results]
    }

# ============================================
# Admin endpoints (internal)
# ============================================

@app.post("/admin/company", include_in_schema=False)
async def create_company(name: str, slug: str, plan: str = "trial"):
    """Create a new company (admin only)"""
    import secrets
    
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Create company
        cur.execute("""
            INSERT INTO companies (name, slug, plan)
            VALUES (%s, %s, %s::plan_type)
            RETURNING id, name, slug, plan, monthly_quota
        """, (name, slug, plan))
        company = dict(cur.fetchone())
        
        # Generate API key
        api_key = f"lak_{secrets.token_hex(24)}"
        key_hash = hashlib.sha256(api_key.encode()).hexdigest()
        
        cur.execute("""
            INSERT INTO api_keys (company_id, name, key_hash, key_prefix)
            VALUES (%s, %s, %s, %s)
        """, (company["id"], f"{name} - Default Key", key_hash, api_key[:8]))
        
        conn.commit()
        
        return {
            "company": company,
            "api_key": api_key,
            "warning": "Save this API key - it cannot be retrieved later"
        }

# ============================================
# Export to DOCX
# ============================================

class ExportRequest(BaseModel):
    content: str  # Markdown content from AI response
    filename: Optional[str] = "legal-document"

@app.post("/v1/legal/export-docx")
async def export_docx(req: ExportRequest):
    """Convert markdown content to a professional .docx file"""
    import io
    import re
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.name = 'Times New Roman'
        h_font.bold = True
        h_font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            h_font.size = Pt(16)
        elif level == 2:
            h_font.size = Pt(14)
        else:
            h_font.size = Pt(13)
    
    lines = req.content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headings
        if line.startswith('### '):
            text = line[4:].strip().strip('*')
            doc.add_heading(text, level=3)
        elif line.startswith('## '):
            text = line[3:].strip().strip('*')
            doc.add_heading(text, level=2)
        elif line.startswith('# '):
            text = line[2:].strip().strip('*')
            doc.add_heading(text, level=1)
        elif line.startswith('---'):
            # Horizontal rule - add thin line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 50)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(150, 150, 150)
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet list
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
        elif re.match(r'^\d+[\.\)] ', line):
            # Numbered list
            text = re.sub(r'^\d+[\.\)] ', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)
        else:
            # Normal paragraph
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
        
        i += 1
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    safe_name = re.sub(r'[^\w\-]', '_', req.filename)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name}.docx"'
        }
    )

def _add_formatted_text(paragraph, text: str):
    """Parse markdown bold/italic and add formatted runs to paragraph"""
    import re
    from docx.shared import Pt
    
    # Split by **bold** and *italic* patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
        elif part.startswith('[') and ']' in part:
            # Placeholder like [TÊN CÔNG TY] - highlight it
            run = paragraph.add_run(part)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(200, 0, 0)  # Red for placeholders
        else:
            # Check for [PLACEHOLDER] within normal text
            sub_parts = re.split(r'(\[.*?\])', part)
            for sp in sub_parts:
                if sp.startswith('[') and sp.endswith(']'):
                    run = paragraph.add_run(sp)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(200, 0, 0)
                else:
                    run = paragraph.add_run(sp)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)


# ============================================
# Feature: Smart Contract Comparison
# ============================================

class ContractCompareRequest(BaseModel):
    contract_ids: List[str] = Field(..., min_length=2, description="Danh sách ID hợp đồng cần so sánh (tối thiểu 2)")

@app.post("/v1/contracts/compare")
async def compare_contracts(req: ContractCompareRequest, company: dict = Depends(verify_api_key)):
    """Compare 2+ contracts side-by-side with AI analysis"""
    check_rate_limit(str(company["company_id"]))

    if len(req.contract_ids) < 2:
        raise HTTPException(status_code=400, detail="Cần ít nhất 2 hợp đồng để so sánh")
    if len(req.contract_ids) > 5:
        raise HTTPException(status_code=400, detail="Tối đa 5 hợp đồng")

    contracts_data = []
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        for cid in req.contract_ids:
            cur.execute("""
                SELECT id, name, contract_type, extracted_text, parties,
                       start_date, end_date, value, status
                FROM contracts
                WHERE id::text = %s AND company_id = %s AND status != 'deleted'
            """, (cid, company["company_id"]))
            contract = cur.fetchone()
            if not contract:
                raise HTTPException(status_code=404, detail=f"Không tìm thấy hợp đồng: {cid}")
            c = dict(contract)
            for key in ["start_date", "end_date"]:
                if c.get(key):
                    c[key] = str(c[key])
            contracts_data.append(c)

    # Build comparison prompt
    contracts_text = ""
    for i, c in enumerate(contracts_data, 1):
        text = (c.get("extracted_text") or "")[:8000]
        contracts_text += f"\n\n--- HỢP ĐỒNG {i}: {c['name']} (Loại: {c.get('contract_type', 'N/A')}) ---\n{text}"

    system_prompt = """Bạn là luật sư chuyên so sánh hợp đồng. Phân tích và so sánh các hợp đồng được cung cấp.

Trả về JSON format:
{
    "summary": "Tóm tắt so sánh tổng quan",
    "contracts": [{"id": "...", "name": "...", "strengths": ["..."], "weaknesses": ["..."]}],
    "differences": [{"aspect": "Khía cạnh", "details": [{"contract": "Tên HĐ", "content": "Nội dung"}], "recommendation": "Đề xuất"}],
    "inconsistencies": ["Điểm không nhất quán giữa các hợp đồng"],
    "best_for_company": {"contract_name": "...", "reason": "Lý do"},
    "recommendations": ["Đề xuất chung"]
}"""

    user_message = f"So sánh {len(contracts_data)} hợp đồng sau:{contracts_text}\n\nHãy phân tích chi tiết sự khác biệt, điểm mạnh/yếu của từng hợp đồng, và đề xuất hợp đồng nào có lợi hơn cho công ty."

    result = await call_claude(system_prompt, user_message, max_tokens=8192)

    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s", (company["company_id"],))
        cur.execute("""
            INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
            VALUES (%s, '/v1/contracts/compare', 'compare', %s, %s, 200)
        """, (company["company_id"], result["input_tokens"], result["output_tokens"]))
        conn.commit()

    try:
        comparison = json.loads(result["content"])
    except:
        comparison = {"raw_analysis": result["content"]}

    return {
        "comparison": comparison,
        "contracts": [{
            "id": str(c["id"]),
            "name": c["name"],
            "contract_type": c.get("contract_type"),
            "start_date": c.get("start_date"),
            "end_date": c.get("end_date"),
            "value": c.get("value")
        } for c in contracts_data],
        "tokens_used": result["input_tokens"] + result["output_tokens"],
        "model": result["model"]
    }


# ============================================
# Feature: Contract Timeline & Deadline Tracking
# ============================================

@app.get("/v1/contracts/timeline")
async def contract_timeline(company: dict = Depends(verify_api_key)):
    """Get all contract deadlines, expiry dates as a timeline"""
    from datetime import datetime, timedelta

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("""
            SELECT id, name, contract_type, start_date, end_date, value, status, parties
            FROM contracts
            WHERE company_id = %s AND status != 'deleted'
              AND (start_date IS NOT NULL OR end_date IS NOT NULL)
            ORDER BY
                CASE WHEN end_date IS NOT NULL THEN end_date ELSE start_date END ASC
        """, (company["company_id"],))
        contracts = cur.fetchall()

    now = datetime.now().date()
    timeline_items = []
    overdue_count = 0
    expiring_30 = 0
    expiring_60 = 0
    expiring_90 = 0

    for c in contracts:
        item = dict(c)
        for key in ["start_date", "end_date"]:
            if item.get(key):
                item[key] = str(item[key])
        if item.get("parties"):
            try:
                if isinstance(item["parties"], str):
                    item["parties"] = json.loads(item["parties"])
            except:
                pass

        end_date = c.get("end_date")
        if end_date:
            days_remaining = (end_date - now).days
            item["days_remaining"] = days_remaining
            if days_remaining < 0:
                item["timeline_status"] = "overdue"
                overdue_count += 1
            elif days_remaining <= 30:
                item["timeline_status"] = "expiring_soon"
                expiring_30 += 1
            elif days_remaining <= 60:
                item["timeline_status"] = "expiring_60"
                expiring_60 += 1
            elif days_remaining <= 90:
                item["timeline_status"] = "expiring_90"
                expiring_90 += 1
            else:
                item["timeline_status"] = "active"
        else:
            item["days_remaining"] = None
            item["timeline_status"] = "no_end_date"

        item["id"] = str(item["id"])
        timeline_items.append(item)

    return {
        "timeline": timeline_items,
        "summary": {
            "total": len(timeline_items),
            "overdue": overdue_count,
            "expiring_30_days": expiring_30,
            "expiring_60_days": expiring_60,
            "expiring_90_days": expiring_90,
            "today": str(now)
        }
    }


# ============================================
# Feature: Document Annotations
# ============================================

@app.on_event("startup")
async def create_annotations_table():
    """Create annotations table if not exists"""
    try:
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute("""
                CREATE TABLE IF NOT EXISTS document_annotations (
                    id UUID DEFAULT gen_random_uuid() PRIMARY KEY,
                    document_id UUID NOT NULL,
                    company_id UUID NOT NULL,
                    user_id UUID,
                    text_selection TEXT,
                    start_offset INTEGER,
                    end_offset INTEGER,
                    comment TEXT NOT NULL,
                    annotation_type VARCHAR(50) DEFAULT 'comment',
                    is_ai_generated BOOLEAN DEFAULT FALSE,
                    created_at TIMESTAMPTZ DEFAULT now(),
                    updated_at TIMESTAMPTZ DEFAULT now()
                )
            """)
            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_annotations_document
                ON document_annotations(document_id, company_id)
            """)
            conn.commit()
    except Exception as e:
        print(f"Annotations table creation: {e}")


class AnnotationCreate(BaseModel):
    text_selection: Optional[str] = None
    start_offset: Optional[int] = None
    end_offset: Optional[int] = None
    comment: str = Field(..., min_length=1, max_length=5000)
    annotation_type: str = Field("comment", description="comment, issue, suggestion, highlight")


@app.post("/v1/documents/{doc_id}/annotate")
async def annotate_document(doc_id: str, annotation: AnnotationCreate, company: dict = Depends(verify_api_key)):
    """Add annotation/comment to a specific part of document"""
    user_id = company.get("user_id")

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        # Verify document belongs to company
        cur.execute("""
            SELECT id FROM documents WHERE id::text = %s AND company_id = %s
            UNION ALL
            SELECT id FROM contracts WHERE id::text = %s AND company_id = %s AND status != 'deleted'
        """, (doc_id, company["company_id"], doc_id, company["company_id"]))
        if not cur.fetchone():
            raise HTTPException(status_code=404, detail="Tài liệu không tồn tại")

        cur.execute("""
            INSERT INTO document_annotations
                (document_id, company_id, user_id, text_selection, start_offset, end_offset, comment, annotation_type)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id, document_id, text_selection, start_offset, end_offset, comment, annotation_type, is_ai_generated, created_at
        """, (
            doc_id, company["company_id"], user_id,
            annotation.text_selection, annotation.start_offset, annotation.end_offset,
            annotation.comment, annotation.annotation_type
        ))
        result = dict(cur.fetchone())
        conn.commit()

    result["id"] = str(result["id"])
    result["document_id"] = str(result["document_id"])
    result["created_at"] = str(result["created_at"])
    return result


@app.get("/v1/documents/{doc_id}/annotations")
async def get_annotations(doc_id: str, company: dict = Depends(verify_api_key)):
    """Get all annotations for a document"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("""
            SELECT a.id, a.document_id, a.user_id, a.text_selection, a.start_offset, a.end_offset,
                   a.comment, a.annotation_type, a.is_ai_generated, a.created_at,
                   u.full_name as author_name
            FROM document_annotations a
            LEFT JOIN users u ON u.id = a.user_id
            WHERE a.document_id::text = %s AND a.company_id = %s
            ORDER BY a.start_offset ASC NULLS LAST, a.created_at ASC
        """, (doc_id, company["company_id"]))
        annotations = cur.fetchall()

    result = []
    for ann in annotations:
        item = dict(ann)
        for key in ["id", "document_id", "user_id"]:
            if item.get(key):
                item[key] = str(item[key])
        if item.get("created_at"):
            item["created_at"] = str(item["created_at"])
        result.append(item)

    return {"annotations": result, "total": len(result)}


@app.delete("/v1/documents/{doc_id}/annotations/{annotation_id}")
async def delete_annotation(doc_id: str, annotation_id: str, company: dict = Depends(verify_api_key)):
    """Delete an annotation"""
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("""
            DELETE FROM document_annotations
            WHERE id::text = %s AND document_id::text = %s AND company_id = %s
        """, (annotation_id, doc_id, company["company_id"]))
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Annotation không tồn tại")
        conn.commit()
    return {"message": "Đã xóa annotation"}


# ============================================
# Feature: Notifications Endpoint
# ============================================

@app.get("/v1/notifications")
async def get_notifications(company: dict = Depends(verify_api_key)):
    """Get notifications for the company — expiring contracts, overdue items"""
    from datetime import date as date_type
    notifications = []

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Expiring contracts (within 30 days)
        cur.execute("""
            SELECT id, name, end_date
            FROM contracts
            WHERE company_id = %s AND status != 'deleted'
            AND end_date IS NOT NULL
            AND end_date <= CURRENT_DATE + INTERVAL '30 days'
            AND end_date >= CURRENT_DATE
            ORDER BY end_date ASC
        """, (company["company_id"],))

        for c in cur.fetchall():
            days = (c['end_date'] - date_type.today()).days
            notifications.append({
                "type": "expiring_contract",
                "severity": "warning" if days > 7 else "critical",
                "title": f"HĐ sắp hết hạn: {c['name']}",
                "detail": f"Còn {days} ngày",
                "action": {"type": "view_contract", "id": str(c['id'])}
            })

        # Overdue contracts
        cur.execute("""
            SELECT id, name, end_date
            FROM contracts
            WHERE company_id = %s AND status != 'deleted'
            AND end_date IS NOT NULL AND end_date < CURRENT_DATE
            ORDER BY end_date DESC LIMIT 5
        """, (company["company_id"],))

        for c in cur.fetchall():
            notifications.append({
                "type": "overdue_contract",
                "severity": "critical",
                "title": f"HĐ đã hết hạn: {c['name']}",
                "detail": f"Hết hạn {c['end_date']}",
                "action": {"type": "view_contract", "id": str(c['id'])}
            })

    return {"notifications": notifications, "count": len(notifications)}


# ============================================
# Feature: Smart Contract Report Export (Word)
# ============================================

@app.post("/v1/contracts/{contract_id}/report")
async def generate_contract_report(contract_id: str, company: dict = Depends(verify_api_key)):
    """Generate professional AI review report for a contract as Word document"""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    company_id = str(company["company_id"])

    try:
        with get_db() as conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute(
                "SELECT * FROM contracts WHERE id::text = %s AND company_id = %s AND status != 'deleted'",
                (contract_id, company_id)
            )
            contract = cur.fetchone()
            if not contract:
                raise HTTPException(404, "Hợp đồng không tồn tại")
    except HTTPException:
        raise
    except Exception as e:
        print(f"DB error in contract report: {e}")
        raise HTTPException(500, "Lỗi truy vấn database")

    # Run AI analysis
    try:
        analysis = await legal_agent.run_agent(
            question=f"Phân tích chi tiết và rà soát toàn diện hợp đồng: {contract['name']}. Đánh giá rủi ro, điều khoản quan trọng, vấn đề cần lưu ý, và đề xuất cụ thể.",
            company_id=company_id
        )
    except Exception as e:
        print(f"AI analysis error: {e}")
        raise HTTPException(502, "Lỗi phân tích AI")

    # Generate Word document
    doc = DocxDocument()

    # Title
    title = doc.add_heading('BÁO CÁO RÀ SOÁT HỢP ĐỒNG', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contract info table
    doc.add_heading('I. Thông tin hợp đồng', level=1)
    table = doc.add_table(rows=5, cols=2)
    try:
        table.style = 'Light Grid Accent 1'
    except Exception:
        pass  # Style may not exist in all templates

    # Parse parties safely
    parties_raw = contract.get('parties', [])
    if isinstance(parties_raw, str):
        try:
            parties_raw = json.loads(parties_raw)
        except Exception:
            parties_raw = [parties_raw]
    parties_str = ', '.join(str(p) for p in (parties_raw or ['N/A']))

    cells = [
        ('Tên hợp đồng', str(contract.get('name', 'N/A'))),
        ('Các bên', parties_str),
        ('Ngày bắt đầu', str(contract.get('start_date', 'N/A'))),
        ('Ngày kết thúc', str(contract.get('end_date', 'N/A'))),
        ('Loại hợp đồng', str(contract.get('contract_type', 'N/A'))),
    ]
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # AI Analysis
    doc.add_heading('II. Kết quả rà soát AI', level=1)
    answer = analysis.get('answer', '')
    for line in answer.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
        else:
            clean = line.replace('**', '').replace('*', '').replace('`', '')
            doc.add_paragraph(clean)

    # Citations
    citations = analysis.get('citations', [])
    if citations:
        doc.add_heading('III. Căn cứ pháp lý', level=1)
        for c in citations[:10]:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{c.get('source', '')} ({c.get('law_number', '')})")
            run.bold = True
            if c.get('content'):
                p.add_run(f"\n{str(c['content'])[:200]}...")

    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run('⚠️ Đây là báo cáo phân tích tự động bằng AI. Vui lòng tham vấn luật sư trực tiếp cho vụ việc cụ thể.')
    run.italic = True
    run.font.size = Pt(9)

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = str(contract.get('name', 'contract'))[:30].replace(' ', '_')
    filename = f"Bao_cao_ra_soat_{safe_name}.docx"

    # Audit log for report export
    log_audit(company_id, str(company.get("user_id")) if company.get("user_id") else None, "export", "report", contract_id)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ============================================
# Feature: Multi-Contract Risk Dashboard
# ============================================

@app.get("/v1/contracts/risk-overview")
async def contract_risk_overview(company: dict = Depends(verify_api_key)):
    """Get risk overview for all company contracts"""
    from datetime import date
    company_id = str(company["company_id"])

    try:
        with get_db() as conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute("""
                SELECT id, name, contract_type, start_date, end_date, status,
                       parties, metadata
                FROM contracts
                WHERE company_id = %s AND status != 'deleted'
                ORDER BY created_at DESC
            """, (company_id,))
            contracts = cur.fetchall()
    except Exception as e:
        print(f"DB error in risk overview: {e}")
        raise HTTPException(500, "Lỗi truy vấn database")

    overview = {
        "total": len(contracts),
        "by_status": {},
        "by_type": {},
        "expiring_soon": [],
        "expired": [],
        "missing_end_date": [],
        "risk_items": []
    }

    today = date.today()

    for c in contracts:
        c = dict(c)
        # Status count
        s = c.get('status', 'unknown') or 'unknown'
        overview["by_status"][s] = overview["by_status"].get(s, 0) + 1

        # Type count
        t = c.get('contract_type', 'Khác') or 'Khác'
        overview["by_type"][t] = overview["by_type"].get(t, 0) + 1

        # Parse parties
        parties_raw = c.get("parties", [])
        if isinstance(parties_raw, str):
            try:
                parties_raw = json.loads(parties_raw)
            except Exception:
                parties_raw = [parties_raw]

        contract_info = {
            "id": str(c["id"]),
            "name": c.get("name", "N/A"),
            "end_date": str(c.get("end_date")) if c.get("end_date") else None,
            "parties": parties_raw or []
        }

        if c.get("end_date"):
            try:
                days_left = (c["end_date"] - today).days
                if days_left < 0:
                    overview["expired"].append({**contract_info, "days_overdue": abs(days_left)})
                    overview["risk_items"].append({
                        "type": "expired",
                        "severity": "critical",
                        "contract": contract_info,
                        "message": f"Đã hết hạn {abs(days_left)} ngày"
                    })
                elif days_left <= 30:
                    overview["expiring_soon"].append({**contract_info, "days_left": days_left})
                    overview["risk_items"].append({
                        "type": "expiring",
                        "severity": "warning" if days_left > 7 else "critical",
                        "contract": contract_info,
                        "message": f"Còn {days_left} ngày"
                    })
            except Exception:
                pass
        else:
            overview["missing_end_date"].append(contract_info)
            overview["risk_items"].append({
                "type": "no_end_date",
                "severity": "info",
                "contract": contract_info,
                "message": "Không có ngày kết thúc"
            })

    # Sort risk items by severity
    severity_order = {"critical": 0, "warning": 1, "info": 2}
    overview["risk_items"].sort(key=lambda x: severity_order.get(x["severity"], 3))

    return overview


# ============================================
# Feature: Chat History Search
# ============================================

@app.get("/v1/chats/search")
async def search_chat_history(
    q: str = Query(..., min_length=1),
    limit: int = Query(20, ge=1, le=100),
    company: dict = Depends(verify_api_key)
):
    """Search through chat history"""
    company_id = str(company["company_id"])

    try:
        with get_db() as conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            # Search in usage_logs - handle missing columns gracefully
            try:
                cur.execute("""
                    SELECT u.id, u.question, u.answer, u.created_at, u.session_id,
                           u.response_time_ms, u.citations_count
                    FROM usage_logs u
                    WHERE u.company_id = %s
                    AND (u.question ILIKE %s OR u.answer ILIKE %s)
                    ORDER BY u.created_at DESC
                    LIMIT %s
                """, (company_id, f"%{q}%", f"%{q}%", limit))
                results = cur.fetchall()
            except Exception:
                # Fallback: some columns may not exist, try simpler query
                conn.rollback()
                cur.execute("""
                    SELECT m.id, m.content, m.role, m.created_at, m.session_id
                    FROM messages m
                    WHERE m.company_id = %s
                    AND m.content ILIKE %s
                    ORDER BY m.created_at DESC
                    LIMIT %s
                """, (company_id, f"%{q}%", limit))
                raw = cur.fetchall()
                results = []
                for r in raw:
                    results.append({
                        "id": r["id"],
                        "question": r["content"] if r.get("role") == "user" else None,
                        "answer": r["content"] if r.get("role") == "assistant" else None,
                        "created_at": r.get("created_at"),
                        "session_id": r.get("session_id"),
                        "response_time_ms": None,
                        "citations_count": None
                    })
    except Exception as e:
        print(f"Search chat error: {e}")
        raise HTTPException(500, "Lỗi tìm kiếm")

    return {
        "query": q,
        "results": [{
            "id": str(r["id"]) if r.get("id") else None,
            "question": r.get("question"),
            "answer_preview": (r.get("answer") or "")[:200] + "..." if r.get("answer") and len(r.get("answer", "")) > 200 else r.get("answer", ""),
            "created_at": r["created_at"].isoformat() if r.get("created_at") and hasattr(r["created_at"], 'isoformat') else str(r.get("created_at", "")),
            "session_id": str(r.get("session_id")) if r.get("session_id") else None,
            "response_time": r.get("response_time_ms")
        } for r in results],
        "total": len(results)
    }


# ============================================
# Feature: Enhanced Contract Comparison
# ============================================

@app.post("/v1/contracts/compare-detailed")
async def compare_contracts_detailed(req: ContractCompareRequest, company: dict = Depends(verify_api_key)):
    """Enhanced contract comparison with detailed side-by-side analysis"""
    check_rate_limit(str(company["company_id"]))

    if len(req.contract_ids) < 2:
        raise HTTPException(status_code=400, detail="Cần ít nhất 2 hợp đồng để so sánh")
    if len(req.contract_ids) > 5:
        raise HTTPException(status_code=400, detail="Tối đa 5 hợp đồng")

    contracts_data = []
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        for cid in req.contract_ids:
            cur.execute("""
                SELECT id, name, contract_type, extracted_text, parties,
                       start_date, end_date, value, status, notes
                FROM contracts
                WHERE id::text = %s AND company_id = %s AND status != 'deleted'
            """, (cid, company["company_id"]))
            contract = cur.fetchone()
            if not contract:
                raise HTTPException(status_code=404, detail=f"Không tìm thấy hợp đồng: {cid}")
            c = dict(contract)
            for key in ["start_date", "end_date"]:
                if c.get(key):
                    c[key] = str(c[key])
            contracts_data.append(c)

    # Build detailed comparison prompt
    contracts_text = ""
    for i, c in enumerate(contracts_data, 1):
        text = (c.get("extracted_text") or "")[:10000]
        parties_raw = c.get("parties", [])
        if isinstance(parties_raw, str):
            try:
                parties_raw = json.loads(parties_raw)
            except Exception:
                parties_raw = [parties_raw]
        contracts_text += f"\n\n--- HỢP ĐỒNG {i}: {c['name']} ---\n"
        contracts_text += f"Loại: {c.get('contract_type', 'N/A')}\n"
        contracts_text += f"Các bên: {', '.join(str(p) for p in (parties_raw or []))}\n"
        contracts_text += f"Thời hạn: {c.get('start_date', 'N/A')} → {c.get('end_date', 'N/A')}\n"
        contracts_text += f"Giá trị: {c.get('value', 'N/A')}\n"
        contracts_text += f"NỘI DUNG:\n{text}"

    system_prompt = """Bạn là luật sư chuyên so sánh hợp đồng chi tiết. Phân tích từng khía cạnh.

Trả về JSON format:
{
    "summary": "Tóm tắt so sánh tổng quan",
    "side_by_side": [
        {
            "aspect": "Tên khía cạnh (VD: Điều khoản thanh toán, Phạt vi phạm, Bảo mật...)",
            "contracts": [
                {"name": "Tên HĐ 1", "content": "Nội dung điều khoản", "rating": "good|neutral|risk"},
                {"name": "Tên HĐ 2", "content": "Nội dung điều khoản", "rating": "good|neutral|risk"}
            ],
            "analysis": "Phân tích so sánh",
            "recommendation": "Đề xuất"
        }
    ],
    "risk_matrix": [
        {"contract_name": "...", "overall_risk": "low|medium|high", "risk_score": 1-100, "key_risks": ["..."]}
    ],
    "missing_clauses": [
        {"clause": "Tên điều khoản thiếu", "affected_contracts": ["..."], "importance": "critical|important|optional"}
    ],
    "inconsistencies": ["Điểm không nhất quán"],
    "best_contract": {"name": "...", "reason": "Lý do"},
    "action_items": ["Việc cần làm cụ thể"]
}"""

    user_message = f"""So sánh CHI TIẾT {len(contracts_data)} hợp đồng sau. Phân tích từng điều khoản quan trọng side-by-side:
{contracts_text}

Hãy phân tích:
1. So sánh từng điều khoản quan trọng (thanh toán, phạt vi phạm, bảo mật, chấm dứt HĐ, giải quyết tranh chấp...)
2. Ma trận rủi ro cho từng hợp đồng
3. Điều khoản nào bị thiếu ở hợp đồng nào
4. Điểm không nhất quán giữa các hợp đồng
5. Đề xuất cụ thể"""

    result = await call_claude(system_prompt, user_message, max_tokens=8192)

    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s", (company["company_id"],))
        cur.execute("""
            INSERT INTO usage_logs (company_id, endpoint, agent_type, input_tokens, output_tokens, status_code)
            VALUES (%s, '/v1/contracts/compare-detailed', 'compare', %s, %s, 200)
        """, (company["company_id"], result["input_tokens"], result["output_tokens"]))
        conn.commit()

    try:
        comparison = json.loads(result["content"])
    except Exception:
        comparison = {"raw_analysis": result["content"]}

    return {
        "comparison": comparison,
        "contracts": [{
            "id": str(c["id"]),
            "name": c["name"],
            "contract_type": c.get("contract_type"),
            "start_date": c.get("start_date"),
            "end_date": c.get("end_date"),
            "value": c.get("value")
        } for c in contracts_data],
        "tokens_used": result["input_tokens"] + result["output_tokens"],
        "model": result["model"]
    }


# ============================================
# Feature: Template Auto-Fill with AI
# ============================================

@app.post("/v1/templates/{template_id}/fill")
async def ai_fill_template(
    template_id: str,
    request: dict = Body(...),
    company: dict = Depends(verify_api_key)
):
    """AI fills template with provided data"""
    company_id = str(company["company_id"])
    context = request.get("context", "")

    try:
        with get_db() as conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            # Try both 'id' and 'template_id' columns
            try:
                cur.execute("SELECT * FROM document_templates WHERE template_id = %s", (template_id,))
                template = cur.fetchone()
            except Exception:
                conn.rollback()
                cur.execute("SELECT * FROM document_templates WHERE id::text = %s", (template_id,))
                template = cur.fetchone()

            if not template:
                raise HTTPException(404, "Template không tồn tại")
    except HTTPException:
        raise
    except Exception as e:
        print(f"DB error in template fill: {e}")
        raise HTTPException(500, "Lỗi truy vấn database")

    # Get template content - handle different column names
    template_content = template.get('template_content') or template.get('content') or ''
    template_name = template.get('name', '')

    # Use AI to fill template
    try:
        result = await legal_agent.run_agent(
            question=f"""Điền vào template sau dựa trên thông tin được cung cấp.

TEMPLATE:
{template_content[:10000]}

THÔNG TIN:
{context}

Hãy điền đầy đủ các trường trong template. Giữ nguyên format. Thay [placeholder] bằng nội dung phù hợp.
Nếu thiếu thông tin, đánh dấu [CẦN BỔ SUNG: mô tả].
Trả về toàn bộ nội dung template đã điền.""",
            company_id=company_id
        )
    except Exception as e:
        print(f"AI fill error: {e}")
        raise HTTPException(502, "Lỗi AI khi điền template")

    # Parse [CẦN BỔ SUNG] markers
    filled_content = result.get("answer", "")
    missing_fields = []
    import re as _re
    for match in _re.finditer(r'\[CẦN BỔ SUNG:\s*([^\]]+)\]', filled_content):
        missing_fields.append(match.group(1).strip())

    return {
        "filled_content": filled_content,
        "template_name": template_name,
        "missing_fields": missing_fields
    }


# ============================================
# Feature: Usage Analytics
# ============================================

@app.get("/v1/analytics")
async def get_analytics(
    days: int = Query(30, ge=1, le=365),
    company: dict = Depends(verify_api_key)
):
    """Get usage analytics for the company"""
    company_id = str(company["company_id"])

    try:
        with get_db() as conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)

            # Daily usage
            try:
                cur.execute("""
                    SELECT DATE(created_at) as date, COUNT(*) as queries,
                           AVG(response_time_ms) as avg_response_time,
                           SUM(CASE WHEN citations_count > 0 THEN 1 ELSE 0 END) as queries_with_citations
                    FROM usage_logs
                    WHERE company_id = %s AND created_at >= NOW() - INTERVAL '%s days'
                    GROUP BY DATE(created_at)
                    ORDER BY date DESC
                """, (company_id, days))
                daily = cur.fetchall()
            except Exception:
                conn.rollback()
                # Fallback without optional columns
                cur.execute("""
                    SELECT DATE(created_at) as date, COUNT(*) as queries
                    FROM usage_logs
                    WHERE company_id = %s AND created_at >= NOW() - INTERVAL '%s days'
                    GROUP BY DATE(created_at)
                    ORDER BY date DESC
                """, (company_id, days))
                daily = [dict(r, avg_response_time=0, queries_with_citations=0) for r in cur.fetchall()]

            # Top queries
            try:
                cur.execute("""
                    SELECT question, COUNT(*) as count
                    FROM usage_logs
                    WHERE company_id = %s AND created_at >= NOW() - INTERVAL '%s days'
                    AND question IS NOT NULL
                    GROUP BY question
                    ORDER BY count DESC
                    LIMIT 10
                """, (company_id, days))
                top_queries = cur.fetchall()
            except Exception:
                conn.rollback()
                top_queries = []

            # Summary
            try:
                cur.execute("""
                    SELECT COUNT(*) as total_queries,
                           AVG(response_time_ms) as avg_response_time
                    FROM usage_logs
                    WHERE company_id = %s AND created_at >= NOW() - INTERVAL '%s days'
                """, (company_id, days))
                summary = cur.fetchone()
            except Exception:
                conn.rollback()
                cur.execute("""
                    SELECT COUNT(*) as total_queries
                    FROM usage_logs
                    WHERE company_id = %s AND created_at >= NOW() - INTERVAL '%s days'
                """, (company_id, days))
                row = cur.fetchone()
                summary = {"total_queries": row["total_queries"] if row else 0, "avg_response_time": 0}

    except Exception as e:
        print(f"Analytics error: {e}")
        raise HTTPException(500, "Lỗi truy vấn analytics")

    return {
        "period_days": days,
        "summary": {
            "total_queries": summary["total_queries"] if summary else 0,
            "avg_response_time_ms": round(summary.get("avg_response_time") or 0) if summary else 0,
        },
        "daily": [{
            "date": str(d["date"]),
            "queries": d["queries"],
            "avg_response_time_ms": round(d.get("avg_response_time") or 0),
            "queries_with_citations": d.get("queries_with_citations", 0) or 0
        } for d in daily],
        "top_queries": [{
            "question": (q.get("question") or "")[:100],
            "count": q["count"]
        } for q in top_queries]
    }


# ============================================
# Feature: Onboarding Status
# ============================================

@app.get("/v1/onboarding/status")
async def get_onboarding_status(company: dict = Depends(verify_api_key)):
    """Check what onboarding steps are completed"""
    company_id = str(company["company_id"])

    steps = {
        "profile_complete": False,
        "first_contract": False,
        "first_document": False,
        "first_chat": False,
        "first_search": False,
        "team_setup": False
    }

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Check company profile
        try:
            cur.execute("SELECT name, tax_code FROM companies WHERE id = %s", (company_id,))
            company_data = cur.fetchone()
            if company_data and company_data.get("name"):
                steps["profile_complete"] = True
        except Exception:
            pass

        # Check contracts
        try:
            cur.execute("SELECT COUNT(*) as cnt FROM contracts WHERE company_id = %s AND status != 'deleted'", (company_id,))
            if cur.fetchone()["cnt"] > 0:
                steps["first_contract"] = True
        except Exception:
            pass

        # Check documents
        try:
            cur.execute("SELECT COUNT(*) as cnt FROM documents WHERE company_id = %s", (company_id,))
            if cur.fetchone()["cnt"] > 0:
                steps["first_document"] = True
        except Exception:
            pass

        # Check chat usage
        try:
            cur.execute("SELECT COUNT(*) as cnt FROM usage_logs WHERE company_id = %s", (company_id,))
            result = cur.fetchone()
            if result and result["cnt"] > 0:
                steps["first_chat"] = True
                if result["cnt"] >= 3:
                    steps["first_search"] = True
        except Exception:
            pass

    completed = sum(1 for v in steps.values() if v)
    total = len(steps)

    return {
        "steps": steps,
        "completed": completed,
        "total": total,
        "progress": round(completed / total * 100),
        "is_complete": completed == total
    }


# ============================================
# Feature: Contract Versioning
# ============================================

@app.post("/v1/contracts/{contract_id}/versions")
async def upload_contract_version(
    contract_id: str,
    file: UploadFile = File(...),
    company: dict = Depends(verify_api_key)
):
    """Upload a new version of a contract"""
    company_id = str(company["company_id"])

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Verify contract exists
        cur.execute("SELECT * FROM contracts WHERE id = %s AND company_id = %s AND status != 'deleted'", (contract_id, company_id))
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(404, "Hợp đồng không tồn tại")

        # Get current version number
        current_version = contract.get("metadata", {}).get("version", 1) if contract.get("metadata") else 1
        new_version = current_version + 1

        # Save new file
        file_ext = os.path.splitext(file.filename)[1]
        new_file_id = str(uuid.uuid4())
        file_path = f"uploads/{new_file_id}{file_ext}"

        os.makedirs("uploads", exist_ok=True)
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)

        # Extract text from new version
        text_content = ""
        try:
            if file_ext.lower() == '.docx':
                from docx import Document as DocxDocument
                doc = DocxDocument(BytesIO(content))
                text_content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            elif file_ext.lower() == '.pdf':
                # FIX 15: Use pypdf instead of PyPDF2
                from pypdf import PdfReader  # type: ignore
                reader = PdfReader(BytesIO(content))
                text_content = "\n".join(page.extract_text() or "" for page in reader.pages)
            elif file_ext.lower() == '.txt':
                text_content = content.decode('utf-8', errors='ignore')
        except Exception as e:
            text_content = f"[Lỗi đọc file: {str(e)}]"

        # Update contract with new version info
        metadata = contract.get("metadata") or {}
        version_history = metadata.get("versions", [])
        version_history.append({
            "version": current_version,
            "file_id": contract.get("file_id", ""),
            "uploaded_at": contract.get("updated_at", "").isoformat() if contract.get("updated_at") else None,
            "filename": contract.get("original_filename", "")
        })
        metadata["version"] = new_version
        metadata["versions"] = version_history

        cur.execute("""
            UPDATE contracts
            SET file_id = %s, content = %s, original_filename = %s,
                metadata = %s, updated_at = NOW()
            WHERE id = %s AND company_id = %s
        """, (new_file_id, text_content[:50000], file.filename,
              json.dumps(metadata), contract_id, company_id))
        conn.commit()

    return {
        "version": new_version,
        "file_id": new_file_id,
        "filename": file.filename,
        "previous_versions": len(version_history)
    }


@app.get("/v1/contracts/{contract_id}/versions")
async def get_contract_versions(contract_id: str, company: dict = Depends(verify_api_key)):
    """Get version history of a contract"""
    company_id = str(company["company_id"])

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT metadata, name, original_filename, updated_at FROM contracts WHERE id = %s AND company_id = %s AND status != 'deleted'", (contract_id, company_id))
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(404, "Hợp đồng không tồn tại")

    metadata = contract.get("metadata") or {}
    versions = metadata.get("versions", [])
    current_version = metadata.get("version", 1)

    # Add current version
    versions.append({
        "version": current_version,
        "filename": contract.get("original_filename", ""),
        "uploaded_at": contract.get("updated_at", "").isoformat() if contract.get("updated_at") else None,
        "is_current": True
    })

    return {
        "contract_name": contract["name"],
        "current_version": current_version,
        "versions": sorted(versions, key=lambda x: x.get("version", 0), reverse=True)
    }


# ============================================
# Feature: Contract Notes/Comments
# ============================================

@app.post("/v1/contracts/{contract_id}/notes")
async def add_contract_note(
    contract_id: str,
    request: dict = Body(...),
    company: dict = Depends(verify_api_key)
):
    """Add a note/comment to a contract"""
    company_id = str(company["company_id"])
    user_id = company.get("user_id")
    note_text = request.get("text", "").strip()
    note_type = request.get("type", "note")  # note, warning, action, resolved

    if not note_text:
        raise HTTPException(400, "Ghi chú không được trống")

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Verify contract
        cur.execute("SELECT metadata FROM contracts WHERE id = %s AND company_id = %s AND status != 'deleted'", (contract_id, company_id))
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(404, "Hợp đồng không tồn tại")

        metadata = contract.get("metadata") or {}
        notes = metadata.get("notes", [])

        note = {
            "id": str(uuid.uuid4()),
            "text": note_text,
            "type": note_type,
            "user_id": str(user_id) if user_id else None,
            "created_at": datetime.utcnow().isoformat(),
            "resolved": False
        }
        notes.append(note)
        metadata["notes"] = notes

        cur.execute("UPDATE contracts SET metadata = %s WHERE id = %s", (json.dumps(metadata), contract_id))
        conn.commit()

    return {"note": note, "total_notes": len(notes)}


@app.get("/v1/contracts/{contract_id}/notes")
async def get_contract_notes(contract_id: str, company: dict = Depends(verify_api_key)):
    """Get all notes for a contract"""
    company_id = str(company["company_id"])

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT metadata FROM contracts WHERE id = %s AND company_id = %s AND status != 'deleted'", (contract_id, company_id))
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(404)

    metadata = contract.get("metadata") or {}
    notes = metadata.get("notes", [])

    return {"notes": sorted(notes, key=lambda x: x.get("created_at", ""), reverse=True)}


@app.put("/v1/contracts/{contract_id}/notes/{note_id}")
async def update_contract_note(
    contract_id: str,
    note_id: str,
    request: dict = Body(...),
    company: dict = Depends(verify_api_key)
):
    """Update or resolve a note"""
    company_id = str(company["company_id"])

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT metadata FROM contracts WHERE id = %s AND company_id = %s AND status != 'deleted'", (contract_id, company_id))
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(404)

        metadata = contract.get("metadata") or {}
        notes = metadata.get("notes", [])

        for note in notes:
            if note.get("id") == note_id:
                if "text" in request:
                    note["text"] = request["text"]
                if "resolved" in request:
                    note["resolved"] = request["resolved"]
                if "type" in request:
                    note["type"] = request["type"]
                break

        metadata["notes"] = notes
        cur.execute("UPDATE contracts SET metadata = %s WHERE id = %s", (json.dumps(metadata), contract_id))
        conn.commit()

    return {"success": True}


# ============================================
# Feature: AI Proactive Insights
# ============================================

@app.get("/v1/insights")
async def get_ai_insights(company: dict = Depends(verify_api_key)):
    """Generate proactive AI insights about the company's legal state"""
    company_id = str(company["company_id"])

    insights = []

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Expiring contracts insight
        try:
            cur.execute("""
                SELECT COUNT(*) as cnt FROM contracts
                WHERE company_id = %s AND status != 'deleted'
                AND end_date IS NOT NULL AND end_date <= CURRENT_DATE + INTERVAL '60 days'
                AND end_date >= CURRENT_DATE
            """, (company_id,))
            expiring = cur.fetchone()["cnt"]
            if expiring > 0:
                insights.append({
                    "type": "expiring_contracts",
                    "severity": "warning",
                    "icon": "⏰",
                    "title": f"{expiring} hợp đồng sắp hết hạn trong 60 ngày",
                    "description": "Nên rà soát và chuẩn bị gia hạn hoặc ký mới.",
                    "action": {"type": "open_risk_dashboard"}
                })
        except Exception:
            pass

        # Unreviewed contracts
        try:
            cur.execute("""
                SELECT COUNT(*) as cnt FROM contracts
                WHERE company_id = %s AND status = 'uploaded'
            """, (company_id,))
            unreviewed = cur.fetchone()["cnt"]
            if unreviewed > 0:
                insights.append({
                    "type": "unreviewed",
                    "severity": "info",
                    "icon": "📋",
                    "title": f"{unreviewed} hợp đồng chưa được rà soát AI",
                    "description": "Rà soát bằng AI để phát hiện rủi ro sớm.",
                    "action": {"type": "open_contracts"}
                })
        except Exception:
            pass

        # Usage trend
        try:
            cur.execute("""
                SELECT COUNT(*) as cnt FROM usage_logs
                WHERE company_id = %s AND created_at >= NOW() - INTERVAL '7 days'
            """, (company_id,))
            recent_usage = cur.fetchone()["cnt"]

            cur.execute("""
                SELECT COUNT(*) as cnt FROM usage_logs
                WHERE company_id = %s
                AND created_at >= NOW() - INTERVAL '14 days'
                AND created_at < NOW() - INTERVAL '7 days'
            """, (company_id,))
            prev_usage = cur.fetchone()["cnt"]

            if prev_usage > 0 and recent_usage > prev_usage * 1.5:
                insights.append({
                    "type": "usage_increase",
                    "severity": "success",
                    "icon": "📈",
                    "title": "Sử dụng tăng mạnh tuần này",
                    "description": f"Tuần này: {recent_usage} câu hỏi (tuần trước: {prev_usage})",
                    "action": {"type": "open_analytics"}
                })
        except Exception:
            pass

        # Missing company profile
        try:
            cur.execute("SELECT name, tax_code, address FROM companies WHERE id = %s", (company_id,))
            profile = cur.fetchone()
            if profile and (not profile.get("tax_code") or not profile.get("address")):
                insights.append({
                    "type": "incomplete_profile",
                    "severity": "info",
                    "icon": "🏢",
                    "title": "Hồ sơ công ty chưa đầy đủ",
                    "description": "Cập nhật mã số thuế và địa chỉ để AI tư vấn chính xác hơn.",
                    "action": {"type": "open_settings"}
                })
        except Exception:
            pass

    return {
        "insights": insights,
        "generated_at": datetime.utcnow().isoformat()
    }


# ============================================
# Feature: Export Chat as Markdown
# ============================================

@app.get("/v1/chats/{session_id}/export")
async def export_chat(session_id: str, company: dict = Depends(verify_api_key)):
    """Export a chat session as markdown"""
    company_id = str(company["company_id"])

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        # Try usage_logs first, fall back to messages table
        messages = []
        try:
            cur.execute("""
                SELECT question, answer, created_at, citations_count
                FROM usage_logs
                WHERE company_id = %s AND session_id = %s
                ORDER BY created_at ASC
            """, (company_id, session_id))
            messages = cur.fetchall()
        except Exception:
            conn.rollback()

        # Fallback: try messages table
        if not messages:
            try:
                cur.execute("""
                    SELECT role, content, created_at
                    FROM messages
                    WHERE company_id = %s AND session_id = %s
                    ORDER BY created_at ASC
                """, (company_id, session_id))
                raw_msgs = cur.fetchall()
                # Pair user/assistant messages
                i = 0
                while i < len(raw_msgs):
                    msg = raw_msgs[i]
                    if msg["role"] == "user":
                        answer = ""
                        citations_count = 0
                        if i + 1 < len(raw_msgs) and raw_msgs[i + 1]["role"] == "assistant":
                            answer = raw_msgs[i + 1]["content"]
                            i += 1
                        messages.append({
                            "question": msg["content"],
                            "answer": answer,
                            "created_at": msg["created_at"],
                            "citations_count": citations_count
                        })
                    i += 1
            except Exception:
                pass

    if not messages:
        raise HTTPException(404, "Không tìm thấy cuộc hội thoại")

    # Build markdown
    md = f"# Cuộc hội thoại - AI Legal Agent\n"
    md += f"**Ngày:** {messages[0]['created_at'].strftime('%d/%m/%Y') if messages[0].get('created_at') and hasattr(messages[0]['created_at'], 'strftime') else 'N/A'}\n"
    md += f"**Số lượt hỏi:** {len(messages)}\n\n---\n\n"

    for msg in messages:
        time_str = msg['created_at'].strftime('%H:%M') if msg.get('created_at') and hasattr(msg['created_at'], 'strftime') else ''
        md += f"## 👤 Người dùng ({time_str})\n{msg.get('question', '')}\n\n"
        md += f"## ⚖️ AI Legal Agent\n{msg.get('answer', '')}\n\n"
        if msg.get('citations_count', 0) and msg['citations_count'] > 0:
            md += f"*({msg['citations_count']} nguồn trích dẫn)*\n\n"
        md += "---\n\n"

    return StreamingResponse(
        BytesIO(md.encode('utf-8')),
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="chat_{session_id[:8]}.md"'}
    )


# ============================================
# Feature: Audit Log API
# ============================================

@app.get("/v1/audit-log")
async def get_audit_log(
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
    action: str = Query(None),
    company: dict = Depends(verify_api_key)
):
    """Get audit log for the company"""
    company_id = str(company["company_id"])
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        query = "SELECT * FROM audit_logs WHERE company_id = %s"
        params = [company_id]
        if action:
            query += " AND action = %s"
            params.append(action)
        query += " ORDER BY created_at DESC LIMIT %s OFFSET %s"
        params.extend([limit, offset])
        cur.execute(query, params)
        logs = cur.fetchall()

    return {"logs": [{
        "id": l["id"],
        "action": l["action"],
        "resource_type": l.get("resource_type"),
        "resource_id": l.get("resource_id"),
        "details": l.get("details", {}),
        "created_at": l["created_at"].isoformat() if l.get("created_at") else None
    } for l in logs], "total": len(logs)}


# ============================================
# Feature: Universal Search
# ============================================

@app.get("/v1/search/all")
async def universal_search(
    q: str = Query(..., min_length=1),
    limit: int = Query(20, ge=1, le=50),
    company: dict = Depends(verify_api_key)
):
    """Search across contracts, documents, laws, and chat history"""
    company_id = str(company["company_id"])
    results = {"contracts": [], "documents": [], "laws": [], "chats": []}

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Search contracts
        try:
            cur.execute("""
                SELECT id, name, contract_type, parties, status
                FROM contracts
                WHERE company_id = %s AND status != 'deleted'
                AND (name ILIKE %s OR content ILIKE %s)
                LIMIT %s
            """, (company_id, f"%{q}%", f"%{q}%", limit // 4 + 1))
            results["contracts"] = [{
                "id": str(r["id"]), "name": r["name"], "type": "contract",
                "subtitle": r.get("contract_type", ""),
                "icon": "📄"
            } for r in cur.fetchall()]
        except: pass

        # Search documents
        try:
            cur.execute("""
                SELECT id, filename, category, doc_status
                FROM documents
                WHERE company_id = %s
                AND (filename ILIKE %s OR content ILIKE %s)
                LIMIT %s
            """, (company_id, f"%{q}%", f"%{q}%", limit // 4 + 1))
            results["documents"] = [{
                "id": str(r["id"]), "name": r["filename"], "type": "document",
                "subtitle": r.get("category", ""),
                "icon": "📁"
            } for r in cur.fetchall()]
        except: pass

        # Search laws
        try:
            cur.execute("""
                SELECT id, title, law_number, effective_date
                FROM law_documents
                WHERE title ILIKE %s OR law_number ILIKE %s
                LIMIT %s
            """, (f"%{q}%", f"%{q}%", limit // 4 + 1))
            results["laws"] = [{
                "id": str(r["id"]), "name": r["title"], "type": "law",
                "subtitle": r.get("law_number", ""),
                "icon": "⚖️"
            } for r in cur.fetchall()]
        except: pass

        # Search chat history
        try:
            cur.execute("""
                SELECT id, question, created_at, session_id
                FROM usage_logs
                WHERE company_id = %s AND question ILIKE %s
                ORDER BY created_at DESC
                LIMIT %s
            """, (company_id, f"%{q}%", limit // 4 + 1))
            results["chats"] = [{
                "id": str(r.get("session_id", r["id"])),
                "name": r["question"][:80], "type": "chat",
                "subtitle": r["created_at"].strftime("%d/%m/%Y %H:%M") if r.get("created_at") else "",
                "icon": "💬"
            } for r in cur.fetchall()]
        except: pass

    total = sum(len(v) for v in results.values())
    all_results = []
    for category in ["contracts", "documents", "laws", "chats"]:
        all_results.extend(results[category])

    return {"query": q, "results": all_results, "by_category": results, "total": total}


# ============================================
# Feature: Batch Contract Upload
# ============================================

@app.post("/v1/contracts/batch-upload")
async def batch_upload_contracts(
    files: List[UploadFile] = File(...),
    company: dict = Depends(verify_api_key)
):
    """Upload multiple contracts at once"""
    company_id = str(company["company_id"])
    user_id = company.get("user_id")

    results = []
    allowed_ext = {'.pdf', '.docx', '.doc', '.txt'}

    for file in files[:10]:  # Max 10 files
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in allowed_ext:
            results.append({"filename": file.filename, "status": "error", "error": f"Loại file không hỗ trợ: {file_ext}"})
            continue

        try:
            content = await file.read()
            if len(content) > 10 * 1024 * 1024:
                results.append({"filename": file.filename, "status": "error", "error": "File quá lớn (>10MB)"})
                continue

            # Extract text
            text_content = ""
            try:
                if file_ext == '.docx':
                    from docx import Document as DocxDocument
                    doc = DocxDocument(BytesIO(content))
                    text_content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                elif file_ext == '.pdf':
                    # FIX 15: Use pypdf instead of PyPDF2
                    from pypdf import PdfReader  # type: ignore
                    reader = PdfReader(BytesIO(content))
                    text_content = "\n".join(page.extract_text() or "" for page in reader.pages)
                elif file_ext == '.txt':
                    text_content = content.decode('utf-8', errors='ignore')
            except Exception as e:
                text_content = f"[Lỗi đọc: {e}]"

            # Save file
            file_id = str(uuid.uuid4())
            file_path = f"uploads/{file_id}{file_ext}"
            os.makedirs("uploads", exist_ok=True)
            with open(file_path, "wb") as f:
                f.write(content)

            # Auto-detect contract name from filename
            name = os.path.splitext(file.filename)[0].replace('_', ' ').replace('-', ' ')

            # Insert to DB
            with get_db() as conn:
                cur = conn.cursor(cursor_factory=RealDictCursor)
                contract_id = str(uuid.uuid4())
                cur.execute("""
                    INSERT INTO contracts (id, company_id, name, content, file_id, original_filename, status, created_at, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s, 'uploaded', NOW(), NOW())
                    RETURNING id
                """, (contract_id, company_id, name, text_content[:50000], file_id, file.filename))
                conn.commit()

            # Audit log for each upload
            log_audit(company_id, str(user_id) if user_id else None, "upload", "contract", contract_id, {"filename": file.filename})

            results.append({
                "filename": file.filename,
                "status": "success",
                "contract_id": contract_id,
                "name": name,
                "text_length": len(text_content)
            })

        except Exception as e:
            results.append({"filename": file.filename, "status": "error", "error": str(e)})

    success_count = sum(1 for r in results if r["status"] == "success")
    return {
        "total": len(results),
        "success": success_count,
        "failed": len(results) - success_count,
        "results": results
    }


# ============================================
# Feature: Contract Calendar/Timeline
# ============================================

@app.get("/v1/contracts/calendar")
async def contract_calendar(
    year: int = Query(None),
    month: int = Query(None),
    company: dict = Depends(verify_api_key)
):
    """Get contract events for calendar view"""
    company_id = str(company["company_id"])

    from datetime import date, datetime as dt_cls
    if not year:
        year = date.today().year
    if not month:
        month = date.today().month

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        events = []

        try:
            cur.execute("""
                SELECT id, name, start_date, end_date, contract_type, status
                FROM contracts
                WHERE company_id = %s AND status != 'deleted'
                AND (
                    (EXTRACT(YEAR FROM start_date) = %s AND EXTRACT(MONTH FROM start_date) = %s)
                    OR (EXTRACT(YEAR FROM end_date) = %s AND EXTRACT(MONTH FROM end_date) = %s)
                    OR (start_date <= make_date(%s, %s, 1) AND (end_date IS NULL OR end_date >= make_date(%s, %s, 1)))
                )
                ORDER BY start_date ASC
            """, (company_id, year, month, year, month, year, month, year, month))

            for c in cur.fetchall():
                if c.get("start_date") and c["start_date"].year == year and c["start_date"].month == month:
                    events.append({
                        "date": str(c["start_date"]),
                        "type": "start",
                        "color": "#34d399",
                        "title": f"📗 Bắt đầu: {c['name']}",
                        "contract_id": str(c["id"])
                    })
                if c.get("end_date") and c["end_date"].year == year and c["end_date"].month == month:
                    days_left = (c["end_date"] - date.today()).days
                    color = "#f87171" if days_left <= 0 else "#fbbf24" if days_left <= 30 else "#60a5fa"
                    events.append({
                        "date": str(c["end_date"]),
                        "type": "end",
                        "color": color,
                        "title": f"📕 Kết thúc: {c['name']}",
                        "contract_id": str(c["id"]),
                        "days_left": days_left
                    })
        except Exception as e:
            print(f"Calendar error: {e}")

    return {
        "year": year,
        "month": month,
        "events": events,
        "total": len(events)
    }


# ============================================
# Feature: Data Export — Export all company data
# ============================================

@app.get("/v1/export/all")
async def export_all_data(company: dict = Depends(verify_api_key)):
    """Export all company data as JSON"""
    company_id = str(company["company_id"])
    from datetime import date as date_cls, datetime as dt_cls

    export_data = {"exported_at": datetime.utcnow().isoformat(), "company_id": company_id}

    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Company profile
        try:
            cur.execute("SELECT name, tax_code, address, industry FROM companies WHERE id = %s", (company_id,))
            export_data["company"] = cur.fetchone() or {}
        except: export_data["company"] = {}

        # Contracts
        try:
            cur.execute("""
                SELECT id, name, contract_type, parties, start_date, end_date, status, metadata, created_at
                FROM contracts WHERE company_id = %s AND status != 'deleted'
                ORDER BY created_at DESC
            """, (company_id,))
            contracts = cur.fetchall()
            export_data["contracts"] = [{
                **{k: (str(v) if isinstance(v, (date_cls, dt_cls, uuid.UUID)) else v) for k, v in c.items()}
            } for c in contracts]
        except: export_data["contracts"] = []

        # Documents
        try:
            cur.execute("""
                SELECT id, filename, category, doc_status, created_at
                FROM documents WHERE company_id = %s
                ORDER BY created_at DESC
            """, (company_id,))
            docs = cur.fetchall()
            export_data["documents"] = [{
                **{k: (str(v) if isinstance(v, (date_cls, dt_cls, uuid.UUID)) else v) for k, v in d.items()}
            } for d in docs]
        except: export_data["documents"] = []

        # Usage stats
        try:
            cur.execute("SELECT COUNT(*) as total_queries FROM usage_logs WHERE company_id = %s", (company_id,))
            export_data["usage"] = {"total_queries": cur.fetchone()["total_queries"]}
        except: export_data["usage"] = {}

    content = json.dumps(export_data, ensure_ascii=False, indent=2, default=str)

    # Audit log for data export
    log_audit(company_id, str(company.get("user_id")) if company.get("user_id") else None, "export", "all_data")

    return StreamingResponse(
        BytesIO(content.encode('utf-8')),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="legal_agent_export_{company_id[:8]}.json"'}
    )


# ============================================
# Feature: Contract Version Diff
# ============================================

@app.post("/v1/contracts/{contract_id}/diff")
async def contract_version_diff(
    contract_id: str,
    request: dict = Body(...),
    company: dict = Depends(verify_api_key)
):
    """Compare two versions of a contract or two different contracts"""
    company_id = str(company["company_id"])
    compare_with = request.get("compare_with")  # Another contract ID
    
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Get first contract
        cur.execute("SELECT id, name, content FROM contracts WHERE id = %s AND company_id = %s AND status != 'deleted'", (contract_id, company_id))
        contract1 = cur.fetchone()
        if not contract1:
            raise HTTPException(404, "Hợp đồng không tồn tại")
        
        if compare_with:
            cur.execute("SELECT id, name, content FROM contracts WHERE id = %s AND company_id = %s AND status != 'deleted'", (compare_with, company_id))
            contract2 = cur.fetchone()
            if not contract2:
                raise HTTPException(404, "Hợp đồng so sánh không tồn tại")
        else:
            # Compare with previous version from metadata
            metadata = contract1.get("metadata") or {}
            versions = metadata.get("versions", [])
            if not versions:
                raise HTTPException(400, "Không có phiên bản trước để so sánh")
            contract2 = {"name": f"v{versions[-1].get('version', '?')}", "content": ""}
    
    # Simple diff: split into paragraphs and compare
    import difflib
    
    text1 = (contract1.get("content") or "").strip()
    text2 = (contract2.get("content") or "").strip()
    
    lines1 = text1.split('\n')
    lines2 = text2.split('\n')
    
    differ = difflib.unified_diff(lines2, lines1, lineterm='', 
                                   fromfile=contract2.get("name", "Phiên bản cũ"),
                                   tofile=contract1.get("name", "Phiên bản mới"))
    diff_text = '\n'.join(differ)
    
    # Also compute similarity ratio
    ratio = difflib.SequenceMatcher(None, text1, text2).ratio()
    
    # Count changes
    added = diff_text.count('\n+') - 1  # Exclude the +++ header
    removed = diff_text.count('\n-') - 1  # Exclude the --- header
    
    return {
        "contract1": {"id": str(contract1["id"]), "name": contract1["name"]},
        "contract2": {"id": str(contract2.get("id", "")), "name": contract2.get("name", "")},
        "diff": diff_text,
        "similarity": round(ratio * 100, 1),
        "changes": {"added": max(added, 0), "removed": max(removed, 0)},
    }


# ============================================
# Feature: Smart Suggestions Based on Contract Content
# ============================================

@app.get("/v1/contracts/{contract_id}/suggestions")
async def get_contract_suggestions(contract_id: str, company: dict = Depends(verify_api_key)):
    """AI-powered suggestions for improving a contract"""
    company_id = str(company["company_id"])
    
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT name, content, contract_type FROM contracts WHERE id = %s AND company_id = %s AND status != 'deleted'", (contract_id, company_id))
        contract = cur.fetchone()
    
    if not contract:
        raise HTTPException(404, "Hợp đồng không tồn tại")
    
    content = (contract.get("content") or "")[:5000].lower()
    suggestions = []
    
    # Check for common missing clauses
    clause_checks = [
        {"name": "Bảo mật thông tin", "keywords": ["bảo mật", "confidential", "bí mật"], "importance": "high",
         "suggestion": "Nên thêm điều khoản bảo mật thông tin để bảo vệ dữ liệu kinh doanh."},
        {"name": "Phạt vi phạm", "keywords": ["phạt", "vi phạm", "chế tài"], "importance": "high",
         "suggestion": "Cần quy định mức phạt cụ thể khi vi phạm hợp đồng (tối đa 8% theo Luật TM 2005)."},
        {"name": "Bất khả kháng", "keywords": ["bất khả kháng", "force majeure"], "importance": "medium",
         "suggestion": "Nên có điều khoản bất khả kháng để xử lý các tình huống ngoài kiểm soát."},
        {"name": "Giải quyết tranh chấp", "keywords": ["tranh chấp", "trọng tài", "tòa án"], "importance": "high",
         "suggestion": "Cần quy định cơ quan giải quyết tranh chấp (tòa án hoặc trọng tài)."},
        {"name": "Bồi thường thiệt hại", "keywords": ["bồi thường", "thiệt hại"], "importance": "medium",
         "suggestion": "Nên quy định rõ phạm vi và giới hạn bồi thường thiệt hại."},
        {"name": "Điều khoản chấm dứt", "keywords": ["chấm dứt", "hủy bỏ", "kết thúc hợp đồng"], "importance": "high",
         "suggestion": "Cần quy định điều kiện và thủ tục chấm dứt hợp đồng trước hạn."},
        {"name": "Bảo hành", "keywords": ["bảo hành", "warranty"], "importance": "low",
         "suggestion": "Xem xét thêm điều khoản bảo hành cho sản phẩm/dịch vụ."},
        {"name": "Sở hữu trí tuệ", "keywords": ["sở hữu trí tuệ", "bản quyền", "intellectual property"], "importance": "medium",
         "suggestion": "Nên quy định quyền sở hữu trí tuệ đối với sản phẩm/tài liệu phát sinh."},
    ]
    
    for check in clause_checks:
        found = any(kw in content for kw in check["keywords"])
        if not found:
            suggestions.append({
                "type": "missing_clause",
                "name": check["name"],
                "importance": check["importance"],
                "suggestion": check["suggestion"],
                "action": {"type": "generate_clause", "clause_name": check["name"]}
            })
    
    # Check for vague terms
    vague_terms = ["hợp lý", "phù hợp", "kịp thời", "nhanh chóng", "sớm nhất"]
    found_vague = [t for t in vague_terms if t in content]
    if found_vague:
        suggestions.append({
            "type": "vague_terms",
            "importance": "medium",
            "suggestion": f"Hợp đồng có các thuật ngữ mơ hồ: {', '.join(found_vague)}. Nên thay bằng con số/thời hạn cụ thể.",
            "terms": found_vague
        })
    
    # Check for missing dates/amounts
    import re
    has_amount = bool(re.search(r'\d+[.,]\d+.*?(VNĐ|đồng|USD|vnđ)', content))
    if not has_amount:
        suggestions.append({
            "type": "missing_amount",
            "importance": "high",
            "suggestion": "Không tìm thấy giá trị hợp đồng cụ thể. Cần ghi rõ số tiền và đơn vị tiền tệ."
        })
    
    # Sort by importance
    importance_order = {"high": 0, "medium": 1, "low": 2}
    suggestions.sort(key=lambda x: importance_order.get(x.get("importance", "low"), 3))
    
    return {
        "contract_name": contract["name"],
        "suggestions": suggestions,
        "total": len(suggestions),
        "score": max(0, 100 - len(suggestions) * 10)  # Simple quality score
    }


# ============================================
# Feature: Bulk AI Analysis
# ============================================

@app.post("/v1/contracts/bulk-analyze")
async def bulk_analyze_contracts(
    request: dict = Body(...),
    company: dict = Depends(verify_api_key)
):
    """Analyze multiple contracts and return summary"""
    company_id = str(company["company_id"])
    contract_ids = request.get("contract_ids", [])
    
    if not contract_ids:
        # Analyze all contracts
        with get_db() as conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute("SELECT id FROM contracts WHERE company_id = %s AND status != 'deleted' LIMIT 20", (company_id,))
            contract_ids = [str(r["id"]) for r in cur.fetchall()]
    
    results = []
    for cid in contract_ids[:20]:
        try:
            with get_db() as conn:
                cur = conn.cursor(cursor_factory=RealDictCursor)
                cur.execute("SELECT id, name, content, contract_type, start_date, end_date, parties FROM contracts WHERE id = %s AND company_id = %s AND status != 'deleted'", (cid, company_id))
                contract = cur.fetchone()
            
            if not contract:
                continue
            
            content = (contract.get("content") or "")[:5000].lower()
            
            # Quick risk assessment without AI call
            risk_score = 0
            risks = []
            
            # Check missing essential clauses
            essential = {
                "phạt vi phạm": ["phạt", "vi phạm"],
                "chấm dứt": ["chấm dứt", "hủy bỏ"],
                "tranh chấp": ["tranh chấp", "tòa án", "trọng tài"],
                "bảo mật": ["bảo mật", "confidential"],
            }
            for clause, keywords in essential.items():
                if not any(kw in content for kw in keywords):
                    risk_score += 15
                    risks.append(f"Thiếu điều khoản {clause}")
            
            # Check dates
            from datetime import date
            if contract.get("end_date"):
                days_left = (contract["end_date"] - date.today()).days
                if days_left < 0:
                    risk_score += 30
                    risks.append(f"Đã hết hạn {abs(days_left)} ngày")
                elif days_left <= 30:
                    risk_score += 20
                    risks.append(f"Sắp hết hạn ({days_left} ngày)")
            
            risk_level = "low" if risk_score < 20 else "medium" if risk_score < 50 else "high"
            
            results.append({
                "id": str(contract["id"]),
                "name": contract["name"],
                "type": contract.get("contract_type", "N/A"),
                "parties": contract.get("parties", []),
                "risk_score": min(risk_score, 100),
                "risk_level": risk_level,
                "risks": risks,
                "start_date": str(contract.get("start_date", "")),
                "end_date": str(contract.get("end_date", "")),
            })
        except Exception as e:
            print(f"Bulk analyze error for {cid}: {e}")
    
    # Summary
    high_risk = sum(1 for r in results if r["risk_level"] == "high")
    medium_risk = sum(1 for r in results if r["risk_level"] == "medium")
    avg_score = sum(r["risk_score"] for r in results) / len(results) if results else 0
    
    return {
        "analyzed": len(results),
        "summary": {
            "avg_risk_score": round(avg_score),
            "high_risk": high_risk,
            "medium_risk": medium_risk,
            "low_risk": len(results) - high_risk - medium_risk
        },
        "results": sorted(results, key=lambda x: x["risk_score"], reverse=True)
    }


# ============================================
# Feature: Company Profile Update
# ============================================

@app.put("/v1/company/profile")
async def update_company_profile(
    request: dict = Body(...),
    company: dict = Depends(verify_api_key)
):
    """Update company profile"""
    company_id = str(company["company_id"])
    
    allowed_fields = ["name", "tax_code", "address", "industry", "phone", "email", "website", "representative"]
    updates = {k: v for k, v in request.items() if k in allowed_fields and v is not None}
    
    if not updates:
        raise HTTPException(400, "Không có thông tin cần cập nhật")
    
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Check which columns exist
        cur.execute("SELECT column_name FROM information_schema.columns WHERE table_name = 'companies'")
        existing_cols = {r["column_name"] for r in cur.fetchall()}
        
        valid_updates = {k: v for k, v in updates.items() if k in existing_cols}
        
        if not valid_updates:
            # Store in metadata
            cur.execute("SELECT metadata FROM companies WHERE id = %s", (company_id,))
            row = cur.fetchone()
            metadata = (row.get("metadata") if row else None) or {}
            metadata.update(updates)
            cur.execute("UPDATE companies SET metadata = %s WHERE id = %s", (json.dumps(metadata), company_id))
        else:
            set_clause = ", ".join(f"{k} = %s" for k in valid_updates)
            values = list(valid_updates.values()) + [company_id]
            cur.execute(f"UPDATE companies SET {set_clause} WHERE id = %s", values)
        
        conn.commit()
    
    return {"updated": list(updates.keys()), "success": True}

@app.get("/v1/company/profile")
async def get_company_profile(company: dict = Depends(verify_api_key)):
    """Get company profile"""
    company_id = str(company["company_id"])
    
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT * FROM companies WHERE id = %s", (company_id,))
        profile = cur.fetchone()
    
    if not profile:
        raise HTTPException(404, "Công ty không tồn tại")
    
    return {
        "id": str(profile["id"]),
        "name": profile.get("name", ""),
        "tax_code": profile.get("tax_code", ""),
        "address": profile.get("address", ""),
        "industry": profile.get("industry", ""),
        "metadata": profile.get("metadata", {}),
        "created_at": profile["created_at"].isoformat() if profile.get("created_at") else None
    }


if __name__ == "__main__":
    import uvicorn  # type: ignore
    uvicorn.run(app, host="0.0.0.0", port=8080)
