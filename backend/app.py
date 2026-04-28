"""
Flask Web Application for Legal Contract Reviewer
MongoDB-only architecture - No Django, Pure Flask + MongoDB
"""
from flask import Flask, render_template, request, jsonify, session, redirect, send_file, make_response
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import os
import shutil
import uuid
from datetime import datetime, timedelta
import secrets
import sys
import hashlib
import json
from functools import wraps
import time
import atexit

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.workflow.graph import build_graph
from vnpay_utils import create_payment_url, verify_return_signature
from src.resource_config import (
    MAX_FILE_SIZE, ALLOWED_EXTENSIONS, UPLOAD_FOLDER,
    AUTO_CLEANUP_UPLOADS, CLEANUP_AFTER_HOURS, SESSION_LIFETIME,
    ENABLE_RATE_LIMIT, RATE_LIMIT_PER_MINUTE
)

# MongoDB Connection
import pymongo
from pymongo.errors import PyMongoError
from dotenv import load_dotenv

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(BASE_DIR, '.env'))

VNPAY_TMN_CODE = os.getenv('VNPAY_TMN_CODE', '')
VNPAY_HASH_SECRET = os.getenv('VNPAY_HASH_SECRET', '')
VNPAY_PAYMENT_URL = os.getenv('VNPAY_PAYMENT_URL', 'https://sandbox.vnpayment.vn/paymentv2/vpcpay.html')
VNPAY_RETURN_URL = os.getenv('VNPAY_RETURN_URL', 'http://localhost:5000/api/payments/vnpay/return')
VNPAY_IPN_URL = os.getenv('VNPAY_IPN_URL', 'http://localhost:5000/api/payments/vnpay/ipn')
FRONTEND_RETURN_URL = os.getenv('FRONTEND_RETURN_URL', 'http://localhost:3000/payment/return')

# Kết nối MongoDB
try:
    MONGODB_URI = os.getenv('MONGODB_URI', 'mongodb://localhost:27017/')
    MONGODB_DB = os.getenv('MONGODB_DB', 'legal_db')
    
    mongo_client = pymongo.MongoClient(MONGODB_URI, serverSelectionTimeoutMS=5000)
    # Test connection
    mongo_client.server_info()
    mongo_db = mongo_client[MONGODB_DB]
    
    # Collections
    users_collection = mongo_db['users']
    analysis_collection = mongo_db['analysis_history']
    contracts_collection = mongo_db['contracts']
    legal_docs_collection = mongo_db['legal_documents']
    payments_collection = mongo_db['payments']

    # Tạo indexes an toàn: không để lỗi index làm app rơi về in-memory.
    try:
        users_collection.create_index('email', unique=True)
    except Exception as e:
        print(f"⚠️ Email index creation failed: {e}")

    try:
        username_indexes = []
        for idx in users_collection.list_indexes():
            key_items = list(idx.get('key', {}).items())
            if key_items == [('username', 1)]:
                username_indexes.append(idx)

        has_desired_username_index = any(
            idx.get('unique') is True and idx.get('partialFilterExpression') == {'username': {'$exists': True, '$ne': None}}
            for idx in username_indexes
        )

        if not has_desired_username_index:
            for idx in username_indexes:
                name = idx.get('name')
                if name:
                    try:
                        users_collection.drop_index(name)
                    except Exception as e:
                        print(f"⚠️ Failed to drop username index {name}: {e}")

            users_collection.create_index(
                [('username', 1)],
                name='username_unique_non_null',
                unique=True,
                partialFilterExpression={'username': {'$type': 'string'}},
            )
    except Exception as e:
        print(f"⚠️ Username index setup failed: {e}")

    try:
        analysis_collection.create_index([('user', 1), ('timestamp', -1)])
        payments_collection.create_index('txn_ref', unique=True)
    except Exception as e:
        print(f"⚠️ Other indexes creation failed: {e}")
    
    print(f"✅ MongoDB connected: {MONGODB_DB}")
    print(f"✅ Collections: users, analysis_history, contracts, legal_documents")
    MONGODB_CONNECTED = True
except Exception as e:
    print(f"⚠️ MongoDB connection failed: {e}")
    print("⚠️ Will use in-memory storage instead")
    MONGODB_CONNECTED = False
    mongo_db = None
    users_collection = None
    analysis_collection = None
    contracts_collection = None
    legal_docs_collection = None
    payments_collection = None

app = Flask(__name__)
app.secret_key = secrets.token_hex(16)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE
app.config['ALLOWED_EXTENSIONS'] = ALLOWED_EXTENSIONS
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = False  # Set True in production with HTTPS
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=SESSION_LIFETIME)

ALLOWED_ORIGINS = [
    'http://localhost:3000',
    'http://localhost:3001',
    'http://127.0.0.1:3000',
    'http://localhost:5173',
    'http://127.0.0.1:5173',
]

# Enable CORS for frontend - More permissive config
CORS(app, 
    resources={r"/api/*": {"origins": ALLOWED_ORIGINS}},
     supports_credentials=True,
     allow_headers=["Content-Type", "Authorization", "X-Requested-With", "x-csrftoken", "X-CSRFToken"],
     methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
     expose_headers=["Content-Type", "Authorization"])

# Additional CORS headers for all responses
@app.after_request
def after_request(response):
    origin = request.headers.get('Origin')
    if origin in ALLOWED_ORIGINS:
        response.headers['Access-Control-Allow-Origin'] = origin
        response.headers['Access-Control-Allow-Credentials'] = 'true'
        response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization, X-Requested-With, x-csrftoken, X-CSRFToken'
    return response

# Tạo thư mục uploads nếu chưa có
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Mock analysis history
ANALYSIS_HISTORY = []

# Cache cho analysis results
ANALYSIS_CACHE = {}
CACHE_TTL = 3600  # 1 giờ

# Rate limiting
RATE_LIMIT_STORE = {}

# Graph singleton - lazy initialization
_app_graph = None

def get_app_graph():
    """Lazy load graph để tiết kiệm tài nguyên"""
    global _app_graph
    if _app_graph is None:
        print("🔄 Building LangGraph (first time)...")
        _app_graph = build_graph()
    return _app_graph

def cleanup_old_files():
    """Xóa các file uploads cũ để tiết kiệm dung lượng"""
    if not AUTO_CLEANUP_UPLOADS:
        return
    
    now = time.time()
    cutoff = now - (CLEANUP_AFTER_HOURS * 3600)
    
    for filename in os.listdir(UPLOAD_FOLDER):
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        if os.path.isfile(filepath):
            file_time = os.path.getmtime(filepath)
            if file_time < cutoff:
                try:
                    os.remove(filepath)
                    print(f"✓ Cleaned up old file: {filename}")
                except Exception as e:
                    print(f"⚠️ Error cleaning {filename}: {e}")

def rate_limit(f):
    """Decorator để giới hạn số requests"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not ENABLE_RATE_LIMIT:
            return f(*args, **kwargs)
        
        user_id = session.get('user_email', request.remote_addr)
        now = time.time()
        
        if user_id not in RATE_LIMIT_STORE:
            RATE_LIMIT_STORE[user_id] = []
        
        # Lọc requests trong 1 phút gần nhất
        RATE_LIMIT_STORE[user_id] = [
            t for t in RATE_LIMIT_STORE[user_id] 
            if now - t < 60
        ]
        
        if len(RATE_LIMIT_STORE[user_id]) >= RATE_LIMIT_PER_MINUTE:
            return jsonify({
                'success': False, 
                'message': 'Quá nhiều requests. Vui lòng thử lại sau.'
            }), 429
        
        RATE_LIMIT_STORE[user_id].append(now)
        return f(*args, **kwargs)
    
    return decorated_function

def get_cache_key(text):
    """Tạo cache key từ contract text"""
    return hashlib.md5(f"v2:{text}".encode()).hexdigest()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def _normalize_text(text):
    return '\n'.join(line.strip() for line in (text or '').splitlines() if line and line.strip())


def _resolve_tesseract_cmd():
    """Find Tesseract executable path on Windows/other platforms."""
    env_cmd = os.getenv('TESSERACT_CMD', '').strip()
    if env_cmd and os.path.exists(env_cmd):
        return env_cmd

    common_windows_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
    ]
    for cmd_path in common_windows_paths:
        if os.path.exists(cmd_path):
            return cmd_path

    return shutil.which('tesseract')


def _extract_text_from_pdf_with_ocr(filepath, lang='vie+eng'):
    """OCR fallback for scanned PDFs when text layer extraction is weak."""
    try:
        import fitz  # PyMuPDF
        from PIL import Image
        import pytesseract
    except Exception as e:
        raise ValueError(f"Không thể dùng OCR cho PDF scan (thiếu thư viện): {e}")

    tesseract_cmd = _resolve_tesseract_cmd()
    if not tesseract_cmd:
        raise ValueError(
            "Không tìm thấy Tesseract OCR. Hãy cài Tesseract và cấu hình biến môi trường TESSERACT_CMD"
        )

    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    doc = fitz.open(filepath)
    try:
        ocr_text_parts = []
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.frombytes('RGB', [pix.width, pix.height], pix.samples)
            page_text = pytesseract.image_to_string(img, lang=lang, config='--psm 6')
            if page_text and page_text.strip():
                ocr_text_parts.append(page_text.strip())
    finally:
        doc.close()

    ocr_text = _normalize_text('\n'.join(ocr_text_parts))
    if len(ocr_text) < 30:
        raise ValueError('OCR không trích xuất đủ nội dung từ PDF scan')

    return ocr_text


def extract_text_from_file(filepath):
    """Trích xuất text theo đúng định dạng file để tăng độ chính xác phân tích."""
    ext = os.path.splitext(filepath)[1].lower()
    text = ""

    if ext in ('.txt', '.doc'):
        for encoding in ('utf-8', 'latin-1'):
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    text = f.read()
                if text:
                    break
            except Exception:
                continue

    elif ext == '.pdf':
        try:
            import PyPDF2
        except Exception as e:
            raise ValueError(f"Thiếu thư viện đọc PDF: {e}")

        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            pages = []
            for page in reader.pages:
                pages.append((page.extract_text() or '').strip())
            text = '\n'.join(p for p in pages if p)

        # OCR fallback for scanned PDFs without text layer
        normalized_pdf_text = _normalize_text(text)
        if len(normalized_pdf_text) < 80:
            try:
                ocr_text = _extract_text_from_pdf_with_ocr(filepath)
                text = f"{normalized_pdf_text}\n{ocr_text}" if normalized_pdf_text else ocr_text
            except Exception as e:
                if len(normalized_pdf_text) < 30:
                    raise ValueError(
                        f"PDF có thể là file scan ảnh và không OCR được: {e}"
                    )
                text = normalized_pdf_text

    elif ext == '.docx':
        try:
            import docx
        except Exception as e:
            raise ValueError(f"Thiếu thư viện đọc DOCX: {e}")

        document = docx.Document(filepath)
        text = '\n'.join(
            p.text.strip() for p in document.paragraphs
            if p.text and p.text.strip()
        )

    else:
        raise ValueError(f"Định dạng file không được hỗ trợ: {ext}")

    normalized = _normalize_text(text)
    if len(normalized) < 30:
        raise ValueError('Không trích xuất được đủ nội dung hợp đồng để phân tích chuyên sâu')

    return normalized


def get_client_ip():
    """Get client ip with proxy support for VNPay vnp_IpAddr."""
    ip = request.headers.get('X-Forwarded-For', '').split(',')[0].strip()
    if ip:
        return ip
    return request.remote_addr or '127.0.0.1'


def get_plan_price(plan, billing_cycle='monthly'):
    prices = {
        'pro': {'monthly': 299000, 'yearly': 2990000},
        'enterprise': {'monthly': 999000, 'yearly': 9990000},
    }
    plan_prices = prices.get(plan)
    if not plan_prices:
        return 0
    return plan_prices.get(billing_cycle, plan_prices['monthly'])


def serialize_payment(doc):
    """Convert payment document to JSON-safe payload for frontend."""
    return {
        'txn_ref': doc.get('txn_ref', ''),
        'plan': doc.get('plan', ''),
        'billing_cycle': doc.get('billing_cycle', 'monthly'),
        'amount': doc.get('amount', 0),
        'status': doc.get('status', 'pending'),
        'provider': doc.get('provider', 'vnpay'),
        'created_at': doc.get('created_at').isoformat() if isinstance(doc.get('created_at'), datetime) else doc.get('created_at'),
        'updated_at': doc.get('updated_at').isoformat() if isinstance(doc.get('updated_at'), datetime) else doc.get('updated_at'),
        'vnp_response_code': doc.get('vnp_response_code'),
        'vnp_transaction_no': doc.get('vnp_transaction_no'),
    }


def normalize_plan(plan_value):
    return 'free' if str(plan_value or '').strip().lower() == 'free' else 'pro'


def resolve_user_plan(user_doc):
    if not user_doc:
        return 'free'
    return normalize_plan(user_doc.get('plan') or user_doc.get('subscription_tier') or 'free')


def serialize_user(doc):
    """Convert a user document to a JSON-safe payload for admin views."""
    created_at = doc.get('created_at')
    updated_at = doc.get('updated_at')
    last_active = doc.get('last_active')
    raw_plan = str(doc.get('plan') or doc.get('subscription_tier') or 'free').lower()
    plan = normalize_plan(raw_plan)

    return {
        'id': str(doc.get('_id', '')),
        'name': doc.get('full_name') or doc.get('username') or doc.get('email', ''),
        'username': doc.get('username', ''),
        'email': doc.get('email', ''),
        'plan': plan,
        'raw_plan': raw_plan,
        'status': 'active' if doc.get('is_active', True) else 'suspended',
        'is_admin': bool(doc.get('is_admin', False)),
        'auth_provider': doc.get('auth_provider', 'local'),
        'joinDate': created_at.isoformat() if isinstance(created_at, datetime) else created_at,
        'lastActive': last_active.isoformat() if isinstance(last_active, datetime) else last_active,
        'updated_at': updated_at.isoformat() if isinstance(updated_at, datetime) else updated_at,
        'analysisCount': int(doc.get('analysis_count', 0) or 0),
    }


def serialize_analysis_summary(doc):
    """Compact analysis record for admin lists and activity feed."""
    data = doc.get('data', {}) or {}
    timestamp = doc.get('timestamp') or doc.get('created_at') or data.get('upload_time')
    issues = data.get('issues', []) or []
    high_risk = sum(1 for item in issues if isinstance(item, dict) and item.get('severity') == 'high')

    return {
        'id': str(doc.get('_id', '')),
        'user': doc.get('user', ''),
        'fileName': data.get('filename', 'Không rõ tên file'),
        'date': timestamp.isoformat() if isinstance(timestamp, datetime) else timestamp,
        'issues': len(issues),
        'highRisk': high_risk,
        'status': 'completed' if data else 'processing',
        'analysisMode': data.get('analysis_mode', 'unknown'),
    }


def serialize_activity(item):
    """Create a simple admin activity entry from real data."""
    return {
        'user': item.get('user', ''),
        'action': item.get('action', ''),
        'time': item.get('time', ''),
        'type': item.get('type', 'analysis'),
    }


def update_vnpay_payment(txn_ref, params, is_success):
    """Persist VNPay return/IPN response and update subscription on success."""
    payment_doc = None
    if MONGODB_CONNECTED and payments_collection is not None and txn_ref:
        payment_doc = payments_collection.find_one({'txn_ref': txn_ref})

    if MONGODB_CONNECTED and payments_collection is not None and txn_ref:
        payments_collection.update_one(
            {'txn_ref': txn_ref},
            {
                '$set': {
                    'status': 'success' if is_success else 'failed',
                    'updated_at': datetime.now(),
                    'vnp_response_code': params.get('vnp_ResponseCode'),
                    'vnp_transaction_status': params.get('vnp_TransactionStatus'),
                    'vnp_transaction_no': params.get('vnp_TransactionNo'),
                    'vnp_bank_code': params.get('vnp_BankCode'),
                    'vnp_pay_date': params.get('vnp_PayDate'),
                    'signature_valid': True,
                }
            }
        )

    if is_success and payment_doc and MONGODB_CONNECTED and users_collection is not None:
        plan = normalize_plan(payment_doc.get('plan', 'pro'))
        users_collection.update_one(
            {'email': payment_doc['user_email']},
            {
                '$set': {
                    'subscription_tier': plan,
                    'plan': plan,
                    'updated_at': datetime.now(),
                }
            }
        )

    return payment_doc


def build_analysis_issues(result):
    """Build structured issue objects from SVM and research results.

    This avoids relying on emoji parsing from the LLM final report.
    """
    issues = []

    svm_results = result.get('svm_results', {}) or {}
    contract_type = svm_results.get('contract_type', {}) or {}
    risk_assessment = svm_results.get('risk_assessment', {}) or {}
    violation_check = svm_results.get('violation_check', {}) or {}

    predicted_risk = str(risk_assessment.get('predicted_risk', '')).lower()
    risk_confidence = float(risk_assessment.get('confidence', 0) or 0)

    if contract_type.get('predicted_type'):
        issues.append({
            'severity': 'low',
            'title': f"Loại hợp đồng: {contract_type.get('predicted_type')}",
            'description': f"Độ tin cậy: {contract_type.get('confidence', 0):.0%}",
            'reference': '',
            'suggestion': 'Kiểm tra lại loại hợp đồng để áp dụng đúng mẫu và điều khoản tương ứng.'
        })

    if predicted_risk in {'high', 'critical', 'very_high'}:
        issues.append({
            'severity': 'high',
            'title': 'Hợp đồng có rủi ro cao',
            'description': f"Mức rủi ro dự đoán: {risk_assessment.get('predicted_risk', 'N/A')} ({risk_confidence:.0%})",
            'reference': '',
            'suggestion': 'Rà soát kỹ các điều khoản chính và tham khảo luật sư trước khi ký.'
        })
    elif predicted_risk in {'medium', 'moderate'}:
        issues.append({
            'severity': 'medium',
            'title': 'Hợp đồng có rủi ro trung bình',
            'description': f"Mức rủi ro dự đoán: {risk_assessment.get('predicted_risk', 'N/A')} ({risk_confidence:.0%})",
            'reference': '',
            'suggestion': 'Bổ sung/điều chỉnh những điều khoản còn mơ hồ để giảm tranh chấp.'
        })
    elif predicted_risk:
        issues.append({
            'severity': 'low',
            'title': 'Mức rủi ro hiện tại thấp',
            'description': f"Mức rủi ro dự đoán: {risk_assessment.get('predicted_risk', 'N/A')} ({risk_confidence:.0%})",
            'reference': '',
            'suggestion': 'Vẫn nên đọc lại các điều khoản quan trọng trước khi ký.'
        })

    if violation_check.get('has_violation'):
        issues.append({
            'severity': 'high' if float(violation_check.get('violation_probability', 0) or 0) >= 0.7 else 'medium',
            'title': 'Phát hiện dấu hiệu vi phạm',
            'description': f"Xác suất vi phạm: {float(violation_check.get('violation_probability', 0) or 0):.0%}",
            'reference': '',
            'suggestion': 'Loại bỏ hoặc sửa ngay điều khoản có dấu hiệu vi phạm.'
        })

    for item in result.get('research_results', []) or []:
        svm_violation = item.get('svm_violation') or {}
        if not svm_violation.get('has_violation'):
            continue

        violation_probability = float(svm_violation.get('violation_probability', 0) or 0)
        issues.append({
            'severity': 'high' if violation_probability >= 0.7 else 'medium',
            'title': 'Điều khoản có khả năng vi phạm',
            'description': item.get('clause', '')[:220],
            'reference': item.get('laws', '')[:220],
            'suggestion': 'So sánh điều khoản này với quy định pháp luật và chỉnh sửa trước khi ký.'
        })

    # Giới hạn số vấn đề để UI không quá dài.
    return issues[:12]


def summarize_analysis(result, issues):
    """Create a short safety summary from structured analysis results."""
    svm_results = result.get('svm_results', {}) or {}
    risk_assessment = svm_results.get('risk_assessment', {}) or {}
    contract_type = svm_results.get('contract_type', {}) or {}
    violation_check = svm_results.get('violation_check', {}) or {}

    predicted_type = contract_type.get('predicted_type', 'Không xác định')
    predicted_risk = risk_assessment.get('predicted_risk', 'Không xác định')
    has_violation = 'Có' if violation_check.get('has_violation') else 'Không'

    summary = [
        f"Loại hợp đồng: {predicted_type}",
        f"Mức rủi ro: {predicted_risk}",
        f"Khả năng vi phạm: {has_violation}",
        f"Số vấn đề phát hiện: {len(issues)}"
    ]

    safety_reasoning = ' '.join(summary)

    return '\n'.join(summary[:3]), safety_reasoning


def build_heuristic_issues(result):
    """Fallback rule-based issue detection when ML models are not available."""
    issues = []
    clause_sources = result.get('extracted_clauses') or []

    high_keywords = [
        'phạt vi phạm', 'bồi thường', 'đơn phương', 'miễn trách nhiệm',
        'chấm dứt ngay', 'không hoàn lại', 'vô thời hạn', 'phạt', 'vi phạm'
    ]
    medium_keywords = [
        'gia hạn', 'bảo mật', 'thanh toán', 'đặt cọc', 'sửa đổi', 'nghĩa vụ',
        'chuyển giao', 'tranh chấp', 'tiền', 'hoàn trả'
    ]

    for clause in clause_sources[:8]:
        normalized = str(clause).lower()
        severity = None
        reason = None

        if any(keyword in normalized for keyword in high_keywords):
            severity = 'high'
            reason = 'Điều khoản có dấu hiệu rủi ro cao hoặc dễ gây tranh chấp.'
        elif any(keyword in normalized for keyword in medium_keywords):
            severity = 'medium'
            reason = 'Điều khoản có điểm cần rà soát thêm để tránh bất lợi.'

        if severity:
            issues.append({
                'severity': severity,
                'title': 'Điều khoản cần rà soát',
                'description': clause[:220],
                'reference': '',
                'suggestion': reason
            })

    # Nếu không bắt được keyword nào thì vẫn trả về 1 cảnh báo thấp để UI không bị 0 toàn bộ.
    if not issues and clause_sources:
        issues.append({
            'severity': 'low',
            'title': 'Cần rà soát thủ công',
            'description': str(clause_sources[0])[:220],
            'reference': '',
            'suggestion': 'Hệ thống chưa phát hiện rủi ro rõ ràng từ model, nên kiểm tra lại điều khoản quan trọng.'
        })

    return issues

@app.route('/')
def index():
    """Trang chu"""
    return render_template('index.html')

@app.route('/api/login', methods=['POST', 'OPTIONS'])
def login():
    """API đăng nhập với MongoDB"""
    # Handle OPTIONS request for CORS preflight
    if request.method == 'OPTIONS':
        return '', 200
    
    if not MONGODB_CONNECTED:
        return jsonify({'success': False, 'message': 'Database không khả dụng'}), 503
    
    try:
        data = request.json
        email = data.get('email', '').strip()
        password = data.get('password', '')
        
        if not email or not password:
            return jsonify({'success': False, 'message': 'Email và mật khẩu không được để trống'}), 400
        
        # Tìm user trong MongoDB
        user = users_collection.find_one({
            '$or': [
                {'email': email},
                {'username': email}
            ]
        })
        
        if not user:
            return jsonify({'success': False, 'message': 'Email hoặc mật khẩu không đúng'}), 401
        
        # Kiểm tra password
        if check_password_hash(user['password'], password):
            session['user_id'] = str(user['_id'])
            session['user_email'] = user['email']
            session['is_admin'] = user.get('is_admin', False)
            session.permanent = True
            
            return jsonify({
                'success': True,
                'email': user['email'],
                'is_admin': user.get('is_admin', False)
            })
        else:
            return jsonify({'success': False, 'message': 'Email hoặc mật khẩu không đúng'}), 401
            
    except Exception as e:
        print(f"Login error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi đăng nhập: {str(e)}'}), 500

@app.route('/api/csrf/', methods=['GET'])
def get_csrf():
    """CSRF token endpoint for frontend"""
    return jsonify({'status': 'ok'}), 200

@app.route('/api/register', methods=['POST', 'OPTIONS'])
def register():
    """API đăng ký với MongoDB"""
    # Handle OPTIONS request for CORS preflight
    if request.method == 'OPTIONS':
        return '', 200
    if not MONGODB_CONNECTED:
        return jsonify({'success': False, 'message': 'Database không khả dụng'}), 503
    
    try:
        data = request.json
        email = data.get('email', '').strip()
        password = data.get('password', '')
        username = data.get('username', email.split('@')[0] if email else '')
        
        if not email or not password:
            return jsonify({'success': False, 'message': 'Email và mật khẩu không được để trống'}), 400
        
        if len(password) < 6:
            return jsonify({'success': False, 'message': 'Mật khẩu phải có ít nhất 6 ký tự'}), 400
        
        # Kiểm tra email đã tồn tại
        if users_collection.find_one({'email': email}):
            return jsonify({'success': False, 'message': 'Email đã tồn tại'}), 400
        
        # Kiểm tra username đã tồn tại
        if users_collection.find_one({'username': username}):
            # Tạo username unique
            count = users_collection.count_documents({})
            username = f"{username}_{count + 1}"
        
        # Tạo user mới trong MongoDB
        user_doc = {
            'username': username,
            'email': email,
            'password': generate_password_hash(password),
            'is_admin': False,
            'is_active': True,
            'created_at': datetime.now(),
            'updated_at': datetime.now()
        }
        
        result = users_collection.insert_one(user_doc)
        user_id = str(result.inserted_id)
        
        # Tự động đăng nhập sau khi đăng ký
        session['user_id'] = user_id
        session['user_email'] = email
        session['is_admin'] = False
        session.permanent = True
        
        return jsonify({
            'success': True, 
            'message': 'Đăng ký thành công',
            'email': email,
            'user_id': user_id
        })
        
    except Exception as e:
        print(f"Register error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi đăng ký: {str(e)}'}), 500

@app.route('/api/logout', methods=['POST'])
def logout():
    """API dang xuat"""
    session.clear()
    return jsonify({'success': True})

@app.route('/api/google-login', methods=['POST', 'OPTIONS'])
def google_login():
    """Đăng nhập bằng Google ID token và tạo session cookie."""
    if request.method == 'OPTIONS':
        return '', 200

    if not MONGODB_CONNECTED or users_collection is None:
        return jsonify({'success': False, 'error': 'Database không khả dụng'}), 503

    data = request.json or {}
    credential = data.get('credential')
    if not credential:
        return jsonify({'success': False, 'error': 'Thiếu credential từ Google'}), 400

    try:
        from google.oauth2 import id_token as google_id_token
        from google.auth.transport import requests as google_requests
    except Exception:
        return jsonify({'success': False, 'error': 'Thiếu package google-auth trên server'}), 500

    google_client_id = os.getenv('GOOGLE_CLIENT_ID')
    try:
        payload = google_id_token.verify_oauth2_token(
            credential,
            google_requests.Request(),
            google_client_id if google_client_id else None,
        )
    except Exception:
        return jsonify({'success': False, 'error': 'Google token không hợp lệ'}), 401

    issuer = payload.get('iss')
    if issuer not in ('accounts.google.com', 'https://accounts.google.com'):
        return jsonify({'success': False, 'error': 'Nguồn token không hợp lệ'}), 401

    email = (payload.get('email') or '').strip().lower()
    full_name = (payload.get('name') or '').strip()
    email_verified = payload.get('email_verified', False)

    if not email or not email_verified:
        return jsonify({'success': False, 'error': 'Email Google chưa xác thực'}), 401

    try:
        user = users_collection.find_one({'email': email})

        if not user:
            base_username = email.split('@')[0] or 'google_user'
            username = base_username
            suffix = 1
            while users_collection.find_one({'username': username}):
                username = f"{base_username}_{suffix}"
                suffix += 1

            user_doc = {
                'username': username,
                'email': email,
                'password': generate_password_hash(secrets.token_hex(16)),
                'full_name': full_name or username,
                'auth_provider': 'google',
                'is_admin': False,
                'is_active': True,
                'created_at': datetime.now(),
                'updated_at': datetime.now(),
            }
            insert_result = users_collection.insert_one(user_doc)
            user_id = str(insert_result.inserted_id)
        else:
            updates = {
                'auth_provider': 'google',
                'updated_at': datetime.now(),
            }
            if full_name and not user.get('full_name'):
                updates['full_name'] = full_name
            users_collection.update_one({'_id': user['_id']}, {'$set': updates})
            user_id = str(user['_id'])

        session['user_id'] = user_id
        session['user_email'] = email
        session['is_admin'] = False
        session.permanent = True

        return jsonify({
            'success': True,
            'message': 'Đăng nhập Google thành công',
            'user': {
                'email': email,
                'full_name': full_name,
                'is_admin': False,
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'error': f'Lỗi server: {str(e)}'}), 500

@app.route('/api/verify', methods=['GET'])
def verify_session():
    """API kiểm tra session"""
    if 'user_email' in session:
        return jsonify({
            'success': True,
            'user': {
                'email': session['user_email'],
                'is_admin': session.get('is_admin', False)
            }
        })
    return jsonify({'success': False, 'message': 'Chưa đăng nhập'}), 401


@app.route('/api/profile', methods=['GET', 'OPTIONS'])
@app.route('/api/profile/', methods=['GET', 'OPTIONS'])
def get_profile():
    """Get current user profile for frontend."""
    if request.method == 'OPTIONS':
        return '', 200

    if 'user_email' not in session:
        return jsonify({'success': False, 'error': 'Chưa đăng nhập'}), 401

    if not MONGODB_CONNECTED or users_collection is None:
        return jsonify({'success': False, 'error': 'Database không khả dụng'}), 503

    user = users_collection.find_one({'email': session['user_email']})
    if not user:
        return jsonify({'success': False, 'error': 'Không tìm thấy người dùng'}), 404

    plan = resolve_user_plan(user)
    analysis_limit = 50 if plan == 'pro' else 5
    usage_this_month = 0

    if analysis_collection is not None:
        now = datetime.now()
        start_of_month = datetime(now.year, now.month, 1)
        try:
            usage_this_month = int(analysis_collection.count_documents({
                'user': session['user_email'],
                'timestamp': {'$gte': start_of_month}
            }))
        except Exception:
            usage_this_month = 0

    profile = {
        'full_name': user.get('full_name', ''),
        'email': user.get('email', ''),
        'phone': user.get('phone', ''),
        'avatar': user.get('avatar', ''),
        'plan': plan,
        'subscription_tier': plan,
        'analysis_limit': analysis_limit,
        'usage_this_month': usage_this_month,
        'remaining_analyses': max(0, analysis_limit - usage_this_month),
        'is_admin': bool(user.get('is_admin', False)),
    }
    return jsonify({'success': True, 'profile': profile}), 200


@app.route('/api/profile', methods=['PUT', 'OPTIONS'])
@app.route('/api/profile/', methods=['PUT', 'OPTIONS'])
def update_profile():
    """Update current user profile fields."""
    if request.method == 'OPTIONS':
        return '', 200

    if 'user_email' not in session:
        return jsonify({'success': False, 'error': 'Chưa đăng nhập'}), 401

    if not MONGODB_CONNECTED or users_collection is None:
        return jsonify({'success': False, 'error': 'Database không khả dụng'}), 503

    data = request.json or {}
    updates = {}

    if 'full_name' in data:
        updates['full_name'] = (data.get('full_name') or '').strip()
    if 'phone' in data:
        updates['phone'] = (data.get('phone') or '').strip()

    if not updates:
        return jsonify({'success': False, 'error': 'Không có dữ liệu để cập nhật'}), 400

    updates['updated_at'] = datetime.now()
    users_collection.update_one(
        {'email': session['user_email']},
        {'$set': updates},
    )

    return jsonify({'success': True, 'message': 'Cập nhật hồ sơ thành công'}), 200


@app.route('/api/upload-avatar/', methods=['POST', 'OPTIONS'])
def upload_avatar():
    """Upload user avatar and persist its path in MongoDB."""
    if request.method == 'OPTIONS':
        return '', 200

    if 'user_email' not in session:
        return jsonify({'success': False, 'error': 'Chưa đăng nhập'}), 401

    if not MONGODB_CONNECTED or users_collection is None:
        return jsonify({'success': False, 'error': 'Database không khả dụng'}), 503

    if 'avatar' not in request.files:
        return jsonify({'success': False, 'error': 'Không tìm thấy file avatar'}), 400

    file = request.files['avatar']
    if not file or file.filename == '':
        return jsonify({'success': False, 'error': 'Không có file được chọn'}), 400

    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
    file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
    if file_ext not in allowed_extensions:
        return jsonify({'success': False, 'error': 'Định dạng file không hợp lệ. Chỉ chấp nhận: png, jpg, jpeg, gif, webp'}), 400

    avatars_dir = os.path.join(app.root_path, 'static', 'avatars')
    os.makedirs(avatars_dir, exist_ok=True)

    user = users_collection.find_one({'email': session['user_email']})
    if not user:
        return jsonify({'success': False, 'error': 'Không tìm thấy người dùng'}), 404

    old_avatar = user.get('avatar', '')
    filename = f"{uuid.uuid4().hex}.{file_ext}"
    filepath = os.path.join(avatars_dir, filename)
    file.save(filepath)

    avatar_url = f"/static/avatars/{filename}"
    result = users_collection.update_one(
        {'email': session['user_email']},
        {'$set': {'avatar': avatar_url, 'updated_at': datetime.now()}},
    )

    if old_avatar and old_avatar.startswith('/static/avatars/'):
        old_filepath = os.path.join(app.root_path, old_avatar.lstrip('/'))
        if os.path.exists(old_filepath):
            try:
                os.remove(old_filepath)
            except Exception as e:
                print(f"⚠️ Failed to delete old avatar: {e}")

    if result.modified_count == 0:
        return jsonify({'success': False, 'error': 'Không thể cập nhật avatar'}), 500

    return jsonify({'success': True, 'message': 'Tải lên avatar thành công', 'avatar_url': avatar_url}), 200


@app.route('/api/delete-avatar/', methods=['DELETE', 'OPTIONS'])
def delete_avatar():
    """Delete the current user's avatar and clear it from MongoDB."""
    if request.method == 'OPTIONS':
        return '', 200

    if 'user_email' not in session:
        return jsonify({'success': False, 'error': 'Chưa đăng nhập'}), 401

    if not MONGODB_CONNECTED or users_collection is None:
        return jsonify({'success': False, 'error': 'Database không khả dụng'}), 503

    user = users_collection.find_one({'email': session['user_email']})
    if not user:
        return jsonify({'success': False, 'error': 'Không tìm thấy người dùng'}), 404

    avatar = user.get('avatar', '')
    if avatar and avatar.startswith('/static/avatars/'):
        avatar_path = os.path.join(app.root_path, avatar.lstrip('/'))
        if os.path.exists(avatar_path):
            try:
                os.remove(avatar_path)
            except Exception as e:
                print(f"⚠️ Failed to remove avatar file: {e}")

    result = users_collection.update_one(
        {'email': session['user_email']},
        {'$set': {'avatar': '', 'updated_at': datetime.now()}},
    )

    if result.modified_count == 0:
        return jsonify({'success': False, 'error': 'Không thể xóa avatar'}), 500

    return jsonify({'success': True, 'message': 'Đã xóa avatar'}), 200


@app.route('/api/payments/vnpay/create', methods=['POST'])
def create_vnpay_payment():
    """Create VNPay payment URL and persist pending transaction."""
    if 'user_email' not in session:
        return jsonify({'success': False, 'error': 'Vui lòng đăng nhập'}), 401

    if not (MONGODB_CONNECTED and payments_collection is not None):
        return jsonify({
            'success': False,
            'error': 'Database không khả dụng, chưa thể tạo giao dịch thanh toán'
        }), 503

    payload = request.json or {}
    plan = payload.get('plan', 'pro')
    billing_cycle = payload.get('billing_cycle', 'monthly')
    amount = payload.get('amount')
    frontend_return_url = payload.get('return_url', FRONTEND_RETURN_URL)

    now = datetime.now()
    txn_ref = f"{int(now.timestamp())}{secrets.randbelow(100000):05d}"

    try:
        amount = int(amount) if amount is not None else get_plan_price(plan, billing_cycle)
    except (TypeError, ValueError):
        return jsonify({'success': False, 'error': 'Số tiền không hợp lệ'}), 400

    if amount <= 0:
        return jsonify({'success': False, 'error': 'Số tiền thanh toán phải lớn hơn 0'}), 400

    if not all([VNPAY_TMN_CODE, VNPAY_HASH_SECRET]):
        failed_doc = {
            'txn_ref': txn_ref,
            'user_email': session['user_email'],
            'plan': plan,
            'billing_cycle': billing_cycle,
            'amount': amount,
            'status': 'failed',
            'provider': 'vnpay',
            'failure_reason': 'missing_vnpay_config',
            'failure_message': 'Thiếu VNPAY_TMN_CODE hoặc VNPAY_HASH_SECRET',
            'return_url': frontend_return_url,
            'vnpay_return_url': VNPAY_RETURN_URL,
            'created_ip': get_client_ip(),
            'created_at': now,
            'updated_at': now,
        }
        try:
            payments_collection.insert_one(failed_doc)
        except Exception as e:
            return jsonify({'success': False, 'error': f'Lỗi lưu giao dịch: {str(e)}'}), 500

        return jsonify({
            'success': False,
            'txn_ref': txn_ref,
            'error': 'Chưa cấu hình VNPAY_TMN_CODE hoặc VNPAY_HASH_SECRET'
        }), 500

    order_info = f"GOI_{str(plan).upper()}_{str(billing_cycle).upper()}_{txn_ref}"

    vnp_params = {
        'vnp_Version': '2.1.0',
        'vnp_Command': 'pay',
        'vnp_TmnCode': VNPAY_TMN_CODE,
        'vnp_Amount': str(amount * 100),
        'vnp_CurrCode': 'VND',
        'vnp_TxnRef': txn_ref,
        'vnp_OrderInfo': order_info,
        'vnp_OrderType': 'billpayment',
        'vnp_Locale': 'vn',
        'vnp_ReturnUrl': VNPAY_RETURN_URL,
        'vnp_IpAddr': get_client_ip(),
        'vnp_CreateDate': now.strftime('%Y%m%d%H%M%S'),
    }

    payment_url, _ = create_payment_url(VNPAY_PAYMENT_URL, vnp_params, VNPAY_HASH_SECRET)

    payment_doc = {
        'txn_ref': txn_ref,
        'user_email': session['user_email'],
        'plan': plan,
        'billing_cycle': billing_cycle,
        'amount': amount,
        'status': 'pending',
        'provider': 'vnpay',
        'payment_url': payment_url,
        'return_url': frontend_return_url,
        'vnpay_return_url': VNPAY_RETURN_URL,
        'created_ip': get_client_ip(),
        'created_at': now,
        'updated_at': now,
    }

    try:
        payments_collection.insert_one(payment_doc)
    except Exception as e:
        return jsonify({'success': False, 'error': f'Lỗi lưu giao dịch: {str(e)}'}), 500

    return jsonify({
        'success': True,
        'method': 'vnpay',
        'payment_url': payment_url,
        'txn_ref': txn_ref,
    })


@app.route('/api/payments/vnpay/return', methods=['GET'])
def vnpay_return():
    """VNPay callback endpoint: verify signature, update payment and user plan, then redirect frontend."""
    params = dict(request.args)
    is_valid_signature = verify_return_signature(params, VNPAY_HASH_SECRET)

    txn_ref = params.get('vnp_TxnRef', '')
    response_code = params.get('vnp_ResponseCode', '')
    transaction_status = params.get('vnp_TransactionStatus', '')
    is_success = is_valid_signature and response_code == '00' and transaction_status == '00'

    payment_doc = update_vnpay_payment(txn_ref, params, is_success)

    redirect_url = (
        f"{FRONTEND_RETURN_URL}"
        f"?status={'success' if is_success else 'failed'}"
        f"&txn_ref={txn_ref}"
        f"&code={response_code}"
    )
    if payment_doc and payment_doc.get('plan'):
        redirect_url += f"&plan={payment_doc.get('plan')}"

    return redirect(redirect_url)


@app.route('/api/payments/vnpay/ipn', methods=['GET', 'POST'])
def vnpay_ipn():
    """VNPay IPN endpoint for server-to-server payment confirmation."""
    params = dict(request.values)
    is_valid_signature = verify_return_signature(params, VNPAY_HASH_SECRET)

    txn_ref = params.get('vnp_TxnRef', '')
    response_code = params.get('vnp_ResponseCode', '')
    transaction_status = params.get('vnp_TransactionStatus', '')

    if not is_valid_signature:
        return jsonify({'RspCode': '97', 'Message': 'Invalid signature'}), 200

    payment_doc = None
    if MONGODB_CONNECTED and payments_collection is not None and txn_ref:
        payment_doc = payments_collection.find_one({'txn_ref': txn_ref})

    if not payment_doc:
        return jsonify({'RspCode': '01', 'Message': 'Order not found'}), 200

    if payment_doc.get('status') == 'success':
        return jsonify({'RspCode': '02', 'Message': 'Order already confirmed'}), 200

    is_success = response_code == '00' and transaction_status == '00'
    update_vnpay_payment(txn_ref, params, is_success)

    return jsonify({'RspCode': '00', 'Message': 'Confirm Success'}), 200


@app.route('/api/payments/history', methods=['GET'])
def payments_history():
    """Get current user's payment transactions."""
    if 'user_email' not in session:
        return jsonify({'success': False, 'error': 'Vui lòng đăng nhập'}), 401

    if not (MONGODB_CONNECTED and payments_collection is not None):
        return jsonify({'success': True, 'payments': []})

    try:
        cursor = payments_collection.find(
            {'user_email': session['user_email']}
        ).sort('created_at', -1).limit(30)
        payments = [serialize_payment(doc) for doc in cursor]
        return jsonify({'success': True, 'payments': payments})
    except Exception as e:
        return jsonify({'success': False, 'error': f'Không thể tải lịch sử thanh toán: {str(e)}'}), 500


@app.route('/api/payments/status/<txn_ref>', methods=['GET'])
def payment_status(txn_ref):
    """Get status of a transaction for current user."""
    if 'user_email' not in session:
        return jsonify({'success': False, 'error': 'Vui lòng đăng nhập'}), 401

    if not txn_ref:
        return jsonify({'success': False, 'error': 'Thiếu mã giao dịch'}), 400

    if not (MONGODB_CONNECTED and payments_collection is not None):
        return jsonify({'success': False, 'error': 'Database không khả dụng'}), 503

    payment = payments_collection.find_one({'txn_ref': txn_ref})
    if not payment:
        return jsonify({'success': False, 'error': 'Không tìm thấy giao dịch'}), 404

    if payment.get('user_email') != session['user_email']:
        return jsonify({'success': False, 'error': 'Không có quyền xem giao dịch này'}), 403

    return jsonify({'success': True, 'payment': serialize_payment(payment)})


@app.route('/api/generate-pdf/', methods=['POST', 'OPTIONS'])
def generate_pdf():
    if request.method == 'OPTIONS':
        response = jsonify({'status': 'ok'})
        response.headers.add('Access-Control-Allow-Origin', request.headers.get('Origin', '*'))
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        return response, 200

    try:
        from pdf_generator import generate_pdf_report
        import uuid

        data = request.get_json(silent=True)
        if data is None:
            raw_body = request.get_data(cache=False)
            if raw_body:
                try:
                    import json
                    data = json.loads(raw_body.decode('utf-8'))
                except Exception:
                    try:
                        import json
                        data = json.loads(raw_body.decode('latin-1'))
                    except Exception:
                        data = None

        if not data:
            return jsonify({
                'success': False,
                'error': 'Không có dữ liệu để tạo báo cáo'
            }), 400

        analysis_content = data.get('analysis', data.get('ai_analysis', ''))
        if not analysis_content:
            return jsonify({
                'success': False,
                'error': 'Thiếu nội dung phân tích (analysis field)'
            }), 400

        contract_name = data.get('filename', data.get('contract_name', 'Hop_Dong'))
        contract_name = contract_name.rsplit('.', 1)[0].replace(' ', '_')

        pdf_filename = f"Bao_Cao_{contract_name}_{uuid.uuid4().hex[:6]}.pdf"
        pdf_path = os.path.join(UPLOAD_FOLDER, pdf_filename)

        print(f"Generating PDF: {pdf_filename}")
        result_path = generate_pdf_report(analysis_content, pdf_path)

        if not result_path or not os.path.exists(result_path):
            return jsonify({
                'success': False,
                'error': 'Lỗi tạo PDF - File không được tạo'
            }), 500

        response = make_response(send_file(
            result_path,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"Bao_Cao_{contract_name}.pdf"
        ))
        response.headers['Access-Control-Allow-Origin'] = request.headers.get('Origin', 'http://localhost:3000')
        response.headers['Access-Control-Allow-Credentials'] = 'true'
        response.headers['Access-Control-Expose-Headers'] = 'Content-Disposition'

        return response

    except Exception as e:
        print(f"PDF Generation Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Lỗi tạo PDF: {str(e)}'
        }), 500

@app.route('/api/upload', methods=['POST'])
@rate_limit
def upload_file():
    """API upload và phân tích hợp đồng - với caching"""
    if 'user_email' not in session:
        return jsonify({'success': False, 'message': 'Vui lòng đăng nhập'}), 401
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'Không có file được tải lên'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'success': False, 'message': 'Chưa chọn file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Đọc nội dung file theo định dạng thực tế
        try:
            contract_text = extract_text_from_file(filepath)
        except Exception as e:
            return jsonify({
                'success': False,
                'message': f'Không đọc được nội dung file: {str(e)}'
            }), 500
        
        # Kiểm tra cache
        cache_key = get_cache_key(contract_text)
        now = time.time()
        
        if cache_key in ANALYSIS_CACHE:
            cached_data, cached_time = ANALYSIS_CACHE[cache_key]
            if now - cached_time < CACHE_TTL:
                print("✓ Using cached analysis result")
                # Cleanup file ngay sau khi đọc
                if AUTO_CLEANUP_UPLOADS:
                    try:
                        os.remove(filepath)
                    except:
                        pass
                return jsonify({
                    'success': True,
                    'data': cached_data,
                    'cached': True
                })
        
        # Gọi Agent để phân tích
        try:
            app_graph = get_app_graph()
            inputs = {"contract_text": contract_text}
            result = app_graph.invoke(inputs)
            
            # Extract SVM results for structured data
            svm_results = result.get('svm_results', {})
            contract_type = svm_results.get('contract_type', {}).get('predicted_type', 'Không xác định')
            risk_level = svm_results.get('risk_assessment', {}).get('predicted_risk', 'medium')
            has_violation = svm_results.get('violation_check', {}).get('has_violation', False)

            # Build structured issues from SVM and PageIndex results
            issues = build_analysis_issues(result)

            if not issues:
                issues = build_heuristic_issues(result)

            # Create summary and reasoning from structured data
            summary, safety_reasoning = summarize_analysis(result, issues)

            high_risk_count = sum(1 for item in issues if item.get('severity') == 'high')
            medium_risk_count = sum(1 for item in issues if item.get('severity') == 'medium')
            low_risk_count = sum(1 for item in issues if item.get('severity') == 'low')

            safety_score = max(0, min(100, 100 - (high_risk_count * 10) - (medium_risk_count * 5) - (low_risk_count * 2)))

            final_report = result.get('final_report', 'Không có kết quả phân tích')
            
            # Parse kết quả - FRONTEND FORMAT
            analysis_data = {
                'filename': filename,
                'upload_time': datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
                'contract_type': contract_type,
                'risk_level': risk_level,
                'has_violation': has_violation,
                'summary': summary,
                'ai_analysis': final_report,
                'issues': issues,
                'extractedClauses': result.get('extracted_clauses', [])[:10],
                'researchResults': result.get('research_results', [])[:5],
                'legal_references': [],  # Can be populated from research_results if needed
                'safety_score': safety_score,
                'safety_reasoning': safety_reasoning,
                'analysis_mode': 'model' if svm_results else 'heuristic',
                'issue_counts': {
                    'high': high_risk_count,
                    'medium': medium_risk_count,
                    'low': low_risk_count,
                    'total': len(issues),
                }
            }
            
            # Cache kết quả
            ANALYSIS_CACHE[cache_key] = (analysis_data, now)
            
            # Lưu vào MongoDB history
            history_doc = {
                'user': session['user_email'],
                'data': analysis_data,
                'timestamp': datetime.now(),
                'created_at': datetime.now()
            }
            
            if MONGODB_CONNECTED and analysis_collection is not None:
                try:
                    analysis_collection.insert_one(history_doc)
                    print("✅ Saved to MongoDB history")
                except Exception as e:
                    print(f"⚠️ Failed to save to MongoDB: {e}")
                    # Fallback to memory
                    ANALYSIS_HISTORY.append({
                        'id': len(ANALYSIS_HISTORY) + 1,
                        'user': session['user_email'],
                        'data': analysis_data,
                        'timestamp': datetime.now().isoformat()
                    })
            else:
                # Fallback: Lưu vào memory nếu MongoDB không khả dụng
                ANALYSIS_HISTORY.append({
                    'id': len(ANALYSIS_HISTORY) + 1,
                    'user': session['user_email'],
                    'data': analysis_data,
                    'timestamp': datetime.now().isoformat()
                })
                # Giới hạn history size
                if len(ANALYSIS_HISTORY) > 100:
                    ANALYSIS_HISTORY.pop(0)
            
            # Cleanup file ngay sau khi phân tích xong
            if AUTO_CLEANUP_UPLOADS:
                try:
                    os.remove(filepath)
                except:
                    pass
            
            return jsonify({
                'success': True,
                'data': analysis_data
            })
        except Exception as e:
            # Cleanup file khi có lỗi
            if AUTO_CLEANUP_UPLOADS:
                try:
                    os.remove(filepath)
                except:
                    pass
            return jsonify({
                'success': False,
                'message': f'Lỗi khi phân tích: {str(e)}'
            }), 500
    
    return jsonify({'success': False, 'message': 'File không hợp lệ'}), 400

@app.route('/api/history', methods=['GET'])
def get_history():
    """API lay lich su phan tich - from MongoDB"""
    if 'user_email' not in session:
        return jsonify({'success': False, 'message': 'Vui long dang nhap'}), 401
    
    user_email = session['user_email']
    
    # Lấy từ MongoDB nếu có kết nối
    if MONGODB_CONNECTED and analysis_collection is not None:
        try:
            # Query MongoDB
            cursor = analysis_collection.find(
                {'user': user_email}
            ).sort('timestamp', -1).limit(100)
            
            user_history = []
            for idx, doc in enumerate(cursor):
                # Convert MongoDB document to dict
                history_item = {
                    'id': idx + 1,
                    'user': doc['user'],
                    'data': doc['data'],
                    'timestamp': doc['timestamp'].isoformat() if isinstance(doc['timestamp'], datetime) else doc['timestamp']
                }
                user_history.append(history_item)
            
            print(f"✅ Loaded {len(user_history)} items from MongoDB")
            return jsonify({
                'success': True,
                'history': user_history
            })
        except Exception as e:
            print(f"⚠️ MongoDB query failed: {e}")
            # Fallback to memory
    
    # Fallback: Lấy từ memory nếu MongoDB không khả dụng
    user_history = [h for h in ANALYSIS_HISTORY if h['user'] == user_email]
    
    return jsonify({
        'success': True,
        'history': user_history
    })

@app.route('/api/admin/stats', methods=['GET'])
def admin_stats():
    """API thong ke cho admin"""
    if 'user_email' not in session or not session.get('is_admin'):
        return jsonify({'success': False, 'message': 'Khong co quyen truy cap'}), 403

    if not MONGODB_CONNECTED or users_collection is None or analysis_collection is None:
        return jsonify({'success': False, 'message': 'Database không khả dụng'}), 503

    total_users = users_collection.count_documents({})
    active_users = users_collection.count_documents({'is_active': True})
    admin_users = users_collection.count_documents({'is_admin': True})
    total_analyses = analysis_collection.count_documents({})
    total_contracts = contracts_collection.count_documents({}) if contracts_collection is not None else total_analyses
    recent_analyses = analysis_collection.count_documents({
        'timestamp': {'$gte': datetime.now() - timedelta(days=1)}
    })

    def month_index_from_value(value):
        if not value:
            return None
        if isinstance(value, str):
            try:
                value = datetime.fromisoformat(value.replace('Z', '+00:00'))
            except Exception:
                try:
                    value = datetime.strptime(value, '%d/%m/%Y %H:%M:%S')
                except Exception:
                    return None
        if not isinstance(value, datetime):
            return None
        return value.month - 1

    users_by_month = [0] * 12
    for doc in users_collection.find({}, {'created_at': 1, 'createdAt': 1}):
        month_index = month_index_from_value(doc.get('created_at') or doc.get('createdAt'))
        if month_index is not None:
            users_by_month[month_index] += 1

    analyses_by_month = [0] * 12
    for doc in analysis_collection.find({}, {'timestamp': 1, 'created_at': 1, 'createdAt': 1, 'data.upload_time': 1}):
        upload_time = ((doc.get('data') or {}).get('upload_time') if isinstance(doc.get('data'), dict) else None)
        month_index = month_index_from_value(doc.get('timestamp') or doc.get('created_at') or doc.get('createdAt') or upload_time)
        if month_index is not None:
            analyses_by_month[month_index] += 1

    payment_revenue = 0
    revenue_by_month = [0] * 12
    revenue_by_user = {}
    if payments_collection is not None:
        payment_pipeline = [
            {'$match': {'status': 'success'}},
            {'$group': {'_id': None, 'total': {'$sum': '$amount'}}},
        ]
        revenue_cursor = list(payments_collection.aggregate(payment_pipeline))
        if revenue_cursor:
            payment_revenue = revenue_cursor[0].get('total', 0) or 0

        monthly_pipeline = [
            {'$match': {'status': 'success', 'created_at': {'$type': 'date'}}},
            {
                '$group': {
                    '_id': {'$month': '$created_at'},
                    'total': {'$sum': '$amount'}
                }
            }
        ]
        for item in payments_collection.aggregate(monthly_pipeline):
            month_index = int(item.get('_id', 1)) - 1
            if 0 <= month_index < 12:
                revenue_by_month[month_index] = item.get('total', 0) or 0

        revenue_per_user_pipeline = [
            {'$match': {'status': 'success'}},
            {
                '$group': {
                    '_id': '$user_email',
                    'revenue': {'$sum': '$amount'}
                }
            },
            {'$sort': {'revenue': -1}},
            {'$limit': 5},
        ]
        for item in payments_collection.aggregate(revenue_per_user_pipeline):
            user_email = item.get('_id') or 'unknown'
            revenue_by_user[user_email] = item.get('revenue', 0) or 0

    analysis_counts_by_user = {}
    for item in analysis_collection.aggregate([
        {'$group': {'_id': '$user', 'analyses': {'$sum': 1}}},
        {'$sort': {'analyses': -1}},
        {'$limit': 10},
    ]):
        analysis_counts_by_user[item.get('_id') or 'unknown'] = int(item.get('analyses', 0) or 0)

    latest_analyses = []
    for doc in analysis_collection.find({}).sort('timestamp', -1).limit(5):
        latest_analyses.append(serialize_analysis_summary(doc))

    latest_users = []
    for doc in users_collection.find({}).sort('created_at', -1).limit(5):
        user_payload = serialize_user(doc)
        user_payload['analysisCount'] = analysis_counts_by_user.get(user_payload['email'], user_payload.get('analysisCount', 0))
        user_payload['revenue'] = f"{int(revenue_by_user.get(user_payload['email'], 0) or 0):,} VNĐ".replace(',', '.')
        latest_users.append(user_payload)

    top_users = sorted(
        [
            {
                'email': doc.get('email', ''),
                'name': doc.get('full_name') or doc.get('username') or doc.get('email', ''),
                'analyses': analysis_counts_by_user.get(doc.get('email', ''), 0),
                'revenue': int(revenue_by_user.get(doc.get('email', ''), 0) or 0),
            }
            for doc in users_collection.find({}, {'email': 1, 'full_name': 1, 'username': 1})
        ],
        key=lambda item: item['analyses'],
        reverse=True,
    )[:5]

    stats = {
        'total_users': total_users,
        'total_analyses': total_analyses,
        'total_contracts': total_contracts,
        'active_users': active_users,
        'admin_users': admin_users,
        'recent_analyses': recent_analyses,
        'monthly_revenue': payment_revenue,
        'revenue_by_month': revenue_by_month,
        'users_by_month': users_by_month,
        'analyses_by_month': analyses_by_month,
        'top_users': top_users,
        'success_rate': round((total_analyses / total_contracts) * 100, 1) if total_contracts else 0,
    }

    return jsonify({
        'success': True,
        'stats': stats,
        'recent_activities': [
            serialize_activity({
                'user': item['user'],
                'action': 'Phân tích hợp đồng',
                'time': item['date'],
                'type': 'analysis'
            })
            for item in latest_analyses
        ],
        'latest_users': latest_users,
        'latest_analyses': latest_analyses,
    })


@app.route('/api/admin/users', methods=['GET'])
def admin_users():
    if 'user_email' not in session or not session.get('is_admin'):
        return jsonify({'success': False, 'message': 'Khong co quyen truy cap'}), 403

    if not MONGODB_CONNECTED or users_collection is None:
        return jsonify({'success': False, 'message': 'Database không khả dụng'}), 503

    users = [serialize_user(doc) for doc in users_collection.find({}).sort('created_at', -1)]
    return jsonify({'success': True, 'users': users})


@app.route('/api/admin/users/<user_id>', methods=['PUT', 'DELETE', 'OPTIONS'])
def admin_mutate_user(user_id):
    if request.method == 'OPTIONS':
        return jsonify({'success': True}), 200

    if 'user_email' not in session or not session.get('is_admin'):
        return jsonify({'success': False, 'message': 'Khong co quyen truy cap'}), 403

    if not MONGODB_CONNECTED or users_collection is None:
        return jsonify({'success': False, 'message': 'Database không khả dụng'}), 503

    try:
        from bson import ObjectId
    except Exception:
        return jsonify({'success': False, 'message': 'Thiếu ObjectId support'}), 500

    user_query = {'_id': user_id}
    try:
        object_id = ObjectId(user_id)
        user_query = {'$or': [{'_id': object_id}, {'_id': user_id}]}
    except Exception:
        pass

    if request.method == 'PUT':
        data = request.get_json(silent=True) or {}
        plan = (data.get('plan') or '').strip().lower()
        is_active = data.get('is_active')

        update_payload = {}
        if plan in ('free', 'pro'):
            update_payload['plan'] = plan
            update_payload['subscription_tier'] = plan

        if isinstance(is_active, bool):
            update_payload['is_active'] = is_active

        if not update_payload:
            return jsonify({'success': False, 'message': 'Không có dữ liệu cập nhật hợp lệ'}), 400

        result = users_collection.update_one(
            user_query,
            {'$set': {**update_payload, 'updated_at': datetime.now()}},
        )

        if result.matched_count == 0:
            return jsonify({'success': False, 'message': 'Không tìm thấy người dùng'}), 404

        updated_user = users_collection.find_one(user_query)
        return jsonify({'success': True, 'user': serialize_user(updated_user)})

    user_doc = users_collection.find_one(user_query)
    if not user_doc:
        return jsonify({'success': False, 'message': 'Không tìm thấy người dùng'}), 404

    if user_doc.get('email') == session.get('user_email'):
        return jsonify({'success': False, 'message': 'Không thể xóa tài khoản đang đăng nhập'}), 400

    if user_doc.get('is_admin'):
        return jsonify({'success': False, 'message': 'Không thể xóa tài khoản admin'}), 400

    users_collection.delete_one({'_id': user_doc.get('_id')})
    return jsonify({'success': True, 'message': 'Đã xóa người dùng'})


@app.route('/api/admin/analyses', methods=['GET'])
def admin_analyses():
    if 'user_email' not in session or not session.get('is_admin'):
        return jsonify({'success': False, 'message': 'Khong co quyen truy cap'}), 403

    if not MONGODB_CONNECTED or analysis_collection is None:
        return jsonify({'success': False, 'message': 'Database không khả dụng'}), 503

    analyses = [serialize_analysis_summary(doc) for doc in analysis_collection.find({}).sort('timestamp', -1).limit(100)]
    return jsonify({'success': True, 'analyses': analyses})

@app.route('/history')
def history():
    """History page"""
    if 'user_email' not in session:
        return redirect('/')
    return render_template('history.html')

@app.route('/settings')
def settings():
    """Trang cài đặt tài khoản"""
    if 'user_email' not in session:
        return redirect('/')
    return render_template('settings.html')

@app.route('/pricing')
def pricing():
    """Trang bảng giá"""
    return render_template('pricing.html')

@app.route('/admin')
def admin():
    """Trang admin dashboard"""
    if 'user_email' not in session or not session.get('is_admin'):
        return redirect('/')
    return render_template('admin.html')

# Cleanup tasks
@atexit.register
def shutdown():
    """Cleanup khi app shutdown"""
    print("🧹 Cleaning up resources...")
    cleanup_old_files()
    # Clear caches
    ANALYSIS_CACHE.clear()
    RATE_LIMIT_STORE.clear()

# Templates seeding
TEMPLATES_CACHE = {}

def seed_templates():
    """Load templates from JSON file and cache them."""
    seed_file = os.path.join(os.path.dirname(__file__), 'src', 'templates_seed.json')
    if not os.path.exists(seed_file):
        print(f"[WARN] Template seed file not found: {seed_file}")
        return
    
    try:
        with open(seed_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        for template in data.get('templates', []):
            template_id = template.get('template_id')
            TEMPLATES_CACHE[template_id] = template
        
        print(f"[OK] Loaded {len(TEMPLATES_CACHE)} templates from seed file")
    except Exception as e:
        print(f"[WARN] Failed to load templates: {e}")


@app.route('/v1/templates', methods=['GET', 'OPTIONS'])
def list_templates():
    """List all templates"""
    if request.method == 'OPTIONS':
        return '', 200
    
    templates = []
    for template_id, template in TEMPLATES_CACHE.items():
        templates.append({
            'template_id': template.get('template_id'),
            'name': template.get('name'),
            'category': template.get('category'),
            'description': template.get('description'),
            'variables': template.get('variables', {})
        })
    
    return jsonify({
        'success': True,
        'templates': sorted(templates, key=lambda x: x.get('name', ''))
    })


@app.route('/v1/templates/<template_id>', methods=['GET', 'OPTIONS'])
def get_template(template_id):
    """Get template details"""
    if request.method == 'OPTIONS':
        return '', 200
    
    template = TEMPLATES_CACHE.get(template_id)
    if not template:
        return jsonify({'success': False, 'error': 'not_found'}), 404
    
    return jsonify({
        'success': True,
        'template': {
            'template_id': template.get('template_id'),
            'name': template.get('name'),
            'category': template.get('category'),
            'description': template.get('description'),
            'template_content': json.dumps(template, ensure_ascii=False),
            'variables': template.get('variables', {})
        }
    })


def _docx_blank_line(length=72):
    return '.' * length


def _docx_add_center_paragraph(doc, text, *, bold=False, size=12):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return paragraph


def _docx_add_text_paragraph(doc, text, *, bold=False, italic=False, size=12):
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return paragraph


def _docx_add_field_block(doc, label, field, index=None):
    field_type = (field.get('type') or 'text').lower()
    required = ' (*)' if field.get('required') else ''
    prefix = f"{index}. " if index is not None else ''
    base_label = f"{prefix}{label}{required}:"

    if field_type == 'textarea':
        _docx_add_text_paragraph(doc, base_label, bold=True, size=12)
        for _ in range(4):
            _docx_add_text_paragraph(doc, _docx_blank_line(92), size=11)
        return

    if field_type == 'date':
        blank = '..... / ..... / ........'
    elif field_type == 'number':
        blank = _docx_blank_line(42)
    elif field_type == 'select' and field.get('options'):
        options = ' / '.join(str(option) for option in field.get('options', []))
        blank = f"{_docx_blank_line(18)} ({options})"
    else:
        blank = _docx_blank_line(72)

    _docx_add_text_paragraph(doc, f"{base_label} {blank}", size=12)


def _docx_add_party_section(doc, heading, fields):
    _docx_add_text_paragraph(doc, heading, bold=True, size=13)
    for item in fields:
        label = item[0] if isinstance(item, tuple) else str(item)
        _docx_add_text_paragraph(doc, f"{label}: {_docx_blank_line(70)}", size=12)


def _docx_add_numbered_heading(doc, number, heading):
    _docx_add_text_paragraph(doc, f"ĐIỀU {number}: {heading}", bold=True, size=13)


def _docx_add_clause_line(doc, number, text, indent=0):
    prefix = f"{number} - " if number else ''
    spaces = ' ' * indent if indent else ''
    _docx_add_text_paragraph(doc, f"{spaces}{prefix}{text}", size=12)


def _docx_add_blank_lines(doc, count=2, width=92):
    for _ in range(count):
        _docx_add_text_paragraph(doc, _docx_blank_line(width), size=11)


def _docx_get_party_layout(template):
    template_id = template.get('template_id')
    variables = template.get('variables') or {}

    if template_id == 'hop_dong_thue':
        return [
            ('BÊN CHO THUÊ MẶT BẰNG', [
                ('Họ tên', variables.get('landlord', {})),
                ('CMND/CCCD số', {}),
                ('Thường trú', {}),
                ('Là chủ sở hữu căn nhà/mặt bằng', {}),
            ]),
            ('BÊN THUÊ MẶT BẰNG', [
                ('Họ tên', variables.get('tenant', {})),
                ('CMND/CCCD số', {}),
                ('Thường trú', {}),
            ]),
        ]

    if template_id == 'hop_dong_lao_dong':
        return [
            ('BÊN SỬ DỤNG LAO ĐỘNG', [
                ('Tên công ty', variables.get('company_name', {})),
                ('Địa chỉ', {}),
                ('Đại diện', {}),
                ('Chức vụ', {}),
            ]),
            ('NGƯỜI LAO ĐỘNG', [
                ('Họ tên', variables.get('employee_name', {})),
                ('Ngày sinh', {}),
                ('Số CCCD/CMND', {}),
                ('Địa chỉ thường trú', {}),
            ]),
        ]

    if template_id == 'hop_dong_dich_vu':
        return [
            ('BÊN CUNG CẤP DỊCH VỤ', [
                ('Tên đơn vị', variables.get('provider_name', {})),
                ('Địa chỉ', {}),
                ('Người đại diện', {}),
            ]),
            ('BÊN SỬ DỤNG DỊCH VỤ', [
                ('Tên đơn vị', variables.get('client_name', {})),
                ('Địa chỉ', {}),
                ('Người đại diện', {}),
            ]),
        ]

    return [
        ('BÊN A', [('Thông tin', {})]),
        ('BÊN B', [('Thông tin', {})]),
    ]


def _docx_build_contract_body(doc, template):
    template_id = template.get('template_id')
    variables = template.get('variables') or {}
    if not isinstance(variables, dict):
        variables = {}

    if template_id == 'hop_dong_thue':
        _docx_add_numbered_heading(doc, 1, 'NỘI DUNG HỢP ĐỒNG')
        _docx_add_clause_line(doc, '1.1', f"Bên A đồng ý cho bên B thuê mặt bằng tại: {_docx_blank_line(52)}")
        _docx_add_blank_lines(doc, 2)
        _docx_add_clause_line(doc, '1.2', f"Diện tích và hiện trạng mặt bằng: {_docx_blank_line(44)}")
        _docx_add_blank_lines(doc, 2)
        _docx_add_clause_line(doc, '1.3', f"Mục đích thuê: {_docx_blank_line(62)}")

        _docx_add_numbered_heading(doc, 2, 'THỜI HẠN HỢP ĐỒNG')
        _docx_add_clause_line(doc, '2.1', f"Thời hạn thuê mặt bằng là: {_docx_blank_line(32)} tháng, từ ngày {_docx_blank_line(18)} đến hết ngày {_docx_blank_line(18)}.")
        _docx_add_clause_line(doc, '2.2', 'Sau khi hết hạn hợp đồng, hai bên có thể thỏa thuận gia hạn hoặc chấm dứt hợp đồng.')

        _docx_add_numbered_heading(doc, 3, 'GIÁ CẢ - PHƯƠNG THỨC THANH TOÁN')
        _docx_add_clause_line(doc, '3.1', f"Giá thuê mặt bằng là: {_docx_blank_line(58)}")
        _docx_add_blank_lines(doc, 2)
        _docx_add_clause_line(doc, '3.2', 'Việc thanh toán được thực hiện theo định kỳ hai bên đã thỏa thuận.')
        _docx_add_clause_line(doc, '3.3', 'Các khoản chi phí phát sinh khác sẽ do hai bên thống nhất bằng văn bản.')

        _docx_add_numbered_heading(doc, 4, 'TRÁCH NHIỆM CỦA HAI BÊN')
        _docx_add_clause_line(doc, '4.1', 'Trách nhiệm của bên A:')
        _docx_add_clause_line(doc, '4.1.1', 'Bên A bảo đảm quyền sử dụng hợp pháp mặt bằng cho bên B.')
        _docx_add_clause_line(doc, '4.1.2', 'Bên A bàn giao mặt bằng, trang thiết bị theo thỏa thuận.')
        _docx_add_clause_line(doc, '4.2', 'Trách nhiệm của bên B:')
        _docx_add_clause_line(doc, '4.2.1', 'Sử dụng mặt bằng đúng mục đích thuê, giữ gìn tài sản và tuân thủ quy định pháp luật.')
        _docx_add_clause_line(doc, '4.2.2', 'Thanh toán tiền thuê và các chi phí liên quan đúng thời hạn.')
        _docx_add_clause_line(doc, '4.2.3', 'Khi chấm dứt hợp đồng, giao trả lại mặt bằng theo hiện trạng ban đầu.')

        _docx_add_numbered_heading(doc, 5, 'CAM KẾT CHUNG')
        _docx_add_clause_line(doc, '', 'Hai bên cam kết thực hiện đúng các điều khoản đã ghi trong hợp đồng.')
        _docx_add_clause_line(doc, '', 'Nếu phát sinh tranh chấp, hai bên ưu tiên thương lượng; nếu không giải quyết được thì đưa ra Tòa án có thẩm quyền.')
    elif template_id == 'hop_dong_lao_dong':
        _docx_add_numbered_heading(doc, 1, 'CÔNG VIỆC VÀ ĐỊA ĐIỂM LÀM VIỆC')
        _docx_add_clause_line(doc, '1.1', f"Vị trí công việc: {_docx_blank_line(58)}")
        _docx_add_clause_line(doc, '1.2', f"Địa điểm làm việc: {_docx_blank_line(58)}")
        _docx_add_clause_line(doc, '1.3', f"Thời gian làm việc: {_docx_blank_line(52)}")

        _docx_add_numbered_heading(doc, 2, 'THỜI HẠN HỢP ĐỒNG')
        _docx_add_clause_line(doc, '2.1', f"Loại hợp đồng: {_docx_blank_line(20)} (xác định thời hạn / không xác định thời hạn)")
        _docx_add_clause_line(doc, '2.2', f"Thời hạn (nếu có): {_docx_blank_line(48)} tháng")

        _docx_add_numbered_heading(doc, 3, 'TIỀN LƯƠNG VÀ CHẾ ĐỘ')
        _docx_add_clause_line(doc, '3.1', f"Mức lương: {_docx_blank_line(60)} VND/tháng")
        _docx_add_clause_line(doc, '3.2', 'Hình thức trả lương và các khoản phụ cấp được thỏa thuận cụ thể giữa hai bên.')

        _docx_add_numbered_heading(doc, 4, 'TRÁCH NHIỆM CỦA CÁC BÊN')
        _docx_add_clause_line(doc, '4.1', 'Bên sử dụng lao động bảo đảm điều kiện làm việc, trả lương và thực hiện các nghĩa vụ theo pháp luật.')
        _docx_add_clause_line(doc, '4.2', 'Người lao động thực hiện đúng công việc được giao, tuân thủ nội quy và bảo mật thông tin.')

        _docx_add_numbered_heading(doc, 5, 'CAM KẾT CHUNG')
        _docx_add_clause_line(doc, '', 'Hai bên cam kết tuân thủ đầy đủ các quy định của Bộ luật Lao động và các thỏa thuận trong hợp đồng.')
    elif template_id == 'hop_dong_dich_vu':
        _docx_add_numbered_heading(doc, 1, 'PHẠM VI DỊCH VỤ')
        _docx_add_clause_line(doc, '1.1', f"Mô tả dịch vụ: {_docx_blank_line(58)}")
        _docx_add_clause_line(doc, '1.2', f"Thời hạn thực hiện: {_docx_blank_line(54)}")

        _docx_add_numbered_heading(doc, 2, 'GIÁ DỊCH VỤ VÀ THANH TOÁN')
        _docx_add_clause_line(doc, '2.1', f"Phí dịch vụ: {_docx_blank_line(58)} VND")
        _docx_add_clause_line(doc, '2.2', 'Phương thức thanh toán: chuyển khoản/tiền mặt theo thỏa thuận.')

        _docx_add_numbered_heading(doc, 3, 'TRÁCH NHIỆM CỦA HAI BÊN')
        _docx_add_clause_line(doc, '3.1', 'Bên cung cấp dịch vụ thực hiện đúng phạm vi và tiến độ đã thỏa thuận.')
        _docx_add_clause_line(doc, '3.2', 'Bên sử dụng dịch vụ thanh toán đầy đủ, đúng hạn và phối hợp cung cấp thông tin cần thiết.')

        _docx_add_numbered_heading(doc, 4, 'CAM KẾT CHUNG')
        _docx_add_clause_line(doc, '', 'Hai bên thống nhất bảo mật thông tin và giải quyết tranh chấp bằng thương lượng trước khi đưa ra cơ quan có thẩm quyền.')
    elif template_id == 'hop_dong_mua_ban':
        _docx_add_numbered_heading(doc, 1, 'HÀNG HÓA - SỐ LƯỢNG - CHẤT LƯỢNG')
        _docx_add_clause_line(doc, '1.1', f"Mô tả hàng hóa: {_docx_blank_line(56)}")
        _docx_add_clause_line(doc, '1.2', f"Số lượng: {_docx_blank_line(66)}")
        _docx_add_clause_line(doc, '1.3', f"Đơn giá: {_docx_blank_line(70)} VND")

        _docx_add_numbered_heading(doc, 2, 'GIAO NHẬN HÀNG HÓA')
        _docx_add_clause_line(doc, '2.1', f"Địa điểm giao hàng: {_docx_blank_line(52)}")
        _docx_add_clause_line(doc, '2.2', f"Thời gian giao hàng: {_docx_blank_line(52)}")

        _docx_add_numbered_heading(doc, 3, 'THANH TOÁN')
        _docx_add_clause_line(doc, '3.1', f"Tổng giá trị hợp đồng: {_docx_blank_line(44)}")
        _docx_add_clause_line(doc, '3.2', 'Phương thức và thời hạn thanh toán theo thỏa thuận của hai bên.')

        _docx_add_numbered_heading(doc, 4, 'CAM KẾT CHUNG')
        _docx_add_clause_line(doc, '', 'Hai bên cam kết thực hiện đúng hợp đồng và chịu trách nhiệm trước pháp luật nếu vi phạm.')
    elif template_id == 'bien_ban_hop':
        _docx_add_numbered_heading(doc, 1, 'THÔNG TIN CUỘC HỌP')
        _docx_add_clause_line(doc, '1.1', f"Tiêu đề cuộc họp: {_docx_blank_line(54)}")
        _docx_add_clause_line(doc, '1.2', f"Thời gian: {_docx_blank_line(64)}")
        _docx_add_clause_line(doc, '1.3', f"Địa điểm: {_docx_blank_line(64)}")

        _docx_add_numbered_heading(doc, 2, 'THÀNH PHẦN THAM DỰ')
        _docx_add_clause_line(doc, '2.1', 'Danh sách người tham dự:')
        _docx_add_blank_lines(doc, 4)

        _docx_add_numbered_heading(doc, 3, 'NỘI DUNG CUỘC HỌP')
        _docx_add_clause_line(doc, '3.1', f"Các nội dung trao đổi: {_docx_blank_line(50)}")
        _docx_add_blank_lines(doc, 4)

        _docx_add_numbered_heading(doc, 4, 'KẾT LUẬN - THỐNG NHẤT')
        _docx_add_clause_line(doc, '', 'Các bên thống nhất các nội dung sau:')
        _docx_add_blank_lines(doc, 4)
    elif template_id == 'cong_van':
        _docx_add_numbered_heading(doc, 1, 'THÔNG TIN CÔNG VĂN')
        _docx_add_clause_line(doc, '1.1', f"Đơn vị gửi: {_docx_blank_line(58)}")
        _docx_add_clause_line(doc, '1.2', f"Đơn vị nhận: {_docx_blank_line(58)}")
        _docx_add_clause_line(doc, '1.3', f"Trích yếu: {_docx_blank_line(60)}")

        _docx_add_numbered_heading(doc, 2, 'NỘI DUNG')
        _docx_add_clause_line(doc, '2.1', 'Nội dung công văn:')
        _docx_add_blank_lines(doc, 5)

        _docx_add_numbered_heading(doc, 3, 'KẾT THÚC')
        _docx_add_clause_line(doc, '', 'Kính đề nghị Quý đơn vị xem xét, phối hợp và phản hồi trong thời gian sớm nhất.')
    elif template_id == 'noi_quy':
        _docx_add_numbered_heading(doc, 1, 'PHẠM VI ÁP DỤNG')
        _docx_add_clause_line(doc, '1.1', f"Tên công ty: {_docx_blank_line(60)}")
        _docx_add_clause_line(doc, '1.2', 'Nội quy này áp dụng đối với toàn bộ người lao động và các bộ phận liên quan.')

        _docx_add_numbered_heading(doc, 2, 'THỜI GIỜ LÀM VIỆC - NGHỈ NGƠI')
        _docx_add_clause_line(doc, '2.1', f"Giờ làm việc: {_docx_blank_line(60)}")
        _docx_add_clause_line(doc, '2.2', f"Chính sách nghỉ phép: {_docx_blank_line(52)}")

        _docx_add_numbered_heading(doc, 3, 'KỶ LUẬT LAO ĐỘNG')
        _docx_add_clause_line(doc, '3.1', 'Người lao động phải tuân thủ nội quy, giữ gìn tài sản và bảo mật thông tin của công ty.')
        _docx_add_clause_line(doc, '3.2', 'Mọi hành vi vi phạm sẽ bị xử lý theo quy định nội bộ và pháp luật hiện hành.')

        _docx_add_numbered_heading(doc, 4, 'CAM KẾT CHUNG')
        _docx_add_clause_line(doc, '', 'Nội quy có hiệu lực kể từ ngày ký và là cơ sở áp dụng thống nhất trong công ty.')
    elif template_id == 'quyet_dinh':
        _docx_add_numbered_heading(doc, 1, 'CĂN CỨ BAN HÀNH')
        _docx_add_clause_line(doc, '', 'Căn cứ vào điều lệ, quy chế tổ chức và nhu cầu thực tế của đơn vị.')
        _docx_add_blank_lines(doc, 2)

        _docx_add_numbered_heading(doc, 2, 'NỘI DUNG QUYẾT ĐỊNH')
        _docx_add_clause_line(doc, '2.1', f"Số quyết định: {_docx_blank_line(56)}")
        _docx_add_clause_line(doc, '2.2', f"Nội dung quyết định: {_docx_blank_line(50)}")
        _docx_add_blank_lines(doc, 4)

        _docx_add_numbered_heading(doc, 3, 'HIỆU LỰC THI HÀNH')
        _docx_add_clause_line(doc, '3.1', f"Quyết định có hiệu lực kể từ ngày {_docx_blank_line(38)}")
        _docx_add_clause_line(doc, '3.2', 'Các bộ phận, cá nhân liên quan chịu trách nhiệm thi hành quyết định này.')
    elif template_id == 'thong_bao':
        _docx_add_numbered_heading(doc, 1, 'NỘI DUNG THÔNG BÁO')
        _docx_add_clause_line(doc, '1.1', f"Tên công ty: {_docx_blank_line(60)}")
        _docx_add_clause_line(doc, '1.2', f"Ngày hiệu lực: {_docx_blank_line(58)}")
        _docx_add_clause_line(doc, '1.3', f"Nội dung thông báo: {_docx_blank_line(50)}")
        _docx_add_blank_lines(doc, 4)

        _docx_add_numbered_heading(doc, 2, 'ĐỐI TƯỢNG ÁP DỤNG')
        _docx_add_clause_line(doc, '', 'Thông báo này áp dụng đối với toàn bộ cá nhân, bộ phận hoặc đơn vị được nêu trong văn bản.')
        _docx_add_blank_lines(doc, 2)
    elif template_id == 'nda':
        _docx_add_numbered_heading(doc, 1, 'PHẠM VI THÔNG TIN BẢO MẬT')
        _docx_add_clause_line(doc, '1.1', f"Phạm vi thông tin: {_docx_blank_line(52)}")
        _docx_add_clause_line(doc, '1.2', f"Thời hạn bảo mật: {_docx_blank_line(48)} năm")

        _docx_add_numbered_heading(doc, 2, 'NGHĨA VỤ CỦA CÁC BÊN')
        _docx_add_clause_line(doc, '2.1', 'Bên nhận thông tin không được tiết lộ, sao chép hoặc sử dụng trái phép thông tin bảo mật.')
        _docx_add_clause_line(doc, '2.2', 'Bên công khai thông tin có quyền yêu cầu hoàn trả, hủy bỏ hoặc ngừng sử dụng thông tin khi cần thiết.')

        _docx_add_numbered_heading(doc, 3, 'XỬ LÝ VI PHẠM')
        _docx_add_clause_line(doc, '', 'Nếu vi phạm, bên vi phạm phải bồi thường thiệt hại và chịu trách nhiệm theo quy định pháp luật.')
    else:
        _docx_add_numbered_heading(doc, 1, 'NỘI DUNG')
        for index, (key, field) in enumerate(variables.items(), start=1):
            label = field.get('label') or key.replace('_', ' ').title()
            _docx_add_field_block(doc, label, field, index=index)

        _docx_add_numbered_heading(doc, 2, 'NỘI DUNG THỎA THUẬN')
        _docx_add_clause_line(doc, '2.1', 'Các bên thống nhất ký kết văn bản này trên cơ sở tự nguyện, trung thực và đúng quy định pháp luật.')
        _docx_add_blank_lines(doc, 2)

        _docx_add_numbered_heading(doc, 3, 'THỜI HẠN - HIỆU LỰC')
        _docx_add_clause_line(doc, '3.1', f"Thời hạn áp dụng: {_docx_blank_line(54)}")
        _docx_add_clause_line(doc, '3.2', f"Ngày hiệu lực: {_docx_blank_line(58)}")

        _docx_add_numbered_heading(doc, 4, 'CAM KẾT CHUNG')
        _docx_add_clause_line(doc, '', 'Hai bên cam kết thực hiện đúng các điều khoản đã ghi trong văn bản.')
        _docx_add_blank_lines(doc, 2)


def _docx_add_signature_block(doc):
    _docx_add_text_paragraph(doc, 'ĐẠI DIỆN BÊN A', bold=True, size=12)
    _docx_add_text_paragraph(doc, 'Ký, ghi rõ họ tên', size=11)
    _docx_add_blank_lines(doc, 3, 36)
    _docx_add_text_paragraph(doc, 'ĐẠI DIỆN BÊN B', bold=True, size=12)
    _docx_add_text_paragraph(doc, 'Ký, ghi rõ họ tên', size=11)
    _docx_add_blank_lines(doc, 3, 36)


@app.route('/v1/templates/<template_id>/download', methods=['GET'])
def download_template_docx(template_id):
    """Generate a contract-style .docx for the given template and return it as attachment."""
    try:
        template = TEMPLATES_CACHE.get(template_id)
        if not template:
            return jsonify({'success': False, 'error': 'not_found'}), 404

        try:
            from io import BytesIO
            from docx import Document as DocxDocument
            from docx.shared import Inches
        except Exception as e:
            return jsonify({'success': False, 'error': f'missing_docx_lib: {e}'}), 500

        doc = DocxDocument()
        section = doc.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        # Vietnamese contract header
        _docx_add_center_paragraph(doc, 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', bold=True, size=13)
        _docx_add_center_paragraph(doc, 'Độc lập - Tự do - Hạnh phúc', bold=True, size=12)
        _docx_add_center_paragraph(doc, '--------------------------------', size=12)
        doc.add_paragraph()

        title = str(template.get('name', 'VĂN BẢN')).upper()
        _docx_add_center_paragraph(doc, title, bold=True, size=18)
        _docx_add_center_paragraph(doc, 'Số: .../...', size=12)
        doc.add_paragraph()
        _docx_add_text_paragraph(doc, 'Hôm nay, ngày ... tháng ... năm ... tại ............................................................', size=12)
        _docx_add_text_paragraph(doc, 'Chúng tôi gồm:', bold=True, size=12)
        doc.add_paragraph()

        for heading, fields in _docx_get_party_layout(template):
            _docx_add_party_section(doc, heading, fields)
            doc.add_paragraph()

        _docx_add_text_paragraph(doc, 'Hai bên thoả thuận ký kết văn bản với nội dung sau:', bold=True, size=12)
        doc.add_paragraph()

        _docx_build_contract_body(doc, template)
        doc.add_paragraph()

        _docx_add_text_paragraph(doc, 'ĐIỀU 5: CAM KẾT CHUNG', bold=True, size=13)
        _docx_add_text_paragraph(doc, 'Hai bên cam kết thực hiện đúng các điều khoản đã ghi trong văn bản. Nếu có tranh chấp phát sinh thì sẽ ưu tiên giải quyết bằng thương lượng; nếu không đạt kết quả thì đưa ra cơ quan có thẩm quyền giải quyết.', size=12)
        _docx_add_text_paragraph(doc, 'Văn bản được lập thành 02 bản có giá trị pháp lý như nhau, mỗi bên giữ 01 bản.', size=12)

        doc.add_paragraph()
        _docx_add_signature_block(doc)

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)

        return send_file(
            bio,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{template_id}.docx"
        )
    except Exception as e:
        print(f"[ERROR] generate docx: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    # Seed templates on startup
    seed_templates()

    # Cleanup old files khi khởi động
    cleanup_old_files()
    app.run(debug=True, use_reloader=False, host='0.0.0.0', port=5000)
